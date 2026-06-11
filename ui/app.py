# ═══════════════════════════════════════════════════════════════════════════════
# Caledonian HR Group — Scottish Employment Contract Intelligence Platform
# Senior Staff Engineer refactor + 10 new premium features
# ═══════════════════════════════════════════════════════════════════════════════
import os
import re
import time
import json
import difflib
import hashlib
import requests
import streamlit as st
import fitz
import plotly.graph_objects as go
import plotly.express as px
from fpdf import FPDF
from datetime import datetime
import io

# ── python-docx for Legal Memo DOCX export (graceful fallback) ────────────────
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ── RAG / FAISS / LangChain imports (graceful fallback if not installed) ──────
try:
    import numpy as np
    from sentence_transformers import SentenceTransformer
    import faiss
    RAG_AVAILABLE = True
except ImportError:
    RAG_AVAILABLE = False

try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Caledonian HR Group · Scottish Employment Intelligence",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Backend ───────────────────────────────────────────────────────────────────
API_URL = os.getenv("API_URL", "https://legal-intelligence-engine-production.up.railway.app")

# ─────────────────────────────────────────────────────────────────────────────
# GLOBAL CSS — Dark luxury Scottish enterprise design
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@400;500;600;700&family=Inter:wght@300;400;500;600;700&family=Playfair+Display:wght@600;700&display=swap');

/* === Global === */
html, body, [class*="css"] { font-family: 'Inter', sans-serif !important; }
.stApp { background: #0B1825 !important; }
.block-container { padding: 1.5rem 2rem 2rem 2rem !important; max-width: 1400px !important; }

/* === Sidebar === */
section[data-testid="stSidebar"] > div {
    background: #071219 !important;
    border-right: 1px solid #1E3448 !important;
    padding-top: 0 !important;
}

/* === Buttons === */
.stButton > button {
    background: linear-gradient(135deg, #C9A84C 0%, #A8862A 100%) !important;
    color: #071219 !important; font-weight: 700 !important;
    font-size: 0.85rem !important; letter-spacing: 0.07em !important;
    text-transform: uppercase !important; border: none !important;
    border-radius: 6px !important; padding: 0.6rem 1.8rem !important;
    transition: all 0.2s ease !important;
    box-shadow: 0 4px 14px rgba(201,168,76,0.25) !important;
}
.stButton > button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 20px rgba(201,168,76,0.4) !important;
}

/* === Download button === */
.stDownloadButton > button {
    background: transparent !important; border: 1px solid #C9A84C !important;
    color: #C9A84C !important; font-size: 0.8rem !important;
    border-radius: 6px !important; letter-spacing: 0.05em !important;
    padding: 0.5rem 1.2rem !important;
}
.stDownloadButton > button:hover { background: rgba(201,168,76,0.1) !important; }

/* === Text areas === */
.stTextArea textarea {
    background: #0D1F30 !important; border: 1px solid #1E3448 !important;
    color: #E8E4DA !important; border-radius: 8px !important;
    font-family: 'Inter', sans-serif !important; font-size: 0.84rem !important;
    line-height: 1.65 !important;
}
.stTextArea textarea:focus { border-color: #C9A84C !important; box-shadow: 0 0 0 2px rgba(201,168,76,0.15) !important; }

/* === File uploader === */
[data-testid="stFileUploader"] {
    background: rgba(201,168,76,0.03) !important; border: 1.5px dashed #1E3448 !important;
    border-radius: 10px !important; transition: border-color 0.2s !important;
}
[data-testid="stFileUploader"]:hover { border-color: #C9A84C !important; }

/* === Metrics === */
div[data-testid="stMetric"] {
    background: #0D1F30 !important; border: 1px solid #1E3448 !important;
    border-top: 2px solid #C9A84C !important; border-radius: 10px !important;
    padding: 1rem 1.2rem !important;
}
div[data-testid="stMetricValue"] { color: #C9A84C !important; font-family: 'Playfair Display', serif !important; font-size: 1.8rem !important; }
div[data-testid="stMetricLabel"] { color: #607D99 !important; font-size: 0.72rem !important; letter-spacing: 0.1em !important; text-transform: uppercase !important; }

/* === Tabs === */
.stTabs [data-baseweb="tab-list"] { background: #0D1F30 !important; border-radius: 8px 8px 0 0 !important; border-bottom: 1px solid #1E3448 !important; gap: 0 !important; }
.stTabs [data-baseweb="tab"] { color: #607D99 !important; font-size: 0.78rem !important; font-weight: 500 !important; letter-spacing: 0.04em !important; padding: 0.65rem 1rem !important; border-radius: 0 !important; border-bottom: 2px solid transparent !important; }
.stTabs [aria-selected="true"] { color: #C9A84C !important; border-bottom: 2px solid #C9A84C !important; background: transparent !important; }
.stTabs [data-baseweb="tab-panel"] { background: #0D1F30 !important; border: 1px solid #1E3448 !important; border-top: none !important; border-radius: 0 0 8px 8px !important; padding: 1.2rem !important; }

/* === Alerts === */
.stSuccess { background: rgba(39,174,96,0.08) !important; border: 1px solid rgba(39,174,96,0.3) !important; border-radius: 8px !important; color: #7DD8A0 !important; }
.stError   { background: rgba(192,57,43,0.08) !important; border: 1px solid rgba(192,57,43,0.3) !important; border-radius: 8px !important; color: #F0A8A8 !important; }
.stWarning { background: rgba(230,126,34,0.08) !important; border: 1px solid rgba(230,126,34,0.3) !important; border-radius: 8px !important; color: #F5C888 !important; }
.stInfo    { background: rgba(52,152,219,0.08) !important; border: 1px solid rgba(52,152,219,0.3) !important; border-radius: 8px !important; color: #A8CCE8 !important; }

/* === Spinner === */
.stSpinner > div { border-top-color: #C9A84C !important; }

/* === Scrollbar === */
::-webkit-scrollbar { width: 5px; height: 5px; }
::-webkit-scrollbar-track { background: #0B1825; }
::-webkit-scrollbar-thumb { background: #1E3448; border-radius: 3px; }
::-webkit-scrollbar-thumb:hover { background: #C9A84C; }

/* === Chat bubbles === */
.chat-user { background: rgba(201,168,76,0.12); border: 1px solid rgba(201,168,76,0.3); border-radius: 12px 12px 4px 12px; padding: 0.7rem 1rem; margin: 0.4rem 0; color: #E8E4DA; font-size: 0.85rem; max-width: 80%; margin-left: auto; }
.chat-ai { background: #0D1F30; border: 1px solid #1E3448; border-left: 3px solid #C9A84C; border-radius: 12px 12px 12px 4px; padding: 0.7rem 1rem; margin: 0.4rem 0; color: #E8E4DA; font-size: 0.85rem; max-width: 90%; }

/* === Premium cards === */
.premium-card { background: linear-gradient(135deg, #0D1F30 0%, #0B1825 100%); border: 1px solid #1E3448; border-radius: 12px; padding: 1.2rem; margin-bottom: 0.8rem; transition: border-color 0.2s; }
.premium-card:hover { border-color: rgba(201,168,76,0.4); }
.kpi-value { font-family: 'Cormorant Garamond', serif; font-size: 2rem; font-weight: 700; color: #C9A84C; }

/* === Hide default chrome === */
#MainMenu, footer, header { visibility: hidden !important; }

/* ═══════════════════════════════════════════════════════════
   GLASSMORPHISM UPGRADE — Harvey AI / Deloitte Executive Style
   ═══════════════════════════════════════════════════════════ */

/* Glassmorphism base card */
.glass-card {
    background: rgba(13, 31, 48, 0.75) !important;
    backdrop-filter: blur(20px) saturate(180%) !important;
    -webkit-backdrop-filter: blur(20px) saturate(180%) !important;
    border: 1px solid rgba(201, 168, 76, 0.15) !important;
    border-radius: 16px !important;
    box-shadow: 0 8px 32px rgba(0,0,0,0.4), inset 0 1px 0 rgba(201,168,76,0.08) !important;
}

/* Animated KPI Card */
@keyframes kpiPulse {
    0%   { box-shadow: 0 0 0 0 rgba(201,168,76,0.25); }
    70%  { box-shadow: 0 0 0 10px rgba(201,168,76,0); }
    100% { box-shadow: 0 0 0 0 rgba(201,168,76,0); }
}
@keyframes fadeInUp {
    from { opacity: 0; transform: translateY(18px); }
    to   { opacity: 1; transform: translateY(0); }
}
@keyframes shimmer {
    0%   { background-position: -200% center; }
    100% { background-position: 200% center; }
}
@keyframes countUp {
    from { opacity: 0; }
    to   { opacity: 1; }
}
@keyframes slideInLeft {
    from { opacity: 0; transform: translateX(-16px); }
    to   { opacity: 1; transform: translateX(0); }
}
@keyframes glowPulse {
    0%, 100% { box-shadow: 0 0 8px rgba(201,168,76,0.2), 0 4px 20px rgba(0,0,0,0.3); }
    50%       { box-shadow: 0 0 20px rgba(201,168,76,0.45), 0 4px 30px rgba(0,0,0,0.4); }
}

.exec-kpi-card {
    background: linear-gradient(135deg, rgba(13,31,48,0.9) 0%, rgba(11,24,37,0.95) 100%);
    backdrop-filter: blur(16px);
    border: 1px solid rgba(30,52,72,0.8);
    border-top: 2px solid;
    border-radius: 14px;
    padding: 1.2rem 1.4rem;
    margin-bottom: 0.6rem;
    transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    animation: fadeInUp 0.5s ease forwards;
    position: relative;
    overflow: hidden;
}
.exec-kpi-card::before {
    content: '';
    position: absolute;
    top: 0; left: -100%;
    width: 80%;
    height: 1px;
    background: linear-gradient(90deg, transparent, rgba(201,168,76,0.6), transparent);
    animation: shimmer 3s infinite;
}
.exec-kpi-card:hover {
    transform: translateY(-3px);
    box-shadow: 0 12px 40px rgba(0,0,0,0.5);
    border-color: rgba(201,168,76,0.3) !important;
}
.exec-kpi-value {
    font-family: 'Cormorant Garamond', serif;
    font-size: 2.2rem;
    font-weight: 700;
    line-height: 1;
    animation: fadeInUp 0.7s ease forwards;
}
.exec-kpi-label {
    font-size: 0.62rem;
    font-weight: 700;
    letter-spacing: 0.14em;
    text-transform: uppercase;
    color: #607D99;
    margin-bottom: 0.5rem;
}
.exec-kpi-trend {
    font-size: 0.7rem;
    margin-top: 0.4rem;
    display: flex;
    align-items: center;
    gap: 0.3rem;
}

/* Premium Sidebar Navigation */
.sidebar-nav-item {
    display: flex;
    align-items: center;
    gap: 0.6rem;
    padding: 0.6rem 0.8rem;
    border-radius: 8px;
    margin-bottom: 0.25rem;
    transition: all 0.2s ease;
    cursor: pointer;
    font-size: 0.78rem;
    color: #607D99;
    border: 1px solid transparent;
    text-decoration: none;
}
.sidebar-nav-item:hover {
    background: rgba(201,168,76,0.07);
    border-color: rgba(201,168,76,0.2);
    color: #C9A84C;
}
.sidebar-nav-item.active {
    background: rgba(201,168,76,0.1);
    border-color: rgba(201,168,76,0.3);
    color: #C9A84C;
    font-weight: 600;
}

/* Executive Dashboard Header */
.exec-dashboard-header {
    background: linear-gradient(135deg, #071219 0%, #0B1825 40%, #0D1F30 70%, #071219 100%);
    border: 1px solid rgba(30,52,72,0.8);
    border-top: 3px solid #C9A84C;
    border-radius: 16px;
    padding: 2rem 2.5rem;
    margin-bottom: 1.5rem;
    position: relative;
    overflow: hidden;
    box-shadow: 0 8px 40px rgba(0,0,0,0.5), inset 0 1px 0 rgba(201,168,76,0.1);
}
.exec-dashboard-header::after {
    content: '';
    position: absolute;
    bottom: 0; right: 0;
    width: 300px; height: 300px;
    background: radial-gradient(circle, rgba(201,168,76,0.04) 0%, transparent 70%);
    pointer-events: none;
}

/* Version History Timeline */
.version-timeline-container {
    position: relative;
    padding-left: 2.5rem;
}
.version-timeline-container::before {
    content: '';
    position: absolute;
    left: 14px;
    top: 0; bottom: 0;
    width: 2px;
    background: linear-gradient(180deg, #C9A84C, rgba(201,168,76,0.1));
}
.version-node {
    position: relative;
    margin-bottom: 1.5rem;
    animation: slideInLeft 0.4s ease forwards;
}
.version-node::before {
    content: '';
    position: absolute;
    left: -2rem;
    top: 0.6rem;
    width: 12px;
    height: 12px;
    border-radius: 50%;
    border: 2px solid #C9A84C;
    background: #071219;
    box-shadow: 0 0 8px rgba(201,168,76,0.4);
}
.version-node.latest::before {
    background: #C9A84C;
    animation: glowPulse 2s infinite;
}

/* Trend Chart Containers */
.chart-container-glass {
    background: rgba(7, 18, 25, 0.6);
    backdrop-filter: blur(12px);
    border: 1px solid rgba(30,52,72,0.7);
    border-radius: 12px;
    padding: 1rem;
    margin-bottom: 0.8rem;
}

/* Risk badge animations */
@keyframes riskPulse {
    0%, 100% { opacity: 1; }
    50%       { opacity: 0.7; }
}
.risk-badge-high { animation: riskPulse 2s infinite; }

/* Premium sidebar logo area */
.sidebar-logo-area {
    background: linear-gradient(135deg, #071219 0%, #0D1F30 100%);
    border-bottom: 1px solid rgba(201,168,76,0.2);
    padding: 1.8rem 1rem 1.4rem 1rem;
    margin: -1rem -1rem 1rem -1rem;
    text-align: center;
    position: relative;
    overflow: hidden;
}
.sidebar-logo-area::before {
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0;
    height: 2px;
    background: linear-gradient(90deg, transparent, #C9A84C, transparent);
    animation: shimmer 4s infinite;
}

/* Diff viewer */
.diff-added    { background: rgba(39,174,96,0.08);  border-left: 3px solid #27AE60; border-radius: 0 6px 6px 0; }
.diff-removed  { background: rgba(231,76,60,0.08);  border-left: 3px solid #E74C3C; border-radius: 0 6px 6px 0; }
.diff-modified { background: rgba(243,156,18,0.08); border-left: 3px solid #F39C12; border-radius: 0 6px 6px 0; }

/* Executive section divider */
.exec-divider {
    height: 1px;
    background: linear-gradient(90deg, transparent, rgba(201,168,76,0.3), transparent);
    margin: 1.2rem 0;
}

/* Smooth section transitions */
.section-enter {
    animation: fadeInUp 0.4s ease forwards;
}

/* Plotly chart override for dark theme */
.js-plotly-plot .plotly .bg { fill: transparent !important; }

/* Selectbox dark theme */
div[data-baseweb="select"] > div {
    background-color: #0D1F30 !important;
    border-color: #1E3448 !important;
    color: #E8E4DA !important;
}

/* Number input */
input[type="number"], input[type="text"] {
    background: #0D1F30 !important;
    border-color: #1E3448 !important;
    color: #E8E4DA !important;
}

/* Expander */
.streamlit-expanderHeader {
    background: #0D1F30 !important;
    border-color: #1E3448 !important;
    color: #E8E4DA !important;
    border-radius: 8px !important;
}

/* Toggle */
.stToggle > label { color: #607D99 !important; }

</style>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# REUSABLE UI HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def h(html: str):
    """Render raw HTML safely."""
    st.markdown(html, unsafe_allow_html=True)


def section_title(text: str, icon: str = ""):
    h(f"""
    <div style="display:flex;align-items:center;gap:0.5rem;
    margin:1.1rem 0 0.7rem 0;padding-bottom:0.45rem;
    border-bottom:1px solid #1E3448;">
        <span style="color:#C9A84C;font-size:0.95rem;">{icon}</span>
        <span style="font-size:0.7rem;font-weight:700;letter-spacing:0.14em;
        text-transform:uppercase;color:#607D99;">{text}</span>
    </div>""")


def kv_row(label: str, value: str, val_color: str = "#E8E4DA"):
    h(f"""
    <div style="display:flex;justify-content:space-between;align-items:center;
    padding:0.42rem 0;border-bottom:1px solid rgba(30,52,72,0.6);font-size:0.83rem;">
        <span style="color:#607D99;">{label}</span>
        <span style="color:{val_color};font-weight:600;">{value}</span>
    </div>""")


def _hex_to_rgb(hx: str) -> str:
    hx = hx.lstrip("#")
    r, g, b = int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16)
    return f"{r},{g},{b}"


def pill(text: str, color: str = "#C9A84C"):
    return (f'<span style="display:inline-block;background:rgba({_hex_to_rgb(color)},0.12);'
            f'border:1px solid rgba({_hex_to_rgb(color)},0.4);border-radius:20px;'
            f'padding:0.18rem 0.65rem;font-size:0.72rem;font-weight:600;'
            f'letter-spacing:0.05em;color:{color};margin:0.15rem;">{text}</span>')


def progress_bar(value: float, color: str = "#C9A84C", height: int = 6) -> str:
    pct = max(0, min(100, int(value * 100)))
    return f"""
    <div style="background:#1E3448;border-radius:4px;height:{height}px;margin-top:0.3rem;overflow:hidden;">
      <div style="width:{pct}%;background:linear-gradient(90deg,{color},{color}aa);
      height:100%;border-radius:4px;transition:width 0.5s ease;"></div>
    </div>"""


def risk_config(level: str):
    return {
        "High":   {"color": "#E74C3C", "bg": "rgba(231,76,60,0.08)",   "border": "rgba(231,76,60,0.3)",   "icon": "🔴"},
        "Medium": {"color": "#F39C12", "bg": "rgba(243,156,18,0.08)",  "border": "rgba(243,156,18,0.3)",  "icon": "🟡"},
        "Low":    {"color": "#27AE60", "bg": "rgba(39,174,96,0.08)",   "border": "rgba(39,174,96,0.3)",   "icon": "🟢"},
    }.get(level, {"color": "#27AE60", "bg": "rgba(39,174,96,0.08)", "border": "rgba(39,174,96,0.3)", "icon": "🟢"})


def compliance_status_icon(status: str) -> str:
    """Return icon and color for compliance status."""
    return {
        "compliant":  ("✓", "#27AE60"),
        "review":     ("⚠", "#F39C12"),
        "missing":    ("✗", "#E74C3C"),
    }.get(status, ("—", "#607D99"))


# ═══════════════════════════════════════════════════════════════════════════════
# API HELPERS (preserved from original)
# ═══════════════════════════════════════════════════════════════════════════════

def check_api_health(url):
    try:
        r = requests.get(f"{url}/health", timeout=20)
        r.raise_for_status()
        return True, r.json()
    except Exception as e:
        return False, str(e)


def warm_backend(url):
    try:
        requests.get(f"{url}/health", timeout=20)
    except Exception:
        pass


def fetch_stats(url):
    try:
        r = requests.get(f"{url}/stats", timeout=8)
        r.raise_for_status()
        return r.json()
    except Exception:
        return {"docs_analyzed": None, "high_risk_flags": None}


def fetch_history(url):
    try:
        r = requests.get(f"{url}/history", timeout=8)
        r.raise_for_status()
        return r.json().get("history", [])
    except Exception:
        return []


def call_api(url, text, retries=2, delay=8):
    text = (text or "").strip()[:3000]
    if not text:
        return None, "No text provided."
    warm_backend(url)
    last_err = None
    for attempt in range(1, retries + 1):
        try:
            r = requests.post(f"{url}/predict", json={"text": text}, timeout=180)
            r.raise_for_status()
            return r.json(), None
        except requests.exceptions.RequestException as e:
            last_err = str(e)
            if attempt < retries:
                time.sleep(delay)
        except Exception as e:
            return None, str(e)
    return None, f"Failed after {retries} attempts: {last_err}"


def extract_pdf_text(file):
    try:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        return "\n".join(p.get_text() for p in doc).strip(), None
    except Exception as e:
        return "", str(e)


# ═══════════════════════════════════════════════════════════════════════════════
# FEATURE 1 — UK EMPLOYMENT COMPLIANCE SCORE
# Derives a compliance score from detected clauses and shows a Plotly gauge
# ═══════════════════════════════════════════════════════════════════════════════

def compute_compliance_score(clauses: dict) -> tuple[int, dict]:
    """
    Map detected clauses to UK Employment compliance checklist.
    Returns (score_0_to_100, checklist_dict).
    """
    checklist = {
        "Holiday Entitlement":   any(k for k in clauses if "holiday" in k.lower() or "annual leave" in k.lower()),
        "Notice Period":         any(k for k in clauses if "notice" in k.lower()),
        "Pension Reference":     any(k for k in clauses if "pension" in k.lower()),
        "Data Protection / GDPR": any(k for k in clauses if "data" in k.lower() or "gdpr" in k.lower()),
        "Equality Clause":       any(k for k in clauses if "equal" in k.lower() or "discriminat" in k.lower()),
        "Grievance Procedure":   any(k for k in clauses if "grievance" in k.lower()),
        "Disciplinary Procedure": any(k for k in clauses if "disciplin" in k.lower()),
        "Confidentiality":       any(k for k in clauses if "confidential" in k.lower()),
    }
    detected = {k: v for k, v in checklist.items() if v}
    # If API returned boolean clause map, merge
    for k, v in clauses.items():
        for ck in checklist:
            if ck.lower().split("/")[0].strip() in k.lower():
                checklist[ck] = checklist[ck] or bool(v)
    score = int(sum(checklist.values()) / len(checklist) * 100)
    return score, checklist


def render_compliance_gauge(score: int, checklist: dict):
    """Feature 1 — circular Plotly gauge + checklist cards."""
    section_title("UK Employment Compliance Score", "🇬🇧")

    gauge_col, list_col = st.columns([1, 1.2])

    with gauge_col:
        color = "#27AE60" if score >= 75 else "#F39C12" if score >= 50 else "#E74C3C"
        fig = go.Figure(go.Indicator(
            mode="gauge+number",
            value=score,
            number={"suffix": "%", "font": {"size": 36, "color": color, "family": "Cormorant Garamond"}},
            gauge={
                "axis": {"range": [0, 100], "tickwidth": 1, "tickcolor": "#1E3448",
                          "tickfont": {"color": "#607D99", "size": 10}},
                "bar": {"color": color, "thickness": 0.25},
                "bgcolor": "#0D1F30",
                "borderwidth": 0,
                "steps": [
                    {"range": [0, 49],  "color": "rgba(231,76,60,0.15)"},
                    {"range": [49, 74], "color": "rgba(243,156,18,0.12)"},
                    {"range": [74, 100],"color": "rgba(39,174,96,0.12)"},
                ],
                "threshold": {"line": {"color": "#C9A84C", "width": 2}, "thickness": 0.75, "value": score},
            },
            title={"text": "Compliance Score", "font": {"color": "#607D99", "size": 12}},
        ))
        fig.update_layout(
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            height=220, margin=dict(l=20, r=20, t=30, b=10),
            font_color="#E8E4DA",
        )
        st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})

    with list_col:
        h('<div style="font-size:0.72rem;font-weight:700;letter-spacing:0.1em;text-transform:uppercase;color:#607D99;margin-bottom:0.6rem;">Compliance Checklist</div>')
        for item, present in checklist.items():
            icon, color_str = ("✓", "#27AE60") if present else ("✗", "#E74C3C")
            h(f"""<div style="display:flex;align-items:center;gap:0.6rem;
            padding:0.35rem 0.7rem;margin-bottom:0.25rem;
            background:rgba({'39,174,96' if present else '231,76,60'},0.06);
            border:1px solid rgba({'39,174,96' if present else '231,76,60'},0.2);
            border-radius:6px;font-size:0.82rem;color:#E8E4DA;">
                <span style="color:{color_str};font-weight:700;min-width:14px;">{icon}</span>
                <span>{item}</span>
            </div>""")


# ═══════════════════════════════════════════════════════════════════════════════
# FEATURE 2 — SCOTTISH EMPLOYMENT LAW CHECKER
# Checks contract text against Scottish-specific legislation
# ═══════════════════════════════════════════════════════════════════════════════

SCOTTISH_LAW_CHECKS = {
    "Equality Act 2010":            ["equality", "discriminat", "protected characteristic"],
    "Working Time Regulations":     ["working time", "rest break", "48 hour", "working hours"],
    "GDPR / Data Protection":       ["gdpr", "data protection", "personal data", "ico"],
    "Pension Auto-Enrolment":       ["auto-enrol", "workplace pension", "nest", "pension contribution"],
    "TUPE Regulations":             ["tupe", "transfer of undertaking", "business transfer"],
    "Restrictive Covenants":        ["non-compete", "non-solicitation", "restrictive covenant", "restraint"],
    "Settlement Clauses":           ["settlement", "compromise agreement", "acas", "without prejudice"],
    "Employment Tribunal Risk":     ["tribunal", "unfair dismissal", "wrongful dismissal", "constructive"],
}


def check_scottish_law(text: str) -> dict:
    """
    Scan contract text for Scottish Employment Law indicators.
    Returns dict of {law: status} where status ∈ {compliant, review, missing}.
    """
    text_lower = text.lower()
    results = {}
    for law, keywords in SCOTTISH_LAW_CHECKS.items():
        matches = sum(1 for kw in keywords if kw in text_lower)
        if matches >= 2:
            results[law] = "compliant"
        elif matches == 1:
            results[law] = "review"
        else:
            results[law] = "missing"
    return results


def render_scottish_law_checker(text: str):
    """Feature 2 — Scottish Employment Law review panel."""
    section_title("Scottish Employment Law Review", "🏴󠁧󠁢󠁳󠁣󠁴󠁿")
    if not text.strip():
        st.info("Upload or paste a contract to run the Scottish Law checker.")
        return

    results = check_scottish_law(text)
    compliant_count = sum(1 for s in results.values() if s == "compliant")
    review_count    = sum(1 for s in results.values() if s == "review")
    missing_count   = sum(1 for s in results.values() if s == "missing")

    # Summary badges
    h(f"""<div style="display:flex;gap:0.8rem;margin-bottom:1rem;">
        <div style="background:rgba(39,174,96,0.1);border:1px solid rgba(39,174,96,0.3);
        border-radius:8px;padding:0.5rem 1rem;text-align:center;">
            <div style="font-size:1.2rem;font-weight:700;color:#27AE60;">{compliant_count}</div>
            <div style="font-size:0.65rem;color:#607D99;text-transform:uppercase;letter-spacing:0.08em;">Compliant</div>
        </div>
        <div style="background:rgba(243,156,18,0.1);border:1px solid rgba(243,156,18,0.3);
        border-radius:8px;padding:0.5rem 1rem;text-align:center;">
            <div style="font-size:1.2rem;font-weight:700;color:#F39C12;">{review_count}</div>
            <div style="font-size:0.65rem;color:#607D99;text-transform:uppercase;letter-spacing:0.08em;">Review</div>
        </div>
        <div style="background:rgba(231,76,60,0.1);border:1px solid rgba(231,76,60,0.3);
        border-radius:8px;padding:0.5rem 1rem;text-align:center;">
            <div style="font-size:1.2rem;font-weight:700;color:#E74C3C;">{missing_count}</div>
            <div style="font-size:0.65rem;color:#607D99;text-transform:uppercase;letter-spacing:0.08em;">Missing</div>
        </div>
    </div>""")

    col1, col2 = st.columns(2)
    items = list(results.items())
    half = len(items) // 2

    for i, (law, status) in enumerate(items):
        icon, color = compliance_status_icon(status)
        label = {"compliant": "Compliant", "review": "Review Required", "missing": "Missing"}[status]
        target_col = col1 if i < half + (len(items) % 2) else col2
        with target_col:
            h(f"""<div style="display:flex;align-items:center;justify-content:space-between;
            padding:0.55rem 0.8rem;margin-bottom:0.35rem;
            background:#0D1F30;border:1px solid #1E3448;border-left:3px solid {color};
            border-radius:0 8px 8px 0;font-size:0.82rem;">
                <span style="color:#E8E4DA;">{law}</span>
                <span style="color:{color};font-weight:700;font-size:0.78rem;white-space:nowrap;margin-left:0.5rem;">
                {icon} {label}</span>
            </div>""")


# ═══════════════════════════════════════════════════════════════════════════════
# FEATURE 3 — HR RISK CATEGORY DASHBOARD (4 Plotly Gauges)
# Replaces single risk visualization with 4 executive gauges
# ═══════════════════════════════════════════════════════════════════════════════

def compute_risk_categories(result: dict) -> dict:
    """
    Derive 4 risk category scores from API result.
    Returns {employee_risk, employer_risk, litigation_risk, compliance_risk}.
    """
    base = result.get("risk_score", 0)
    clauses = result.get("clauses", {})
    label = result.get("label", "").lower()

    # Derive scores heuristically from base risk + clause presence
    n_missing = sum(1 for v in clauses.values() if not v)
    n_total   = max(len(clauses), 1)
    missing_ratio = n_missing / n_total

    employee_risk  = min(100, int(base * 0.8 + missing_ratio * 20))
    employer_risk  = min(100, int(base * 0.6 + (20 if "nda" in label else 10)))
    litigation_risk= min(100, int(base * 0.9 + (15 if "terminat" in label else 5)))
    compliance_risk= min(100, int(missing_ratio * 70 + base * 0.3))

    return {
        "Employee Risk":   employee_risk,
        "Employer Risk":   employer_risk,
        "Litigation Risk": litigation_risk,
        "Compliance Risk": compliance_risk,
    }


def make_gauge(title: str, value: int) -> go.Figure:
    """Build a single executive Plotly gauge."""
    color = "#27AE60" if value < 40 else "#F39C12" if value < 70 else "#E74C3C"
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=value,
        number={"font": {"size": 28, "color": color, "family": "Cormorant Garamond"}},
        gauge={
            "axis": {"range": [0, 100], "tickwidth": 1, "tickcolor": "#1E3448",
                     "tickfont": {"color": "#607D99", "size": 9}},
            "bar": {"color": color, "thickness": 0.22},
            "bgcolor": "#0D1F30",
            "borderwidth": 0,
            "steps": [
                {"range": [0, 39],  "color": "rgba(39,174,96,0.1)"},
                {"range": [39, 69], "color": "rgba(243,156,18,0.1)"},
                {"range": [69, 100],"color": "rgba(231,76,60,0.12)"},
            ],
        },
        title={"text": title, "font": {"color": "#607D99", "size": 11}},
    ))
    fig.update_layout(
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        height=190, margin=dict(l=15, r=15, t=35, b=5),
    )
    return fig


def render_risk_gauges(result: dict):
    """Feature 3 — 4-gauge 2×2 executive risk dashboard."""
    section_title("HR Risk Category Dashboard", "📊")
    risks = compute_risk_categories(result)
    keys = list(risks.keys())

    row1 = st.columns(2)
    row2 = st.columns(2)
    for i, (label, val) in enumerate(risks.items()):
        target = row1[i] if i < 2 else row2[i - 2]
        with target:
            fig = make_gauge(label, val)
            st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})
            # Tooltip-style description below
            desc = {
                "Employee Risk":   "Risk exposure for the employee — missing protections, unfair terms.",
                "Employer Risk":   "Employer liability — indemnities, restrictive obligations.",
                "Litigation Risk": "Probability of employment tribunal or legal challenge.",
                "Compliance Risk": "Regulatory non-compliance — UK/Scottish employment law gaps.",
            }
            level = "High" if val >= 70 else "Medium" if val >= 40 else "Low"
            lc = "#E74C3C" if val >= 70 else "#F39C12" if val >= 40 else "#27AE60"
            h(f"""<div style="background:#071219;border:1px solid #1E3448;border-radius:6px;
            padding:0.45rem 0.7rem;margin-top:-0.5rem;margin-bottom:0.3rem;">
                <div style="font-size:0.7rem;color:#607D99;line-height:1.5;">{desc[label]}</div>
                <div style="font-size:0.7rem;color:{lc};font-weight:600;margin-top:0.2rem;">
                Level: {level}</div>
            </div>""")


# ═══════════════════════════════════════════════════════════════════════════════
# FEATURE 4 — SALARY & BENEFITS EXTRACTION (Compensation Intelligence Tab)
# Uses regex + keyword matching on contract text
# ═══════════════════════════════════════════════════════════════════════════════

def extract_compensation(text: str) -> dict:
    """
    Extract salary, benefits, and leave data from contract text via regex.
    Returns dict of {field: extracted_value_or_None}.
    """
    text_s = text[:8000]
    fields = {
        "Salary":          None,
        "Bonus":           None,
        "Pension":         None,
        "Healthcare":      None,
        "Annual Leave":    None,
        "Sick Pay":        None,
        "Notice Period":   None,
        "Probation Period": None,
    }

    # Salary — match £xx,xxx or £xxx,xxx or GBP patterns
    salary_m = re.search(r"(?:salary|remuneration|basic pay)[^\n£]*£\s*([\d,]+)", text_s, re.I)
    if not salary_m:
        salary_m = re.search(r"£\s*([\d,]+)\s*(?:per annum|p\.?a\.?|annually)", text_s, re.I)
    if salary_m:
        fields["Salary"] = f"£{salary_m.group(1)} per annum"

    # Bonus
    bonus_m = re.search(r"(?:bonus|incentive)[^\n£]*£\s*([\d,]+)|(\d+)\s*%\s*(?:annual )?bonus", text_s, re.I)
    if bonus_m:
        fields["Bonus"] = bonus_m.group(0)[:60]

    # Pension
    pension_m = re.search(r"pension[^\n]*?(\d+\.?\d*)\s*%", text_s, re.I)
    if pension_m:
        fields["Pension"] = f"{pension_m.group(1)}% contribution"
    elif "pension" in text_s.lower():
        fields["Pension"] = "Referenced (amount not specified)"

    # Healthcare
    if re.search(r"private\s*(?:medical|health|healthcare)|bupa|vitality|axa health", text_s, re.I):
        fields["Healthcare"] = "Private Medical — included"

    # Annual Leave
    leave_m = re.search(r"(\d+)\s*(?:days?|working days?)\s*(?:annual|holiday|paid leave)", text_s, re.I)
    if not leave_m:
        leave_m = re.search(r"(?:annual leave|holiday entitlement)[^\n]*?(\d+)\s*days?", text_s, re.I)
    if leave_m:
        fields["Annual Leave"] = f"{leave_m.group(1)} days per year"

    # Sick Pay
    sick_m = re.search(r"(?:sick pay|sickness)[^\n]*?(\d+)\s*(?:weeks?|days?|months?)", text_s, re.I)
    if sick_m:
        fields["Sick Pay"] = sick_m.group(0)[:60]
    elif re.search(r"(?:ssp|statutory sick pay)", text_s, re.I):
        fields["Sick Pay"] = "SSP only"

    # Notice Period
    notice_m = re.search(r"notice[^\n]*?(\d+)\s*(week|month|day)", text_s, re.I)
    if notice_m:
        fields["Notice Period"] = f"{notice_m.group(1)} {notice_m.group(2)}s"

    # Probation
    prob_m = re.search(r"probati(?:on|onary)[^\n]*?(\d+)\s*(week|month)", text_s, re.I)
    if prob_m:
        fields["Probation Period"] = f"{prob_m.group(1)} {prob_m.group(2)}s"

    return fields


def benefits_completeness(comp: dict) -> int:
    """Calculate benefits package completeness score (0-100)."""
    filled = sum(1 for v in comp.values() if v is not None)
    return int(filled / len(comp) * 100)


def render_compensation_intelligence(text: str):
    """Feature 4 — Compensation Intelligence tab."""
    h("""<div style="font-size:0.72rem;color:#607D99;margin-bottom:1rem;letter-spacing:0.06em;">
    Automated extraction of salary, benefits, and employment terms from the contract text.</div>""")

    if not text.strip():
        st.info("Upload or paste a contract to extract compensation data.")
        return

    comp = extract_compensation(text)
    score = benefits_completeness(comp)

    # Completeness score bar
    color = "#27AE60" if score >= 70 else "#F39C12" if score >= 40 else "#E74C3C"
    h(f"""<div style="background:#0D1F30;border:1px solid #1E3448;border-radius:10px;
    padding:1rem 1.2rem;margin-bottom:1rem;">
        <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:0.5rem;">
            <span style="font-size:0.72rem;font-weight:700;letter-spacing:0.1em;text-transform:uppercase;color:#607D99;">
            Benefits Package Completeness</span>
            <span style="font-family:'Cormorant Garamond',serif;font-size:1.6rem;font-weight:700;color:{color};">{score}%</span>
        </div>
        <div style="background:#1E3448;border-radius:4px;height:8px;">
            <div style="width:{score}%;background:linear-gradient(90deg,{color},{color}99);
            height:100%;border-radius:4px;"></div>
        </div>
    </div>""")

    # Individual HR cards
    icons = {
        "Salary": "💷", "Bonus": "🎯", "Pension": "🏦",
        "Healthcare": "🏥", "Annual Leave": "🏖️", "Sick Pay": "🩺",
        "Notice Period": "📅", "Probation Period": "⏱️",
    }
    cols = st.columns(2)
    for i, (field, value) in enumerate(comp.items()):
        col = cols[i % 2]
        with col:
            present = value is not None
            badge_color = "#27AE60" if present else "#3A5570"
            badge_text  = "Extracted" if present else "Not Found"
            display_val = value if present else "—"
            h(f"""<div class="premium-card" style="margin-bottom:0.6rem;">
                <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:0.3rem;">
                    <span style="font-size:0.78rem;font-weight:600;color:#C9A84C;">
                    {icons.get(field,'📋')} &nbsp;{field}</span>
                    <span style="background:rgba({_hex_to_rgb(badge_color)},0.15);
                    border:1px solid rgba({_hex_to_rgb(badge_color)},0.35);
                    border-radius:10px;padding:0.1rem 0.5rem;font-size:0.62rem;
                    font-weight:700;color:{badge_color};">{badge_text}</span>
                </div>
                <div style="font-size:0.88rem;color:#E8E4DA;font-weight:500;">{display_val}</div>
            </div>""")


# ═══════════════════════════════════════════════════════════════════════════════
# FEATURE 5 — CONTRACT COMPARISON TOOL
# Diff two PDF contracts and show added/removed/modified clauses
# ═══════════════════════════════════════════════════════════════════════════════

def diff_contracts(old_text: str, new_text: str) -> dict:
    """
    Compare two contract texts and return categorised diffs.
    Returns {added, removed, modified} sentence lists.
    """
    # Split into sentence-ish blocks
    def sentences(t):
        return [s.strip() for s in re.split(r'(?<=[.!?])\s+|\n{2,}', t) if len(s.strip()) > 20]

    old_s = set(sentences(old_text))
    new_s = set(sentences(new_text))

    added   = list(new_s - old_s)[:20]
    removed = list(old_s - new_s)[:20]

    # Modified — pairs with high similarity but not identical
    modified = []
    for ns in list(new_s - old_s)[:30]:
        for os_ in list(old_s - new_s)[:30]:
            ratio = difflib.SequenceMatcher(None, ns, os_).ratio()
            if 0.5 < ratio < 0.95:
                modified.append({"old": os_[:200], "new": ns[:200]})
                break
    return {"added": added[:10], "removed": removed[:10], "modified": modified[:6]}


def render_contract_comparison():
    """Feature 5 — Contract Comparison tab."""
    h("""<div style="font-size:0.72rem;color:#607D99;margin-bottom:1rem;">
    Upload two versions of a contract to identify changes.</div>""")

    col_old, col_new = st.columns(2)
    with col_old:
        h('<div style="font-size:0.7rem;font-weight:700;letter-spacing:0.1em;text-transform:uppercase;color:#607D99;margin-bottom:0.4rem;">📂 Original Contract</div>')
        old_file = st.file_uploader("Old Contract", type=["pdf"], key="diff_old", label_visibility="collapsed")
    with col_new:
        h('<div style="font-size:0.7rem;font-weight:700;letter-spacing:0.1em;text-transform:uppercase;color:#607D99;margin-bottom:0.4rem;">📂 Revised Contract</div>')
        new_file = st.file_uploader("New Contract", type=["pdf"], key="diff_new", label_visibility="collapsed")

    if old_file and new_file:
        old_text, _ = extract_pdf_text(old_file)
        new_text, _ = extract_pdf_text(new_file)

        with st.spinner("Comparing contracts…"):
            diff = diff_contracts(old_text, new_text)

        added_n   = len(diff["added"])
        removed_n = len(diff["removed"])
        modified_n= len(diff["modified"])

        # Summary
        h(f"""<div style="display:flex;gap:1rem;margin:1rem 0;">
            <div style="background:rgba(39,174,96,0.1);border:1px solid rgba(39,174,96,0.3);
            border-radius:8px;padding:0.6rem 1.2rem;text-align:center;">
                <div style="font-size:1.4rem;font-weight:700;color:#27AE60;">+{added_n}</div>
                <div style="font-size:0.65rem;color:#607D99;text-transform:uppercase;letter-spacing:0.08em;">Added</div>
            </div>
            <div style="background:rgba(231,76,60,0.1);border:1px solid rgba(231,76,60,0.3);
            border-radius:8px;padding:0.6rem 1.2rem;text-align:center;">
                <div style="font-size:1.4rem;font-weight:700;color:#E74C3C;">−{removed_n}</div>
                <div style="font-size:0.65rem;color:#607D99;text-transform:uppercase;letter-spacing:0.08em;">Removed</div>
            </div>
            <div style="background:rgba(243,156,18,0.1);border:1px solid rgba(243,156,18,0.3);
            border-radius:8px;padding:0.6rem 1.2rem;text-align:center;">
                <div style="font-size:1.4rem;font-weight:700;color:#F39C12;">~{modified_n}</div>
                <div style="font-size:0.65rem;color:#607D99;text-transform:uppercase;letter-spacing:0.08em;">Modified</div>
            </div>
        </div>""")

        # Diff display tabs
        dt_add, dt_rem, dt_mod = st.tabs(["  ✚ Added Clauses  ", "  ✖ Removed Clauses  ", "  ⟳ Modified Clauses  "])

        with dt_add:
            if diff["added"]:
                for s in diff["added"]:
                    h(f"""<div style="background:rgba(39,174,96,0.07);border:1px solid rgba(39,174,96,0.25);
                    border-left:3px solid #27AE60;border-radius:0 7px 7px 0;
                    padding:0.5rem 0.8rem;margin-bottom:0.35rem;font-size:0.82rem;color:#D4EDE0;">
                    <span style="color:#27AE60;font-weight:700;">+</span> {s}</div>""")
            else:
                st.info("No new clauses detected.")

        with dt_rem:
            if diff["removed"]:
                for s in diff["removed"]:
                    h(f"""<div style="background:rgba(231,76,60,0.07);border:1px solid rgba(231,76,60,0.25);
                    border-left:3px solid #E74C3C;border-radius:0 7px 7px 0;
                    padding:0.5rem 0.8rem;margin-bottom:0.35rem;font-size:0.82rem;color:#F0A8A8;">
                    <span style="color:#E74C3C;font-weight:700;">−</span> {s}</div>""")
            else:
                st.info("No removed clauses detected.")

        with dt_mod:
            if diff["modified"]:
                for item in diff["modified"]:
                    h(f"""<div style="background:#0D1F30;border:1px solid #1E3448;border-radius:8px;
                    padding:0.7rem;margin-bottom:0.5rem;">
                        <div style="font-size:0.68rem;font-weight:700;color:#F39C12;margin-bottom:0.3rem;">~ MODIFIED</div>
                        <div style="font-size:0.78rem;color:#F0A8A8;margin-bottom:0.3rem;
                        padding:0.3rem 0.6rem;background:rgba(231,76,60,0.05);border-radius:4px;">
                        <b>Before:</b> {item['old']}</div>
                        <div style="font-size:0.78rem;color:#D4EDE0;
                        padding:0.3rem 0.6rem;background:rgba(39,174,96,0.05);border-radius:4px;">
                        <b>After:</b> {item['new']}</div>
                    </div>""")
            else:
                st.info("No modified clauses detected.")
    else:
        h("""<div style="background:#0D1F30;border:1px solid #1E3448;border-radius:10px;
        padding:2rem;text-align:center;color:#3A5570;font-size:0.85rem;">
            Upload both contract versions above to begin comparison analysis.
        </div>""")


# ═══════════════════════════════════════════════════════════════════════════════
# FEATURE 6 — RECRUITMENT SUITABILITY ANALYSIS
# Assesses offer attractiveness for recruitment decision-making
# ═══════════════════════════════════════════════════════════════════════════════

def assess_recruitment(comp: dict, text: str) -> dict:
    """
    Score the contract against recruitment benchmarks.
    Returns scores and a recommendation verdict.
    """
    scores = {}

    # Salary competitiveness — basic heuristic
    salary_str = comp.get("Salary") or ""
    salary_num = 0
    m = re.search(r"[\d,]+", salary_str.replace(",", ""))
    if m:
        try:
            salary_num = int(m.group(0).replace(",", ""))
        except Exception:
            pass
    scores["Salary Competitiveness"]  = min(100, max(20, int(salary_num / 700))) if salary_num else 40
    scores["Offer Attractiveness"]    = min(100, int(sum(1 for v in comp.values() if v) / len(comp) * 100))
    scores["Notice Period Assessment"] = 80 if comp.get("Notice Period") else 30
    scores["Probation Review"]        = 70 if comp.get("Probation Period") else 50
    scores["Benefits Competitiveness"]= min(100, int(sum(1 for k, v in comp.items()
                                           if k not in ("Salary","Notice Period","Probation Period") and v) / 5 * 100))

    avg = int(sum(scores.values()) / len(scores))
    verdict = "Approved" if avg >= 70 else "Review Required" if avg >= 45 else "High Risk"
    return {"scores": scores, "average": avg, "verdict": verdict}


def render_recruitment_intelligence(text: str):
    """Feature 6 — Recruitment Intelligence tab."""
    h("""<div style="font-size:0.72rem;color:#607D99;margin-bottom:1rem;">
    Recruitment suitability analysis and offer competitiveness scoring.</div>""")

    if not text.strip():
        st.info("Upload a contract to run recruitment analysis.")
        return

    comp = extract_compensation(text)
    rec  = assess_recruitment(comp, text)
    scores  = rec["scores"]
    verdict = rec["verdict"]
    avg     = rec["average"]

    verdict_color = {"Approved": "#27AE60", "Review Required": "#F39C12", "High Risk": "#E74C3C"}[verdict]
    verdict_icon  = {"Approved": "✅", "Review Required": "⚠️", "High Risk": "🔴"}[verdict]

    # Verdict banner
    h(f"""<div style="background:rgba({_hex_to_rgb(verdict_color)},0.1);
    border:1px solid rgba({_hex_to_rgb(verdict_color)},0.3);border-left:5px solid {verdict_color};
    border-radius:0 10px 10px 0;padding:1rem 1.4rem;margin-bottom:1rem;
    display:flex;align-items:center;justify-content:space-between;">
        <div>
            <div style="font-size:0.68rem;letter-spacing:0.12em;text-transform:uppercase;
            color:#607D99;margin-bottom:0.2rem;">Recruitment Recommendation</div>
            <div style="font-size:1.2rem;font-weight:700;color:{verdict_color};">
            {verdict_icon} &nbsp;{verdict}</div>
        </div>
        <div style="text-align:right;">
            <div style="font-family:'Cormorant Garamond',serif;font-size:2.5rem;
            font-weight:700;color:{verdict_color};line-height:1;">{avg}</div>
            <div style="font-size:0.7rem;color:#607D99;">Overall Score</div>
        </div>
    </div>""")

    # Score breakdown
    for label, val in scores.items():
        bar_color = "#27AE60" if val >= 70 else "#F39C12" if val >= 45 else "#E74C3C"
        h(f"""<div style="margin-bottom:0.7rem;">
            <div style="display:flex;justify-content:space-between;font-size:0.8rem;margin-bottom:0.2rem;">
                <span style="color:#E8E4DA;">{label}</span>
                <span style="color:{bar_color};font-weight:600;">{val}</span>
            </div>
            {progress_bar(val/100, bar_color, 8)}
        </div>""")


# ═══════════════════════════════════════════════════════════════════════════════
# FEATURE 7 — CONTRACT AI CHAT ASSISTANT
# LLM-ready chat interface with conversation history
# ═══════════════════════════════════════════════════════════════════════════════

def call_chat_api(url: str, question: str, contract_text: str, history: list) -> str:
    """
    Placeholder LLM chat function.
    In production, replace with OpenAI/Anthropic/custom LLM endpoint.
    Falls back to keyword-based demo responses when backend unavailable.
    """
    try:
        payload = {
            "question": question,
            "contract_text": contract_text[:2000],
            "history": history[-6:],
        }
        r = requests.post(f"{url}/chat", json=payload, timeout=30)
        r.raise_for_status()
        return r.json().get("answer", "No response from AI.")
    except Exception:
        # Demo fallback responses based on keywords
        q = question.lower()
        if "notice" in q:
            nm = re.search(r"(\d+)\s*(week|month)", contract_text.lower())
            return f"Based on the contract, the notice period is **{nm.group(0)}**." if nm else \
                   "I could not find a specific notice period in this contract. This may be an omission worth addressing."
        if "summar" in q:
            return "This contract establishes an employment relationship. Key provisions include the terms of engagement, remuneration, and obligations of both parties. I recommend a full legal review before signing."
        if "risk" in q:
            return "The primary risks identified are: missing standard protections, ambiguous termination clauses, and potential gaps in GDPR compliance. Review the Risk Dashboard tab for a detailed breakdown."
        if "missing" in q or "clause" in q:
            return "Clauses that may be missing include: explicit grievance procedure, equality/diversity statement, and whistleblowing policy reference. These are recommended under Scottish employment best practice."
        return "I've analysed your query. For detailed contract intelligence, please ensure the contract text is loaded in the main analysis panel and run a full analysis first."


def render_chat_assistant(contract_text: str):
    """Feature 7 — AI Chat Assistant tab."""
    if "chat_history" not in st.session_state:
        st.session_state["chat_history"] = []

    h("""<div style="font-size:0.72rem;color:#607D99;margin-bottom:0.8rem;">
    Ask questions about the contract. Powered by Caledonian AI — Scottish Employment Law specialist.</div>""")

    # Suggested questions
    h('<div style="font-size:0.68rem;letter-spacing:0.1em;text-transform:uppercase;color:#3A5570;margin-bottom:0.4rem;">Quick Questions</div>')
    q_cols = st.columns(4)
    quick_qs = [
        "What is the notice period?",
        "Summarise this contract.",
        "Explain the key risks.",
        "What clauses are missing?",
    ]
    for i, qcol in enumerate(q_cols):
        with qcol:
            if st.button(quick_qs[i], key=f"quick_q_{i}", use_container_width=True):
                st.session_state["chat_input_val"] = quick_qs[i]

    # Chat display area
    chat_html = '<div style="background:#071219;border:1px solid #1E3448;border-radius:10px;padding:1rem;min-height:280px;max-height:380px;overflow-y:auto;margin-bottom:0.8rem;">'
    if not st.session_state["chat_history"]:
        chat_html += '<div style="color:#3A5570;font-size:0.82rem;text-align:center;padding-top:2rem;">Ask a question about your contract above.</div>'
    for msg in st.session_state["chat_history"]:
        if msg["role"] == "user":
            chat_html += f'<div class="chat-user">{msg["content"]}</div>'
        else:
            chat_html += f'<div class="chat-ai"><span style="font-size:0.65rem;color:#C9A84C;font-weight:700;display:block;margin-bottom:0.25rem;">⚖️ Caledonian AI</span>{msg["content"]}</div>'
    chat_html += '</div>'
    h(chat_html)

    # Input
    user_input = st.text_input(
        "Ask a question…",
        value=st.session_state.get("chat_input_val", ""),
        key="chat_input",
        label_visibility="collapsed",
        placeholder="e.g. What is the termination clause? Is this contract GDPR compliant?",
    )

    send_col, clear_col = st.columns([3, 1])
    with send_col:
        send = st.button("✉️  Send", use_container_width=True, key="chat_send")
    with clear_col:
        if st.button("🗑 Clear", use_container_width=True, key="chat_clear"):
            st.session_state["chat_history"] = []
            st.rerun()

    if send and user_input.strip():
        st.session_state["chat_history"].append({"role": "user", "content": user_input.strip()})
        with st.spinner("Caledonian AI is analysing…"):
            answer = call_chat_api(API_URL, user_input.strip(), contract_text,
                                   st.session_state["chat_history"])
        st.session_state["chat_history"].append({"role": "assistant", "content": answer})
        st.session_state["chat_input_val"] = ""
        st.rerun()


# ═══════════════════════════════════════════════════════════════════════════════
# FEATURE 8 — EXECUTIVE BOARD REPORT
# McKinsey/Deloitte-style boardroom report
# ═══════════════════════════════════════════════════════════════════════════════

def render_board_report(result: dict, text: str):
    """Feature 8 — Executive Board Report."""
    label      = result.get("label", "N/A")
    risk_score = result.get("risk_score", 0)
    risk_level = result.get("risk_level", "Low")
    clauses    = result.get("clauses", {})
    insights   = result.get("insights", [])
    recs       = result.get("recommendations", [])
    exec_sum   = result.get("executive_summary", {})
    rc         = risk_config(risk_level)

    compliance_score, checklist = compute_compliance_score(clauses)
    scottish_law = check_scottish_law(text) if text else {}
    comp = extract_compensation(text) if text else {}

    # Board report header
    h(f"""
    <div style="background:linear-gradient(135deg,#071219 0%,#0D1F30 60%,#071219 100%);
    border:1px solid #1E3448;border-top:3px solid #C9A84C;border-radius:12px;
    padding:2rem;margin-bottom:1.5rem;">
        <div style="display:flex;align-items:flex-start;justify-content:space-between;">
            <div>
                <div style="font-size:0.65rem;letter-spacing:0.18em;text-transform:uppercase;
                color:#607D99;margin-bottom:0.4rem;">CONFIDENTIAL — BOARD LEVEL REVIEW</div>
                <div style="font-family:'Cormorant Garamond',serif;font-size:2rem;font-weight:700;
                color:#E8E4DA;line-height:1.2;">Employment Contract<br>Risk Assessment Report</div>
                <div style="font-size:0.78rem;color:#607D99;margin-top:0.5rem;">
                Caledonian HR Group · Scottish Employment Intelligence Platform</div>
            </div>
            <div style="text-align:right;">
                <div style="font-size:0.65rem;letter-spacing:0.12em;text-transform:uppercase;
                color:#607D99;">Report Date</div>
                <div style="font-size:0.9rem;color:#C9A84C;font-weight:600;">
                {datetime.now().strftime('%d %B %Y')}</div>
                <div style="margin-top:0.6rem;background:rgba({_hex_to_rgb(rc['color'])},0.15);
                border:1px solid rgba({_hex_to_rgb(rc['color'])},0.4);border-radius:6px;
                padding:0.3rem 0.8rem;font-size:0.8rem;font-weight:700;color:{rc['color']};">
                {rc['icon']} {risk_level} Risk</div>
            </div>
        </div>
    </div>""")

    # KPI row
    kpi1, kpi2, kpi3, kpi4 = st.columns(4)
    kpis = [
        ("Contract Type", label, "#C9A84C"),
        ("Risk Score", f"{risk_score}/100", rc["color"]),
        ("Compliance", f"{compliance_score}%", "#27AE60" if compliance_score >= 75 else "#F39C12"),
        ("Law Checks", f"{sum(1 for s in scottish_law.values() if s=='compliant')}/{len(scottish_law)}", "#4A90D9"),
    ]
    for col, (lbl, val, color) in zip([kpi1, kpi2, kpi3, kpi4], kpis):
        with col:
            h(f"""<div style="background:#0D1F30;border:1px solid #1E3448;border-top:2px solid {color};
            border-radius:10px;padding:1rem;text-align:center;margin-bottom:0.8rem;">
                <div style="font-size:0.62rem;letter-spacing:0.1em;text-transform:uppercase;color:#607D99;margin-bottom:0.3rem;">{lbl}</div>
                <div class="kpi-value" style="color:{color};">{val}</div>
            </div>""")

    # Sections
    s1, s2 = st.columns(2)
    with s1:
        section_title("Key Findings", "🔍")
        for item in insights[:5]:
            h(f"""<div style="padding:0.45rem 0.7rem;margin-bottom:0.3rem;
            background:rgba(201,168,76,0.05);border-left:2px solid #C9A84C;
            border-radius:0 6px 6px 0;font-size:0.8rem;color:#E8E4DA;">{item}</div>""")
        if not insights:
            h('<div style="color:#3A5570;font-size:0.8rem;">No findings available. Run analysis first.</div>')

    with s2:
        section_title("Strategic Risks", "⚠️")
        risk_items = [
            f"Compliance Score: {compliance_score}% — {'below threshold' if compliance_score < 75 else 'acceptable'}",
            f"Missing clauses: {sum(1 for v in clauses.values() if not v)} of {len(clauses)}",
            f"Scottish Law: {sum(1 for s in scottish_law.values() if s=='missing')} items missing",
        ]
        for item in risk_items:
            color = "#F39C12"
            h(f"""<div style="padding:0.45rem 0.7rem;margin-bottom:0.3rem;
            background:rgba(243,156,18,0.05);border-left:2px solid {color};
            border-radius:0 6px 6px 0;font-size:0.8rem;color:#E8E4DA;">{item}</div>""")

    section_title("Recommendations", "✅")
    rec_cols = st.columns(2)
    for i, item in enumerate(recs[:6]):
        with rec_cols[i % 2]:
            h(f"""<div style="padding:0.5rem 0.8rem;margin-bottom:0.35rem;
            background:rgba(39,174,96,0.05);border-left:3px solid #27AE60;
            border-radius:0 7px 7px 0;font-size:0.8rem;color:#E8E4DA;">
            <b style="color:#C9A84C;">{i+1}.</b> {item}</div>""")
    if not recs:
        h('<div style="color:#3A5570;font-size:0.8rem;">No recommendations available.</div>')

    # Executive summary table
    if exec_sum:
        section_title("Executive Summary", "📑")
        for k, v in exec_sum.items():
            kv_row(k.replace("_", " ").title(), str(v))

    # Board report PDF export
    st.markdown("")
    if st.button("📄 Export Board Report PDF", use_container_width=False, key="board_pdf"):
        pdf_bytes = generate_board_pdf(result, text, compliance_score, scottish_law)
        st.download_button(
            "⬇️ Download Board Report",
            data=pdf_bytes,
            file_name=f"caledonian_board_report_{datetime.now().strftime('%Y%m%d')}.pdf",
            mime="application/pdf",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# FEATURE 11 — ENTERPRISE RAG CONTRACT ASSISTANT
# LangChain + FAISS + Sentence Transformers
# ═══════════════════════════════════════════════════════════════════════════════

# Clause-type questions the RAG assistant is tuned to answer
RAG_QUICK_QUESTIONS = [
    "Summarise this contract",
    "What is the notice period?",
    "Explain the key risks",
    "What clauses are missing?",
]

# Question → intent keywords for confidence hinting
RAG_INTENT_MAP = {
    "notice":       ("Notice Period", ["notice period", "weeks notice", "notice of termination"]),
    "terminat":     ("Termination",   ["terminat", "dismissal", "end of employment"]),
    "confidential": ("Confidentiality", ["confidential", "non-disclosure"]),
    "compete":      ("Non-Compete",   ["non-compete", "restraint of trade"]),
    "gdpr":         ("GDPR",          ["gdpr", "data protection", "personal data"]),
    "pension":      ("Pension",       ["pension", "auto-enrol"]),
    "tupe":         ("TUPE",          ["tupe", "transfer of undertaking"]),
    "salary":       ("Salary",        ["salary", "remuneration", "£"]),
    "holiday":      ("Holiday/Leave", ["holiday", "annual leave", "days per year"]),
    "restrict":     ("Restrictive Covenants", ["restrictive covenant", "non-solicit"]),
}


def chunk_contract_text(text: str, chunk_size: int = 300, overlap: int = 60) -> list[str]:
    """
    Split contract into overlapping chunks.
    Uses LangChain RecursiveCharacterTextSplitter when available,
    otherwise falls back to a simple sliding-window splitter.
    """
    if LANGCHAIN_AVAILABLE:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        return splitter.split_text(text)
    # Fallback: paragraph-aware chunking
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    chunks, buf = [], ""
    for para in paragraphs:
        if len(buf) + len(para) < chunk_size:
            buf = (buf + " " + para).strip()
        else:
            if buf:
                chunks.append(buf)
            buf = para
    if buf:
        chunks.append(buf)
    return chunks or [text[:chunk_size]]


@st.cache_resource(show_spinner=False)
def _load_sentence_model():
    """Load (and cache) the sentence-transformer model."""
    if not RAG_AVAILABLE:
        return None
    return SentenceTransformer("all-MiniLM-L6-v2")


def build_faiss_index(chunks: list[str]):
    """
    Encode chunks and return (index, embeddings).
    Returns (None, None) if dependencies missing.
    """
    if not RAG_AVAILABLE:
        return None, None
    model = _load_sentence_model()
    if model is None:
        return None, None
    embeddings = model.encode(chunks, convert_to_numpy=True, show_progress_bar=False)
    dim = embeddings.shape[1]
    index = faiss.IndexFlatL2(dim)
    index.add(embeddings.astype("float32"))
    return index, embeddings


def retrieve_relevant_chunks(
    question: str, chunks: list[str], index, top_k: int = 4
) -> list[tuple[str, float]]:
    """
    Retrieve top_k most relevant chunks for a question.
    Returns list of (chunk_text, similarity_score).
    """
    if not RAG_AVAILABLE or index is None:
        # Keyword fallback — score by word overlap
        q_words = set(question.lower().split())
        scored = []
        for c in chunks:
            c_words = set(c.lower().split())
            score = len(q_words & c_words) / max(len(q_words), 1)
            scored.append((c, score))
        scored.sort(key=lambda x: x[1], reverse=True)
        return scored[:top_k]

    model = _load_sentence_model()
    q_emb = model.encode([question], convert_to_numpy=True).astype("float32")
    distances, indices = index.search(q_emb, top_k)
    # Convert L2 distance to a 0-1 similarity score
    results = []
    for dist, idx in zip(distances[0], indices[0]):
        if idx < len(chunks):
            sim = float(1 / (1 + dist))
            results.append((chunks[idx], sim))
    return results


def _compute_confidence(question: str, retrieved: list[tuple[str, float]]) -> float:
    """Estimate answer confidence from retrieval scores + intent match."""
    if not retrieved:
        return 0.0
    avg_sim = sum(s for _, s in retrieved) / len(retrieved)
    q_lower = question.lower()
    intent_boost = 0.0
    for kw, (_, clause_kws) in RAG_INTENT_MAP.items():
        if kw in q_lower:
            if any(ckw in " ".join(c for c, _ in retrieved).lower() for ckw in clause_kws):
                intent_boost = 0.15
                break
    return min(1.0, avg_sim + intent_boost)


def rag_answer(question: str, contract_text: str) -> dict:
    """
    Core RAG pipeline: chunk → embed → retrieve → synthesise answer.
    Returns dict with keys: answer, confidence, sources, retrieved_context.
    """
    if not contract_text.strip():
        return {"answer": "No contract text loaded.", "confidence": 0.0, "sources": [], "retrieved_context": []}

    # Build / cache the index per unique contract hash
    contract_hash = hashlib.md5(contract_text[:5000].encode()).hexdigest()
    cache_key = f"rag_index_{contract_hash}"

    if cache_key not in st.session_state:
        chunks = chunk_contract_text(contract_text, chunk_size=400, overlap=80)
        index, _ = build_faiss_index(chunks)
        st.session_state[cache_key] = (chunks, index)
    else:
        chunks, index = st.session_state[cache_key]

    retrieved = retrieve_relevant_chunks(question, chunks, index, top_k=4)
    confidence = _compute_confidence(question, retrieved)

    # Build context string for the answer synthesiser
    context = "\n\n---\n\n".join(c for c, _ in retrieved)

    # Answer synthesis — rule-based with context extraction
    q_lower = question.lower()
    answer = _synthesise_answer(q_lower, context, contract_text)

    # Source citations: list chunk snippets with scores
    sources = [
        {"snippet": c[:120] + ("…" if len(c) > 120 else ""), "score": round(s, 3)}
        for c, s in retrieved
    ]

    return {
        "answer": answer,
        "confidence": confidence,
        "sources": sources,
        "retrieved_context": [c for c, _ in retrieved],
    }


def _synthesise_answer(q_lower: str, context: str, full_text: str) -> str:
    """
    Produce a structured answer from retrieved context.
    Extracts real contract data via regex where possible.
    """
    ctx_lower = context.lower()

    # ── Summarise ──────────────────────────────────────────────────────────────
    if any(w in q_lower for w in ["summar", "overview", "what is this", "describe"]):
        parts = []
        # Contract type hint
        for typ in ["employment contract", "nda", "service agreement", "consultancy", "settlement"]:
            if typ in full_text.lower():
                parts.append(f"**Document type:** {typ.title()}")
                break
        # Parties
        party_m = re.search(r"between\s+([A-Z][^,\n]{3,40})\s+(?:and|&)\s+([A-Z][^,\n]{3,40})", full_text)
        if party_m:
            parts.append(f"**Parties:** {party_m.group(1).strip()} and {party_m.group(2).strip()}")
        # Key financials
        sal_m = re.search(r"£\s*([\d,]+)\s*(?:per annum|p\.?a\.?)", full_text, re.I)
        if sal_m:
            parts.append(f"**Salary:** £{sal_m.group(1)} per annum")
        notice_m = re.search(r"(\d+)\s*(week|month)s?\s*(?:written\s*)?notice", full_text, re.I)
        if notice_m:
            parts.append(f"**Notice period:** {notice_m.group(1)} {notice_m.group(2)}s")
        # Clause presence summary
        present = [k for k in ["confidentiality", "pension", "gdpr", "termination", "non-compete", "tupe"]
                   if k in full_text.lower()]
        if present:
            parts.append(f"**Key clauses detected:** {', '.join(c.title() for c in present)}")
        if parts:
            return "\n\n".join(parts)
        return ("This document appears to be an employment-related contract. "
                "Upload additional text for a more detailed summary.")

    # ── Notice period ──────────────────────────────────────────────────────────
    if "notice" in q_lower:
        nm = re.search(r"(\d+)\s*(week|month)s?\s*(?:written\s*)?notice", ctx_lower)
        if nm:
            snippet_text = retrieved_snippet(context, 'notice')
            return (f"The contract specifies a notice period of **{nm.group(1)} {nm.group(2)}s**.\n\n"
                    f"*Extracted from:* \"{snippet_text}\"")
        return ("A specific notice period was not clearly identified in the retrieved clauses. "
                "The contract may use non-standard wording — check Section 7 or the termination clauses manually.")

    # ── Risks ──────────────────────────────────────────────────────────────────
    if any(w in q_lower for w in ["risk", "danger", "concern", "problem", "issue"]):
        risks = []
        if "non-compete" in ctx_lower or "restraint" in ctx_lower:
            risks.append("🔴 **Non-compete / Restraint of Trade** — may be broadly drafted; enforceability risk under Scottish law.")
        if "gdpr" not in ctx_lower and "data protection" not in ctx_lower:
            risks.append("🟠 **GDPR gap** — no data protection clause detected in the retrieved context.")
        if "pension" not in ctx_lower:
            risks.append("🟠 **Pension** — auto-enrolment obligations may not be referenced.")
        if "terminat" in ctx_lower:
            risks.append("🔵 **Termination clauses present** — review for mutual vs employer-only rights.")
        if "tupe" in ctx_lower:
            risks.append("🟠 **TUPE** — transfer provisions detected; verify employee protections are adequate.")
        if not risks:
            risks.append("✅ No immediately obvious high-risk clauses detected in the retrieved context.")
        return "\n\n".join(risks)

    # ── Missing clauses ────────────────────────────────────────────────────────
    if any(w in q_lower for w in ["missing", "absent", "omit", "lack", "not include"]):
        expected = {
            "Grievance Procedure":    ["grievance"],
            "Disciplinary Procedure": ["disciplin"],
            "Equality Clause":        ["equal", "discriminat"],
            "GDPR / Data Protection": ["gdpr", "data protection"],
            "Pension Auto-Enrolment": ["pension", "auto-enrol"],
            "Whistleblowing Policy":  ["whistleblow", "public interest disclosure"],
            "Health & Safety":        ["health and safety", "h&s"],
            "Intellectual Property":  ["intellectual property", "ip rights"],
        }
        missing = [clause for clause, kws in expected.items()
                   if not any(kw in full_text.lower() for kw in kws)]
        if missing:
            lines = [f"• **{m}**" for m in missing]
            return ("The following clauses appear to be missing or not clearly referenced:\n\n"
                    + "\n".join(lines)
                    + "\n\n*These are recommended under UK/Scottish employment best practice.*")
        return "No obvious missing clauses detected — the contract appears to cover standard provisions."

    # ── Default: return most relevant retrieved context ────────────────────────
    snippet = context[:600] + ("…" if len(context) > 600 else "")
    return (f"Based on the most relevant sections of the contract:\n\n"
            f"*{snippet}*\n\n"
            "For a more targeted answer, try a specific question such as: "
            "\"What is the notice period?\" or \"Summarise this contract\".")


def retrieved_snippet(context: str, keyword: str, window: int = 150) -> str:
    """Extract a short snippet around a keyword from context."""
    idx = context.lower().find(keyword)
    if idx == -1:
        return context[:100]
    start = max(0, idx - 40)
    end = min(len(context), idx + window)
    return context[start:end].strip()


def render_rag_assistant(contract_text: str):
    """Feature 11 — Enterprise RAG Contract Assistant."""
    h("""<div style="font-size:0.72rem;color:#607D99;margin-bottom:1rem;letter-spacing:0.04em;">
    Vector-powered contract Q&amp;A — LangChain · FAISS · Sentence Transformers · Source Citations</div>""")

    if not contract_text.strip():
        st.info("Upload or paste a contract to activate the RAG assistant.")
        return

    # Dependency warning
    if not RAG_AVAILABLE:
        h("""<div style="background:rgba(243,156,18,0.08);border:1px solid rgba(243,156,18,0.3);
        border-left:4px solid #F39C12;border-radius:0 8px 8px 0;padding:0.7rem 1rem;
        font-size:0.8rem;color:#F5C888;margin-bottom:1rem;">
        ⚠️ <b>sentence-transformers</b> and <b>faiss-cpu</b> not installed —
        running in keyword-fallback mode. Install with:
        <code>pip install sentence-transformers faiss-cpu langchain</code></div>""")

    # ── Quick questions ────────────────────────────────────────────────────────
    h('<div style="font-size:0.68rem;letter-spacing:0.1em;text-transform:uppercase;'
      'color:#3A5570;margin-bottom:0.4rem;">Quick Questions</div>')
    qq_cols = st.columns(4)
    for i, (col, q) in enumerate(zip(qq_cols, RAG_QUICK_QUESTIONS)):
        with col:
            if st.button(q, key=f"rag_qq_{i}", use_container_width=True):
                st.session_state["rag_question"] = q

    # ── Question input ─────────────────────────────────────────────────────────
    rag_q = st.text_input(
        "RAG Question",
        value=st.session_state.get("rag_question", ""),
        key="rag_input",
        label_visibility="collapsed",
        placeholder="e.g. What are the termination conditions? Is there a GDPR clause?",
    )

    ask_col, clear_col = st.columns([4, 1])
    with ask_col:
        ask_btn = st.button("🔍  Ask RAG Assistant", use_container_width=True, key="rag_ask")
    with clear_col:
        if st.button("🗑", use_container_width=True, key="rag_clear"):
            st.session_state.pop("rag_result", None)
            st.session_state.pop("rag_question", None)
            st.rerun()

    if ask_btn and rag_q.strip():
        st.session_state["rag_question"] = rag_q.strip()
        with st.spinner("Building vector index & retrieving relevant clauses…"):
            result = rag_answer(rag_q.strip(), contract_text)
        st.session_state["rag_result"] = result

    # ── Display result ─────────────────────────────────────────────────────────
    rag_result = st.session_state.get("rag_result")
    if rag_result:
        confidence = rag_result["confidence"]
        conf_pct   = int(confidence * 100)
        conf_color = "#27AE60" if confidence >= 0.65 else "#F39C12" if confidence >= 0.35 else "#E74C3C"
        conf_label = "High" if confidence >= 0.65 else "Medium" if confidence >= 0.35 else "Low"

        # Answer card
        h(f"""<div style="background:#071219;border:1px solid #1E3448;border-left:4px solid #C9A84C;
        border-radius:0 10px 10px 0;padding:1.2rem 1.4rem;margin:1rem 0;">
            <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:0.7rem;">
                <span style="font-size:0.7rem;font-weight:700;letter-spacing:0.12em;
                text-transform:uppercase;color:#C9A84C;">⚖️ RAG Answer</span>
                <span style="background:rgba({_hex_to_rgb(conf_color)},0.15);
                border:1px solid rgba({_hex_to_rgb(conf_color)},0.4);
                border-radius:10px;padding:0.15rem 0.7rem;
                font-size:0.7rem;font-weight:700;color:{conf_color};">
                {conf_label} Confidence · {conf_pct}%</span>
            </div>
            <div style="font-size:0.85rem;color:#E8E4DA;line-height:1.75;white-space:pre-wrap;">{rag_result['answer']}</div>
        </div>""")

        # Confidence progress bar
        h(f"""<div style="margin-bottom:1rem;">
            <div style="display:flex;justify-content:space-between;font-size:0.72rem;
            color:#607D99;margin-bottom:0.3rem;">
                <span>Answer Confidence Score</span><span style="color:{conf_color};">{conf_pct}%</span>
            </div>
            <div style="background:#1E3448;border-radius:4px;height:6px;">
                <div style="width:{conf_pct}%;background:linear-gradient(90deg,{conf_color},{conf_color}99);
                height:100%;border-radius:4px;"></div>
            </div>
        </div>""")

        # Source citations + context viewer
        ctx_tab, src_tab = st.tabs(["  📄 Retrieved Context  ", "  🔗 Source Citations  "])

        with ctx_tab:
            h('<div style="font-size:0.7rem;letter-spacing:0.1em;text-transform:uppercase;'
              'color:#607D99;margin-bottom:0.5rem;">Top Retrieved Chunks (Vector Search)</div>')
            for i, chunk_text in enumerate(rag_result["retrieved_context"], 1):
                h(f"""<div style="background:#0D1F30;border:1px solid #1E3448;
                border-left:3px solid #4A90D9;border-radius:0 8px 8px 0;
                padding:0.7rem 1rem;margin-bottom:0.5rem;">
                    <div style="font-size:0.62rem;color:#4A90D9;font-weight:700;
                    letter-spacing:0.08em;text-transform:uppercase;margin-bottom:0.3rem;">
                    Chunk {i}</div>
                    <div style="font-size:0.8rem;color:#B0C4D8;line-height:1.65;">
                    {chunk_text[:300]}{"…" if len(chunk_text) > 300 else ""}</div>
                </div>""")

        with src_tab:
            h('<div style="font-size:0.7rem;letter-spacing:0.1em;text-transform:uppercase;'
              'color:#607D99;margin-bottom:0.5rem;">Retrieval Similarity Scores</div>')
            for i, src in enumerate(rag_result["sources"], 1):
                score = src["score"]
                sc = "#27AE60" if score >= 0.5 else "#F39C12" if score >= 0.25 else "#607D99"
                bar_w = min(100, int(score * 200))
                h(f"""<div style="background:#0D1F30;border:1px solid #1E3448;border-radius:8px;
                padding:0.6rem 0.9rem;margin-bottom:0.4rem;">
                    <div style="display:flex;justify-content:space-between;font-size:0.75rem;margin-bottom:0.3rem;">
                        <span style="color:#E8E4DA;font-style:italic;">"{src['snippet']}"</span>
                        <span style="color:{sc};font-weight:700;white-space:nowrap;margin-left:0.5rem;">
                        {score:.3f}</span>
                    </div>
                    <div style="background:#1E3448;border-radius:3px;height:3px;">
                        <div style="width:{bar_w}%;background:{sc};height:100%;border-radius:3px;"></div>
                    </div>
                </div>""")


# ═══════════════════════════════════════════════════════════════════════════════
# FEATURE 12 — ENHANCED PDF CLAUSE HIGHLIGHTING (PyMuPDF + Navigator)
# Clickable clause navigator · Risk colour coding · 8 clause types
# ═══════════════════════════════════════════════════════════════════════════════

ENHANCED_CLAUSES = {
    "Termination Clause": {
        "keywords":  ["terminat", "dismissal", "end of employment", "notice of termination"],
        "color":     "#E74C3C",
        "risk":      "High",
        "risk_label": "🔴 High Risk",
        "icon":      "🔴",
        "desc":      "Governs how employment ends; high risk if one-sided.",
    },
    "Notice Period": {
        "keywords":  ["notice period", "weeks notice", "months notice", "notice of termination",
                      "written notice", "notice in writing"],
        "color":     "#E67E22",
        "risk":      "High",
        "risk_label": "🔴 High Risk",
        "icon":      "🔴",
        "desc":      "Defines departure timelines. Missing or asymmetric periods are high risk.",
    },
    "Confidentiality": {
        "keywords":  ["confidential", "non-disclosure", "proprietary information", "trade secret"],
        "color":     "#F39C12",
        "risk":      "Medium",
        "risk_label": "🟠 Medium Risk",
        "icon":      "🟠",
        "desc":      "Protects business secrets; review scope and duration.",
    },
    "Non-Compete": {
        "keywords":  ["non-compete", "restraint of trade", "compete with", "competitive activity",
                      "not to solicit"],
        "color":     "#E67E22",
        "risk":      "Medium",
        "risk_label": "🟠 Medium Risk",
        "icon":      "🟠",
        "desc":      "Restricts post-employment activities; enforceability varies under Scottish law.",
    },
    "GDPR": {
        "keywords":  ["gdpr", "data protection", "personal data", "ico", "data controller",
                      "lawful basis"],
        "color":     "#9B59B6",
        "risk":      "Medium",
        "risk_label": "🟠 Medium Risk",
        "icon":      "🟠",
        "desc":      "Regulatory compliance; missing clause is a legal risk.",
    },
    "TUPE": {
        "keywords":  ["tupe", "transfer of undertaking", "business transfer", "acquired rights"],
        "color":     "#27AE60",
        "risk":      "Low",
        "risk_label": "🟢 Low Risk",
        "icon":      "🟢",
        "desc":      "Transfer protections — informational but important for employees.",
    },
    "Pension": {
        "keywords":  ["pension", "auto-enrol", "nest", "workplace pension", "pension contribution"],
        "color":     "#27AE60",
        "risk":      "Low",
        "risk_label": "🟢 Low Risk",
        "icon":      "🟢",
        "desc":      "Auto-enrolment compliance — should be present in all UK contracts.",
    },
    "Restrictive Covenants": {
        "keywords":  ["restrictive covenant", "non-solicitation", "non-dealing", "garden leave",
                      "poach"],
        "color":     "#4A90D9",
        "risk":      "Info",
        "risk_label": "🔵 Informational",
        "icon":      "🔵",
        "desc":      "Post-employment restrictions — review reasonableness and geography.",
    },
}


def find_clause_occurrences(text: str, clause_name: str) -> list[dict]:
    """
    Find all occurrences of a clause's keywords in the text.
    Returns list of {start, end, keyword, excerpt}.
    """
    cfg = ENHANCED_CLAUSES[clause_name]
    occurrences = []
    for kw in cfg["keywords"]:
        for m in re.finditer(re.escape(kw), text, re.IGNORECASE):
            start = max(0, m.start() - 80)
            end   = min(len(text), m.end() + 120)
            occurrences.append({
                "start":   m.start(),
                "end":     m.end(),
                "keyword": m.group(0),
                "excerpt": text[start:end].strip(),
            })
    # Deduplicate by start position
    seen, deduped = set(), []
    for occ in sorted(occurrences, key=lambda x: x["start"]):
        if occ["start"] not in seen:
            seen.add(occ["start"])
            deduped.append(occ)
    return deduped[:8]


def render_enhanced_clause_highlighter(text: str, uploaded_file=None):
    """Feature 12 — Enhanced PDF Clause Highlighting with Navigator."""
    h("""<div style="font-size:0.72rem;color:#607D99;margin-bottom:1rem;">
    Real clause detection with risk colour coding · Clickable navigator ·
    🔴 High &nbsp;🟠 Medium &nbsp;🟢 Low &nbsp;🔵 Informational</div>""")

    if not text.strip():
        st.info("Upload or paste contract text to activate clause highlighting.")
        return

    # ── Risk legend ────────────────────────────────────────────────────────────
    legend_html = '<div style="display:flex;flex-wrap:wrap;gap:0.5rem;margin-bottom:1.2rem;">'
    for cname, cfg in ENHANCED_CLAUSES.items():
        found = any(kw in text.lower() for kw in cfg["keywords"])
        opacity = "1" if found else "0.35"
        legend_html += (
            f'<span style="opacity:{opacity};background:rgba({_hex_to_rgb(cfg["color"])},0.15);'
            f'border:1px solid rgba({_hex_to_rgb(cfg["color"])},0.5);border-radius:4px;'
            f'padding:0.18rem 0.6rem;font-size:0.7rem;color:{cfg["color"]};font-weight:600;">'
            f'{cfg["icon"]} {cname}</span>'
        )
    legend_html += "</div>"
    h(legend_html)

    # ── Detected clauses grid ──────────────────────────────────────────────────
    section_title("Clause Navigator", "🧭")

    detected   = {n: cfg for n, cfg in ENHANCED_CLAUSES.items()
                  if any(kw in text.lower() for kw in cfg["keywords"])}
    undetected = {n: cfg for n, cfg in ENHANCED_CLAUSES.items() if n not in detected}

    nav_cols = st.columns(4)
    for i, (cname, cfg) in enumerate(ENHANCED_CLAUSES.items()):
        found = cname in detected
        col   = nav_cols[i % 4]
        btn_key = f"nav_clause_{i}"
        with col:
            if found:
                occurrences = find_clause_occurrences(text, cname)
                n_found = len(occurrences)
                h(f"""<div style="background:rgba({_hex_to_rgb(cfg['color'])},0.08);
                border:1px solid rgba({_hex_to_rgb(cfg['color'])},0.35);
                border-top:3px solid {cfg['color']};border-radius:8px;
                padding:0.7rem 0.8rem;margin-bottom:0.5rem;cursor:pointer;"
                onclick="">
                    <div style="font-size:0.75rem;font-weight:700;color:{cfg['color']};
                    margin-bottom:0.15rem;">{cfg['icon']} {cname}</div>
                    <div style="font-size:0.65rem;color:#607D99;">{cfg['risk_label']}</div>
                    <div style="font-size:0.62rem;color:#3A5570;margin-top:0.2rem;">
                    {n_found} occurrence{"s" if n_found != 1 else ""} found</div>
                </div>""")
            else:
                h(f"""<div style="background:rgba(30,52,72,0.2);
                border:1px dashed #1E3448;border-radius:8px;
                padding:0.7rem 0.8rem;margin-bottom:0.5rem;opacity:0.55;">
                    <div style="font-size:0.75rem;font-weight:700;color:#3A5570;
                    margin-bottom:0.15rem;">✗ {cname}</div>
                    <div style="font-size:0.65rem;color:#3A5570;">{cfg['risk_label']}</div>
                    <div style="font-size:0.62rem;color:#2A3F52;margin-top:0.2rem;">
                    Not detected</div>
                </div>""")

    # ── Clause detail expanders ────────────────────────────────────────────────
    if detected:
        section_title("Clause Detail & Excerpts", "📋")
        for cname, cfg in detected.items():
            occurrences = find_clause_occurrences(text, cname)
            with st.expander(f"{cfg['icon']} {cname}  ·  {cfg['risk_label']}  ·  {len(occurrences)} occurrence(s)", expanded=False):
                h(f"""<div style="font-size:0.78rem;color:#607D99;margin-bottom:0.7rem;
                padding-bottom:0.5rem;border-bottom:1px solid #1E3448;">
                {cfg['desc']}</div>""")
                if occurrences:
                    for j, occ in enumerate(occurrences, 1):
                        excerpt = occ["excerpt"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        # Highlight the trigger keyword in the excerpt
                        excerpt_hl = re.sub(
                            f"({re.escape(occ['keyword'])})",
                            f'<mark style="background:rgba({_hex_to_rgb(cfg["color"])},0.3);'
                            f'color:{cfg["color"]};border-radius:3px;padding:0 3px;'
                            f'border-bottom:1.5px solid {cfg["color"]};">\\1</mark>',
                            excerpt,
                            flags=re.IGNORECASE,
                        )
                        h(f"""<div style="background:#071219;border:1px solid #1E3448;
                        border-left:3px solid {cfg['color']};border-radius:0 8px 8px 0;
                        padding:0.65rem 0.9rem;margin-bottom:0.45rem;">
                            <div style="font-size:0.62rem;color:{cfg['color']};font-weight:700;
                            letter-spacing:0.08em;text-transform:uppercase;margin-bottom:0.3rem;">
                            Occurrence {j} · Position {occ['start']:,}</div>
                            <div style="font-size:0.8rem;color:#B0C4D8;line-height:1.65;">
                            …{excerpt_hl}…</div>
                        </div>""")
                else:
                    st.info("No specific excerpts found.")

    # ── Full highlighted text ──────────────────────────────────────────────────
    section_title("Full Contract Viewer — Highlighted", "🔍")

    # Filter clause selector
    active_clauses = st.multiselect(
        "Show highlights for:",
        options=list(ENHANCED_CLAUSES.keys()),
        default=list(detected.keys()),
        key="highlight_filter",
    )

    display_text = text[:5000]
    highlighted  = display_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    for cname in active_clauses:
        cfg = ENHANCED_CLAUSES[cname]
        risk_lbl = cfg["risk_label"]
        c_color  = cfg["color"]
        for kw in cfg["keywords"]:
            highlighted = re.compile(f"({re.escape(kw)})", re.IGNORECASE).sub(
                f'<mark style="background:rgba({_hex_to_rgb(c_color)},0.22);'
                f'color:{c_color};border-radius:3px;padding:0 3px;'
                f'border-bottom:2px solid {c_color};" '
                f'title="{cname} - {risk_lbl}">\\1</mark>',
                highlighted,
            )

    h(f"""<div style="background:#071219;border:1px solid #1E3448;border-radius:10px;
    padding:1.2rem;max-height:500px;overflow-y:auto;
    font-size:0.82rem;color:#B0C4D8;line-height:1.85;font-family:'Inter',sans-serif;
    white-space:pre-wrap;">
    {highlighted}
    {"<div style='color:#3A5570;font-size:0.72rem;margin-top:0.5rem;text-align:right;'>[First 5,000 characters shown]</div>" if len(text) > 5000 else ""}
    </div>""")

    # ── Risk summary ───────────────────────────────────────────────────────────
    section_title("Risk Summary", "⚠️")
    risk_buckets = {"High": [], "Medium": [], "Low": [], "Info": []}
    for cname, cfg in ENHANCED_CLAUSES.items():
        found = cname in detected
        bucket = cfg["risk"] if found else None
        if bucket:
            risk_buckets[bucket].append(cname)

    rc1, rc2, rc3, rc4 = st.columns(4)
    for col, (bucket, color, icon) in zip(
        [rc1, rc2, rc3, rc4],
        [("High", "#E74C3C", "🔴"), ("Medium", "#F39C12", "🟠"),
         ("Low", "#27AE60", "🟢"), ("Info", "#4A90D9", "🔵")],
    ):
        clauses_in = risk_buckets[bucket]
        with col:
            h(f"""<div style="background:rgba({_hex_to_rgb(color)},0.08);
            border:1px solid rgba({_hex_to_rgb(color)},0.3);border-top:3px solid {color};
            border-radius:8px;padding:0.8rem;text-align:center;margin-bottom:0.5rem;">
                <div style="font-size:1.5rem;font-weight:700;color:{color};">{icon} {len(clauses_in)}</div>
                <div style="font-size:0.65rem;color:#607D99;text-transform:uppercase;
                letter-spacing:0.08em;margin-top:0.2rem;">{bucket} Risk</div>
            </div>""")
            for c in clauses_in:
                h(f"""<div style="font-size:0.72rem;color:{color};padding:0.2rem 0.4rem;
                background:rgba({_hex_to_rgb(color)},0.06);border-radius:4px;
                margin-bottom:0.2rem;">{c}</div>""")


# ═══════════════════════════════════════════════════════════════════════════════
# FEATURE 9 — PDF CLAUSE HIGHLIGHTING (Original — preserved)

HIGHLIGHT_CLAUSES = {
    "Termination Clause":    (["terminat", "dismissal", "end of employment"], "#E74C3C"),
    "Confidentiality":       (["confidential", "non-disclosure", "proprietary"], "#F39C12"),
    "Non-Compete":           (["non-compete", "restraint of trade", "restrictive covenant"], "#E67E22"),
    "Notice Period":         (["notice period", "notice of termination", "weeks notice"], "#4A90D9"),
    "GDPR References":       (["gdpr", "data protection", "personal data", "ico"], "#9B59B6"),
    "TUPE References":       (["tupe", "transfer of undertaking"], "#16A085"),
}


def render_clause_highlighter(text: str):
    """Feature 9 — Interactive clause highlighting viewer."""
    h("""<div style="font-size:0.72rem;color:#607D99;margin-bottom:1rem;">
    Key contract clauses highlighted by risk level. Upload a contract to activate.</div>""")

    if not text.strip():
        st.info("Upload or paste contract text to enable clause highlighting.")
        return

    # Legend
    legend_html = '<div style="display:flex;flex-wrap:wrap;gap:0.5rem;margin-bottom:1rem;">'
    for clause_name, (_, color) in HIGHLIGHT_CLAUSES.items():
        legend_html += f'<span style="background:rgba({_hex_to_rgb(color)},0.15);border:1px solid rgba({_hex_to_rgb(color)},0.4);border-radius:4px;padding:0.15rem 0.55rem;font-size:0.7rem;color:{color};font-weight:600;">{clause_name}</span>'
    legend_html += '</div>'
    h(legend_html)

    # Build highlighted HTML
    display_text = text[:4000]
    highlighted  = display_text

    # Escape HTML first
    highlighted = highlighted.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    # Apply highlights (case-insensitive, word boundary)
    for clause_name, (keywords, color) in HIGHLIGHT_CLAUSES.items():
        for kw in keywords:
            pattern = re.compile(f"({re.escape(kw)})", re.IGNORECASE)
            highlighted = pattern.sub(
                f'<mark style="background:rgba({_hex_to_rgb(color)},0.25);'
                f'color:{color};border-radius:3px;padding:0 3px;'
                f'border-bottom:1.5px solid {color};" '
                f'title="{clause_name}">\\1</mark>',
                highlighted
            )

    h(f"""<div style="background:#071219;border:1px solid #1E3448;border-radius:10px;
    padding:1.2rem;max-height:480px;overflow-y:auto;
    font-size:0.82rem;color:#B0C4D8;line-height:1.8;font-family:'Inter',sans-serif;
    white-space:pre-wrap;">
    {highlighted}
    </div>""")

    # Clause detection summary
    section_title("Clause Detection Summary", "🔎")
    det_cols = st.columns(3)
    for i, (clause_name, (keywords, color)) in enumerate(HIGHLIGHT_CLAUSES.items()):
        found = any(kw in text.lower() for kw in keywords)
        with det_cols[i % 3]:
            icon = "✓" if found else "✗"
            ic   = color if found else "#3A5570"
            h(f"""<div style="display:flex;align-items:center;gap:0.5rem;
            padding:0.4rem 0.7rem;margin-bottom:0.3rem;
            background:rgba({_hex_to_rgb(ic)},0.06);border:1px solid rgba({_hex_to_rgb(ic)},0.2);
            border-radius:6px;font-size:0.78rem;color:#E8E4DA;">
                <span style="color:{ic};font-weight:700;">{icon}</span> {clause_name}
            </div>""")




# ═══════════════════════════════════════════════════════════════════════════════
# FEATURE 13 — MULTI-DOCUMENT PORTFOLIO ANALYSIS
# Upload 1-100 contracts, compute portfolio-level intelligence dashboard
# ═══════════════════════════════════════════════════════════════════════════════

def _parse_salary_num(salary_str: str) -> float:
    """Extract numeric salary value from formatted string like '£45,000 per annum'."""
    if not salary_str:
        return 0.0
    m = re.search(r"[\d,]+", salary_str.replace(",", ""))
    try:
        return float(m.group(0)) if m else 0.0
    except Exception:
        return 0.0


def _parse_notice_weeks(notice_str: str) -> float:
    """Convert notice period string to weeks."""
    if not notice_str:
        return 0.0
    m = re.search(r"(\d+)\s*(week|month|day)", notice_str, re.I)
    if not m:
        return 0.0
    n, unit = int(m.group(1)), m.group(2).lower()
    if "month" in unit:
        return n * 4.33
    if "day" in unit:
        return n / 5.0
    return float(n)


def _contract_type_from_text(text: str) -> str:
    """Best-guess contract type from text keywords."""
    t = text.lower()
    if "non-disclosure" in t or "nda" in t:
        return "NDA"
    if "service agreement" in t or "services agreement" in t:
        return "Service Agreement"
    if "consultancy" in t or "consulting" in t:
        return "Consultancy"
    if "settlement" in t:
        return "Settlement"
    if "zero hour" in t or "zero-hour" in t:
        return "Zero Hours"
    if "fixed term" in t or "fixed-term" in t:
        return "Fixed Term"
    if "employment" in t or "employee" in t:
        return "Employment Contract"
    return "Other"


def analyse_portfolio_contract(text: str, filename: str) -> dict:
    """
    Run lightweight local analysis on a single contract for portfolio aggregation.
    Returns a dict with all metrics needed for dashboard.
    """
    comp  = extract_compensation(text)
    score, checklist = compute_compliance_score({})

    # Compute compliance from keywords present in text
    text_lower = text.lower()
    checklist_kw = {
        "Holiday Entitlement":    any(kw in text_lower for kw in ["holiday", "annual leave"]),
        "Notice Period":          any(kw in text_lower for kw in ["notice period", "notice of"]),
        "Pension Reference":      "pension" in text_lower,
        "Data Protection / GDPR": any(kw in text_lower for kw in ["gdpr", "data protection"]),
        "Equality Clause":        any(kw in text_lower for kw in ["equal", "discriminat"]),
        "Grievance Procedure":    "grievance" in text_lower,
        "Disciplinary Procedure": "disciplin" in text_lower,
        "Confidentiality":        "confidential" in text_lower,
    }
    compliance_score = int(sum(checklist_kw.values()) / len(checklist_kw) * 100)

    # Risk score heuristic
    risk_factors = [
        30 if not checklist_kw["Notice Period"]          else 0,
        20 if not checklist_kw["Data Protection / GDPR"] else 0,
        15 if not checklist_kw["Disciplinary Procedure"] else 0,
        15 if not checklist_kw["Grievance Procedure"]    else 0,
        10 if "non-compete" in text_lower                else 0,
        10 if "restraint" in text_lower                  else 0,
    ]
    risk_score = min(100, sum(risk_factors))
    risk_level = "High" if risk_score >= 60 else "Medium" if risk_score >= 30 else "Low"

    salary_num   = _parse_salary_num(comp.get("Salary") or "")
    notice_weeks = _parse_notice_weeks(comp.get("Notice Period") or "")

    # Missing clause list
    missing_clauses = [k for k, v in checklist_kw.items() if not v]

    # Clause presence flags (8 standard)
    clause_flags = {
        "Confidentiality":    "confidential" in text_lower,
        "Non-Compete":        any(kw in text_lower for kw in ["non-compete", "restraint"]),
        "GDPR":               any(kw in text_lower for kw in ["gdpr", "data protection"]),
        "Notice Period":      any(kw in text_lower for kw in ["notice period", "notice of"]),
        "Pension":            "pension" in text_lower,
        "TUPE":               "tupe" in text_lower,
        "Grievance":          "grievance" in text_lower,
        "Disciplinary":       "disciplin" in text_lower,
    }

    return {
        "filename":         filename,
        "contract_type":    _contract_type_from_text(text),
        "risk_score":       risk_score,
        "risk_level":       risk_level,
        "compliance_score": compliance_score,
        "salary_num":       salary_num,
        "salary_str":       comp.get("Salary") or "—",
        "notice_weeks":     notice_weeks,
        "notice_str":       comp.get("Notice Period") or "—",
        "missing_clauses":  missing_clauses,
        "clause_flags":     clause_flags,
        "benefits_score":   benefits_completeness(comp),
        "text_length":      len(text),
        "has_pension":      checklist_kw["Pension Reference"],
        "has_gdpr":         checklist_kw["Data Protection / GDPR"],
        "has_notice":       checklist_kw["Notice Period"],
    }


def _portfolio_kpi_card(label: str, value: str, color: str, icon: str, sub: str = "") -> str:
    return f"""
    <div style="background:linear-gradient(135deg,#0D1F30,#0B1825);
    border:1px solid #1E3448;border-top:3px solid {color};border-radius:10px;
    padding:1rem 1.1rem;text-align:center;">
        <div style="font-size:1.4rem;margin-bottom:0.2rem;">{icon}</div>
        <div style="font-family:'Cormorant Garamond',serif;font-size:1.9rem;
        font-weight:700;color:{color};line-height:1.1;">{value}</div>
        <div style="font-size:0.65rem;letter-spacing:0.1em;text-transform:uppercase;
        color:#607D99;margin-top:0.2rem;">{label}</div>
        {f'<div style="font-size:0.7rem;color:#3A5570;margin-top:0.15rem;">{sub}</div>' if sub else ''}
    </div>"""


def render_portfolio_analysis():
    """Feature 13 — Multi-Document Portfolio Intelligence Dashboard."""
    h("""<div style="font-size:0.72rem;color:#607D99;margin-bottom:0.3rem;letter-spacing:0.04em;">
    Upload 1–100 contracts for portfolio-level risk &amp; compliance intelligence.</div>""")

    # ── Upload section ─────────────────────────────────────────────────────────
    h('<div style="font-size:0.7rem;font-weight:700;letter-spacing:0.12em;text-transform:uppercase;'
      'color:#607D99;margin:0.8rem 0 0.4rem 0;">📂 Upload Contract Portfolio (PDF, max 100 files)</div>')

    portfolio_files = st.file_uploader(
        "Portfolio Upload",
        type=["pdf"],
        accept_multiple_files=True,
        key="portfolio_upload",
        label_visibility="collapsed",
    )

    if not portfolio_files:
        h("""<div style="background:#0D1F30;border:1px solid #1E3448;border-radius:12px;
        padding:3rem;text-align:center;margin-top:1rem;">
            <div style="font-size:2.5rem;margin-bottom:0.8rem;">📂</div>
            <div style="font-family:'Cormorant Garamond',serif;font-size:1.3rem;
            color:#E8E4DA;margin-bottom:0.5rem;">Portfolio Dashboard</div>
            <div style="font-size:0.82rem;color:#607D99;line-height:1.7;">
            Upload multiple PDF contracts above to generate a portfolio-wide<br>
            risk assessment, compliance overview, and workforce intelligence report.
            </div>
            <div style="display:flex;justify-content:center;gap:1rem;margin-top:1.2rem;flex-wrap:wrap;">
                <span style="background:rgba(201,168,76,0.08);border:1px solid rgba(201,168,76,0.2);
                border-radius:5px;padding:0.25rem 0.75rem;font-size:0.72rem;color:#C9A84C;">
                Portfolio Risk Score</span>
                <span style="background:rgba(201,168,76,0.08);border:1px solid rgba(201,168,76,0.2);
                border-radius:5px;padding:0.25rem 0.75rem;font-size:0.72rem;color:#C9A84C;">
                Clause Heatmap</span>
                <span style="background:rgba(201,168,76,0.08);border:1px solid rgba(201,168,76,0.2);
                border-radius:5px;padding:0.25rem 0.75rem;font-size:0.72rem;color:#C9A84C;">
                Salary Distribution</span>
                <span style="background:rgba(201,168,76,0.08);border:1px solid rgba(201,168,76,0.2);
                border-radius:5px;padding:0.25rem 0.75rem;font-size:0.72rem;color:#C9A84C;">
                Risk Patterns</span>
            </div>
        </div>""")
        return

    # ── Parse all contracts ────────────────────────────────────────────────────
    if len(portfolio_files) > 100:
        st.warning("Maximum 100 files. Only the first 100 will be analysed.")
        portfolio_files = portfolio_files[:100]

    # Cache portfolio results by file names+sizes hash
    cache_sig = hashlib.md5(
        "".join(f"{f.name}{f.size}" for f in portfolio_files).encode()
    ).hexdigest()

    if st.session_state.get("portfolio_cache_sig") != cache_sig:
        records = []
        prog = st.progress(0, text="Parsing contracts…")
        for i, f in enumerate(portfolio_files):
            text, err = extract_pdf_text(f)
            if text.strip():
                rec = analyse_portfolio_contract(text, f.name)
                records.append(rec)
            prog.progress((i + 1) / len(portfolio_files),
                          text=f"Parsed {i+1}/{len(portfolio_files)}: {f.name[:40]}")
        prog.empty()
        st.session_state["portfolio_records"]  = records
        st.session_state["portfolio_cache_sig"] = cache_sig
    else:
        records = st.session_state.get("portfolio_records", [])

    if not records:
        st.error("No readable text found in uploaded PDFs.")
        return

    n = len(records)

    # ── Aggregate metrics ──────────────────────────────────────────────────────
    avg_risk        = int(sum(r["risk_score"]       for r in records) / n)
    avg_compliance  = int(sum(r["compliance_score"] for r in records) / n)
    salaries        = [r["salary_num"] for r in records if r["salary_num"] > 0]
    avg_salary      = int(sum(salaries) / len(salaries)) if salaries else 0
    notice_vals     = [r["notice_weeks"] for r in records if r["notice_weeks"] > 0]
    avg_notice      = round(sum(notice_vals) / len(notice_vals), 1) if notice_vals else 0
    high_risk_n     = sum(1 for r in records if r["risk_level"] == "High")
    compliant_n     = sum(1 for r in records if r["compliance_score"] >= 75)

    risk_color      = "#E74C3C" if avg_risk >= 60 else "#F39C12" if avg_risk >= 30 else "#27AE60"
    comp_color      = "#27AE60" if avg_compliance >= 75 else "#F39C12" if avg_compliance >= 50 else "#E74C3C"

    # ── KPI Row ────────────────────────────────────────────────────────────────
    section_title("Portfolio Intelligence Dashboard", "📊")
    k1, k2, k3, k4, k5, k6 = st.columns(6)
    kpis = [
        (k1, "Contracts", str(n),                            "#C9A84C", "📂", ""),
        (k2, "Avg Risk Score", str(avg_risk),                risk_color, "⚠️", f"{high_risk_n} high-risk"),
        (k3, "Avg Compliance", f"{avg_compliance}%",         comp_color, "✅", f"{compliant_n} compliant"),
        (k4, "Avg Salary", f"£{avg_salary:,}" if avg_salary else "N/A", "#C9A84C", "💷", f"{len(salaries)} extracted"),
        (k5, "Avg Notice", f"{avg_notice}w" if avg_notice else "N/A",   "#4A90D9", "📅", "weeks"),
        (k6, "High Risk", str(high_risk_n),                  "#E74C3C", "🔴", f"of {n} contracts"),
    ]
    for col, label, val, color, icon, sub in kpis:
        with col:
            h(_portfolio_kpi_card(label, val, color, icon, sub))

    # ── Charts row 1: Risk Distribution + Compliance Distribution ─────────────
    section_title("Risk & Compliance Distributions", "📈")
    ch1, ch2 = st.columns(2)

    with ch1:
        # Risk score histogram
        risk_scores = [r["risk_score"] for r in records]
        fig_risk = go.Figure()
        fig_risk.add_trace(go.Histogram(
            x=risk_scores,
            nbinsx=10,
            marker=dict(
                color=risk_scores,
                colorscale=[[0, "#27AE60"], [0.4, "#F39C12"], [1.0, "#E74C3C"]],
                line=dict(color="#0B1825", width=1),
            ),
            hovertemplate="Risk Score: %{x}<br>Count: %{y}<extra></extra>",
        ))
        fig_risk.add_vline(x=avg_risk, line_dash="dash", line_color="#C9A84C",
                           annotation_text=f"Avg: {avg_risk}", annotation_font_color="#C9A84C")
        fig_risk.update_layout(
            title=dict(text="Risk Score Distribution", font=dict(color="#607D99", size=12)),
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            height=280, margin=dict(l=30, r=10, t=40, b=30),
            font=dict(color="#E8E4DA", size=10),
            xaxis=dict(gridcolor="#1E3448", title="Risk Score"),
            yaxis=dict(gridcolor="#1E3448", title="# Contracts"),
            bargap=0.1,
        )
        st.plotly_chart(fig_risk, use_container_width=True, config={"displayModeBar": False})

    with ch2:
        # Compliance score histogram
        comp_scores = [r["compliance_score"] for r in records]
        fig_comp = go.Figure()
        fig_comp.add_trace(go.Histogram(
            x=comp_scores,
            nbinsx=10,
            marker=dict(
                color=comp_scores,
                colorscale=[[0, "#E74C3C"], [0.5, "#F39C12"], [1.0, "#27AE60"]],
                line=dict(color="#0B1825", width=1),
            ),
            hovertemplate="Compliance: %{x}%<br>Count: %{y}<extra></extra>",
        ))
        fig_comp.add_vline(x=avg_compliance, line_dash="dash", line_color="#C9A84C",
                           annotation_text=f"Avg: {avg_compliance}%", annotation_font_color="#C9A84C")
        fig_comp.update_layout(
            title=dict(text="Compliance Score Distribution", font=dict(color="#607D99", size=12)),
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            height=280, margin=dict(l=30, r=10, t=40, b=30),
            font=dict(color="#E8E4DA", size=10),
            xaxis=dict(gridcolor="#1E3448", title="Compliance Score (%)"),
            yaxis=dict(gridcolor="#1E3448", title="# Contracts"),
            bargap=0.1,
        )
        st.plotly_chart(fig_comp, use_container_width=True, config={"displayModeBar": False})

    # ── Charts row 2: Contract Type Distribution + Salary Distribution ─────────
    section_title("Contract Type & Salary Intelligence", "💷")
    ch3, ch4 = st.columns(2)

    with ch3:
        # Contract type pie/donut
        type_counts: dict = {}
        for r in records:
            ct = r["contract_type"]
            type_counts[ct] = type_counts.get(ct, 0) + 1

        colors_pie = ["#C9A84C", "#4A90D9", "#27AE60", "#E74C3C", "#F39C12",
                      "#9B59B6", "#16A085", "#E67E22", "#2ECC71", "#3498DB"]
        fig_type = go.Figure(go.Pie(
            labels=list(type_counts.keys()),
            values=list(type_counts.values()),
            hole=0.55,
            marker=dict(colors=colors_pie[:len(type_counts)],
                        line=dict(color="#0B1825", width=2)),
            hovertemplate="%{label}<br>%{value} contracts (%{percent})<extra></extra>",
            textfont=dict(color="#E8E4DA", size=10),
        ))
        fig_type.update_layout(
            title=dict(text="Contract Type Distribution", font=dict(color="#607D99", size=12)),
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            height=280, margin=dict(l=10, r=10, t=40, b=10),
            font=dict(color="#E8E4DA", size=10),
            legend=dict(font=dict(color="#607D99", size=9), bgcolor="rgba(0,0,0,0)"),
            annotations=[dict(text=f"<b>{n}</b>", font=dict(size=18, color="#C9A84C"),
                               showarrow=False)],
        )
        st.plotly_chart(fig_type, use_container_width=True, config={"displayModeBar": False})

    with ch4:
        if salaries:
            # Salary distribution with box overlay
            fig_sal = go.Figure()
            fig_sal.add_trace(go.Histogram(
                x=salaries,
                nbinsx=8,
                name="Salary",
                marker=dict(color="#C9A84C", opacity=0.75,
                            line=dict(color="#0B1825", width=1)),
                hovertemplate="£%{x:,.0f}<br>Count: %{y}<extra></extra>",
            ))
            fig_sal.add_vline(x=avg_salary, line_dash="dash", line_color="#4A90D9",
                              annotation_text=f"Avg: £{avg_salary:,}",
                              annotation_font_color="#4A90D9")
            fig_sal.update_layout(
                title=dict(text="Salary Distribution", font=dict(color="#607D99", size=12)),
                paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                height=280, margin=dict(l=30, r=10, t=40, b=30),
                font=dict(color="#E8E4DA", size=10),
                xaxis=dict(gridcolor="#1E3448", title="Annual Salary (£)",
                           tickformat="£,.0f"),
                yaxis=dict(gridcolor="#1E3448", title="# Contracts"),
                bargap=0.1,
            )
            st.plotly_chart(fig_sal, use_container_width=True, config={"displayModeBar": False})
        else:
            h("""<div style="background:#0D1F30;border:1px solid #1E3448;border-radius:10px;
            padding:2rem;text-align:center;color:#3A5570;height:280px;
            display:flex;align-items:center;justify-content:center;">
                <div>💷<br><br>No salary data extracted from portfolio</div>
            </div>""")

    # ── Clause Presence Heatmap ────────────────────────────────────────────────
    section_title("Portfolio Clause Heatmap", "🔥")

    clause_keys = ["Confidentiality", "Non-Compete", "GDPR", "Notice Period",
                   "Pension", "TUPE", "Grievance", "Disciplinary"]

    # Build matrix: rows = contracts (up to 30 for readability), cols = clauses
    display_records = records[:30]
    z_matrix   = [[1 if r["clause_flags"].get(ck, False) else 0 for ck in clause_keys]
                  for r in display_records]
    y_labels   = [r["filename"][:25] + ("…" if len(r["filename"]) > 25 else "")
                  for r in display_records]

    # Color: 1 = green (present), 0 = red (missing)
    colorscale_heatmap = [[0.0, "#3A1010"], [0.5, "#1E3448"], [1.0, "#1A4A2A"]]

    hover_text = []
    for r in display_records:
        row_hover = []
        for ck in clause_keys:
            present = r["clause_flags"].get(ck, False)
            row_hover.append(f"{ck}: {'✓ Present' if present else '✗ Missing'}<br>{r['filename']}")
        hover_text.append(row_hover)

    fig_heat = go.Figure(go.Heatmap(
        z=z_matrix,
        x=clause_keys,
        y=y_labels,
        colorscale=colorscale_heatmap,
        showscale=False,
        hoverinfo="text",
        text=hover_text,
        xgap=2,
        ygap=2,
    ))

    # Add text annotations on cells
    for i, row in enumerate(z_matrix):
        for j, val in enumerate(row):
            fig_heat.add_annotation(
                x=clause_keys[j], y=y_labels[i],
                text="✓" if val else "✗",
                showarrow=False,
                font=dict(size=10, color="#27AE60" if val else "#E74C3C"),
            )

    cell_h = max(18, min(26, 600 // max(len(display_records), 1)))
    fig_heat.update_layout(
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        height=max(300, len(display_records) * cell_h + 80),
        margin=dict(l=180, r=20, t=20, b=60),
        font=dict(color="#E8E4DA", size=9),
        xaxis=dict(side="top", tickfont=dict(size=10, color="#C9A84C"),
                   tickangle=-20, gridcolor="#1E3448"),
        yaxis=dict(tickfont=dict(size=8, color="#607D99"), autorange="reversed"),
    )
    st.plotly_chart(fig_heat, use_container_width=True, config={"displayModeBar": False})
    if n > 30:
        h(f'<div style="font-size:0.72rem;color:#3A5570;margin-top:-0.5rem;">'
          f'Showing first 30 of {n} contracts for readability.</div>')

    # ── Missing Clause Analytics ───────────────────────────────────────────────
    section_title("Missing Clause Analytics", "⚠️")
    mc1, mc2 = st.columns([1.3, 1])

    with mc1:
        # Bar chart: most commonly missing clauses
        missing_counts: dict = {}
        for r in records:
            for mc in r["missing_clauses"]:
                missing_counts[mc] = missing_counts.get(mc, 0) + 1

        if missing_counts:
            sorted_missing = sorted(missing_counts.items(), key=lambda x: x[1], reverse=True)
            m_labels = [x[0] for x in sorted_missing]
            m_vals   = [x[1] for x in sorted_missing]
            m_pcts   = [int(v / n * 100) for v in m_vals]

            bar_colors = ["#E74C3C" if p >= 60 else "#F39C12" if p >= 30 else "#4A90D9"
                          for p in m_pcts]

            fig_missing = go.Figure(go.Bar(
                x=m_vals,
                y=m_labels,
                orientation="h",
                marker=dict(color=bar_colors, line=dict(color="#0B1825", width=1)),
                text=[f"{v} ({p}%)" for v, p in zip(m_vals, m_pcts)],
                textposition="outside",
                textfont=dict(color="#E8E4DA", size=10),
                hovertemplate="%{y}<br>Missing in %{x} contracts<extra></extra>",
            ))
            fig_missing.update_layout(
                paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                height=280, margin=dict(l=180, r=60, t=10, b=20),
                font=dict(color="#E8E4DA", size=10),
                xaxis=dict(gridcolor="#1E3448", title="# Contracts Missing Clause"),
                yaxis=dict(gridcolor="rgba(0,0,0,0)", tickfont=dict(size=10)),
                bargap=0.25,
            )
            st.plotly_chart(fig_missing, use_container_width=True,
                            config={"displayModeBar": False})
        else:
            st.info("No missing clauses detected across portfolio.")

    with mc2:
        section_title("Common Risk Patterns", "🔍")
        # Identify patterns across portfolio
        patterns = []
        gdpr_miss   = sum(1 for r in records if not r["has_gdpr"])
        notice_miss = sum(1 for r in records if not r["has_notice"])
        pension_miss= sum(1 for r in records if not r["has_pension"])
        high_risk_c = sum(1 for r in records if r["risk_score"] >= 60)
        non_compete = sum(1 for r in records if r["clause_flags"].get("Non-Compete", False))

        if gdpr_miss > 0:
            pct = int(gdpr_miss / n * 100)
            patterns.append(("GDPR Gap", f"{gdpr_miss} contracts ({pct}%) missing GDPR clause",
                             "#9B59B6" if pct < 40 else "#E74C3C"))
        if notice_miss > 0:
            pct = int(notice_miss / n * 100)
            patterns.append(("Notice Period Gap", f"{notice_miss} contracts ({pct}%) lack clear notice",
                             "#F39C12" if pct < 40 else "#E74C3C"))
        if pension_miss > 0:
            pct = int(pension_miss / n * 100)
            patterns.append(("Pension Gap", f"{pension_miss} contracts ({pct}%) missing pension reference",
                             "#4A90D9" if pct < 40 else "#F39C12"))
        if high_risk_c > 0:
            pct = int(high_risk_c / n * 100)
            patterns.append(("High Risk Concentration", f"{high_risk_c} contracts ({pct}%) are high risk",
                             "#E74C3C"))
        if non_compete > 0:
            patterns.append(("Non-Compete Prevalence",
                             f"{non_compete} contracts ({int(non_compete/n*100)}%) have restraint clauses",
                             "#F39C12"))
        if not patterns:
            patterns.append(("Portfolio Healthy", "No major common risk patterns detected.", "#27AE60"))

        for pat_name, pat_desc, pat_color in patterns:
            h(f"""<div style="background:rgba({_hex_to_rgb(pat_color)},0.07);
            border:1px solid rgba({_hex_to_rgb(pat_color)},0.3);
            border-left:4px solid {pat_color};border-radius:0 8px 8px 0;
            padding:0.6rem 0.9rem;margin-bottom:0.5rem;">
                <div style="font-size:0.75rem;font-weight:700;color:{pat_color};
                margin-bottom:0.15rem;">{pat_name}</div>
                <div style="font-size:0.78rem;color:#B0C4D8;">{pat_desc}</div>
            </div>""")

    # ── Risk Score Trend (ordered by file index as proxy for time) ─────────────
    section_title("Portfolio Risk Score Trend", "📉")

    fig_trend = go.Figure()
    x_labels = [r["filename"][:20] + ("…" if len(r["filename"]) > 20 else "")
                for r in records]
    risk_vals    = [r["risk_score"]       for r in records]
    compliance_v = [r["compliance_score"] for r in records]

    fig_trend.add_trace(go.Scatter(
        x=list(range(1, n + 1)), y=risk_vals,
        mode="lines+markers",
        name="Risk Score",
        line=dict(color="#E74C3C", width=2),
        marker=dict(size=6, color="#E74C3C"),
        hovertemplate="Contract %{x}: " + "%{customdata}<br>Risk: %{y}<extra></extra>",
        customdata=x_labels,
    ))
    fig_trend.add_trace(go.Scatter(
        x=list(range(1, n + 1)), y=compliance_v,
        mode="lines+markers",
        name="Compliance Score",
        line=dict(color="#27AE60", width=2),
        marker=dict(size=6, color="#27AE60"),
        hovertemplate="Contract %{x}: " + "%{customdata}<br>Compliance: %{y}%<extra></extra>",
        customdata=x_labels,
    ))
    fig_trend.add_hrect(y0=60, y1=100, fillcolor="rgba(231,76,60,0.05)",
                        line_width=0, annotation_text="High Risk Zone",
                        annotation_font_color="#E74C3C", annotation_font_size=9)
    fig_trend.add_hrect(y0=0, y1=30, fillcolor="rgba(39,174,96,0.05)",
                        line_width=0)
    fig_trend.update_layout(
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        height=280, margin=dict(l=40, r=20, t=20, b=40),
        font=dict(color="#E8E4DA", size=10),
        xaxis=dict(gridcolor="#1E3448", title="Contract # (upload order)"),
        yaxis=dict(gridcolor="#1E3448", title="Score", range=[0, 105]),
        legend=dict(font=dict(color="#607D99", size=10), bgcolor="rgba(0,0,0,0)",
                    orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        hovermode="x unified",
    )
    st.plotly_chart(fig_trend, use_container_width=True, config={"displayModeBar": False})

    # ── Contract table ─────────────────────────────────────────────────────────
    section_title("Individual Contract Summary", "📋")

    # Filter controls
    f1, f2, f3 = st.columns(3)
    with f1:
        filter_risk = st.selectbox("Filter by Risk", ["All", "High", "Medium", "Low"],
                                   key="pf_risk_filter")
    with f2:
        filter_type = st.selectbox("Filter by Type",
                                   ["All"] + sorted(set(r["contract_type"] for r in records)),
                                   key="pf_type_filter")
    with f3:
        sort_by = st.selectbox("Sort by", ["Risk Score ↓", "Compliance ↓", "Salary ↓",
                                           "Filename A-Z"],
                               key="pf_sort")

    filtered = records[:]
    if filter_risk != "All":
        filtered = [r for r in filtered if r["risk_level"] == filter_risk]
    if filter_type != "All":
        filtered = [r for r in filtered if r["contract_type"] == filter_type]
    if sort_by == "Risk Score ↓":
        filtered.sort(key=lambda x: x["risk_score"], reverse=True)
    elif sort_by == "Compliance ↓":
        filtered.sort(key=lambda x: x["compliance_score"], reverse=True)
    elif sort_by == "Salary ↓":
        filtered.sort(key=lambda x: x["salary_num"], reverse=True)
    else:
        filtered.sort(key=lambda x: x["filename"])

    # Table header
    h("""<div style="display:grid;grid-template-columns:2fr 1.2fr 0.9fr 0.9fr 1fr 1.2fr;
    gap:0.3rem;padding:0.5rem 0.8rem;background:#071219;border:1px solid #1E3448;
    border-radius:8px 8px 0 0;margin-top:0.5rem;">
        <span style="font-size:0.65rem;font-weight:700;letter-spacing:0.1em;
        text-transform:uppercase;color:#607D99;">Filename</span>
        <span style="font-size:0.65rem;font-weight:700;letter-spacing:0.1em;
        text-transform:uppercase;color:#607D99;">Type</span>
        <span style="font-size:0.65rem;font-weight:700;letter-spacing:0.1em;
        text-transform:uppercase;color:#607D99;">Risk</span>
        <span style="font-size:0.65rem;font-weight:700;letter-spacing:0.1em;
        text-transform:uppercase;color:#607D99;">Compliance</span>
        <span style="font-size:0.65rem;font-weight:700;letter-spacing:0.1em;
        text-transform:uppercase;color:#607D99;">Salary</span>
        <span style="font-size:0.65rem;font-weight:700;letter-spacing:0.1em;
        text-transform:uppercase;color:#607D99;">Notice</span>
    </div>""")

    for r in filtered[:50]:
        rc_col   = "#E74C3C" if r["risk_level"] == "High" else "#F39C12" if r["risk_level"] == "Medium" else "#27AE60"
        cc_col   = "#27AE60" if r["compliance_score"] >= 75 else "#F39C12" if r["compliance_score"] >= 50 else "#E74C3C"
        fname    = r["filename"][:32] + ("…" if len(r["filename"]) > 32 else "")
        h(f"""<div style="display:grid;grid-template-columns:2fr 1.2fr 0.9fr 0.9fr 1fr 1.2fr;
        gap:0.3rem;padding:0.42rem 0.8rem;border:1px solid #1E3448;border-top:none;
        background:#0D1F30;font-size:0.78rem;color:#E8E4DA;
        transition:background 0.15s;"
        onmouseover="this.style.background='#112233'"
        onmouseout="this.style.background='#0D1F30'">
            <span style="color:#B0C4D8;overflow:hidden;text-overflow:ellipsis;
            white-space:nowrap;" title="{r['filename']}">{fname}</span>
            <span style="color:#C9A84C;">{r['contract_type']}</span>
            <span style="color:{rc_col};font-weight:700;">{r['risk_score']} <span style="font-size:0.65rem;">{r['risk_level']}</span></span>
            <span style="color:{cc_col};font-weight:700;">{r['compliance_score']}%</span>
            <span style="color:#B0C4D8;">{r['salary_str']}</span>
            <span style="color:#607D99;">{r['notice_str']}</span>
        </div>""")

    if len(filtered) > 50:
        h(f'<div style="font-size:0.72rem;color:#3A5570;padding:0.4rem 0.8rem;'
          f'border:1px solid #1E3448;border-top:none;border-radius:0 0 8px 8px;'
          f'background:#071219;">Showing first 50 of {len(filtered)} filtered contracts.</div>')
    elif filtered:
        h('<div style="height:4px;background:linear-gradient(90deg,#C9A84C,transparent);'
          'border-radius:0 0 4px 4px;"></div>')


# ═══════════════════════════════════════════════════════════════════════════════
# FEATURE 14 — WORKFORCE RISK HEATMAP
# Rows: HR / Finance / Operations / Sales
# Cols: Employee Risk / Litigation Risk / Compliance Risk / Employer Risk
# Interactive Plotly heatmap with tooltips, department drilldown, scenario sliders
# ═══════════════════════════════════════════════════════════════════════════════

WORKFORCE_DEPARTMENTS = ["HR", "Finance", "Operations", "Sales"]
WORKFORCE_RISK_COLS   = ["Employee Risk", "Litigation Risk", "Compliance Risk", "Employer Risk"]

# Default baseline risk matrix (department × risk_type) — editable via sliders
_DEFAULT_HEATMAP = {
    #                 EmpRisk  LitigRisk  CompRisk  EmplorRisk
    "HR":          [  28,       22,        18,        30  ],
    "Finance":     [  35,       45,        55,        40  ],
    "Operations":  [  60,       50,        42,        55  ],
    "Sales":       [  70,       65,        48,        72  ],
}

# Tooltip descriptions per cell
_CELL_DESCRIPTIONS = {
    ("HR",         "Employee Risk"):   "HR contracts typically have strong protections; moderate risk from policy gaps.",
    ("HR",         "Litigation Risk"): "HR dept. faces low litigation risk due to internal expertise.",
    ("HR",         "Compliance Risk"): "High awareness of compliance requirements; usually well-covered.",
    ("HR",         "Employer Risk"):   "Balanced risk profile; HR contracts usually mutual.",
    ("Finance",    "Employee Risk"):   "Finance roles often include restrictive covenants increasing employee risk.",
    ("Finance",    "Litigation Risk"): "Elevated litigation risk from IP, bonus disputes and incentive clawback.",
    ("Finance",    "Compliance Risk"): "Regulatory obligations (FCA, AML) elevate compliance risk significantly.",
    ("Finance",    "Employer Risk"):   "Sensitive data access increases employer exposure.",
    ("Operations", "Employee Risk"):   "Shift patterns, health & safety gaps create moderate-high employee risk.",
    ("Operations", "Litigation Risk"): "Personal injury and wrongful dismissal claims more prevalent.",
    ("Operations", "Compliance Risk"): "Working Time Regulations compliance often fragmented in ops roles.",
    ("Operations", "Employer Risk"):   "Employer liability elevated due to equipment and site responsibilities.",
    ("Sales",      "Employee Risk"):   "Commission clawback and non-solicitation clauses are common risk drivers.",
    ("Sales",      "Litigation Risk"): "Highest litigation risk — non-compete disputes, commission disagreements.",
    ("Sales",      "Compliance Risk"): "GDPR risks from CRM and client data handling in sales contracts.",
    ("Sales",      "Employer Risk"):   "Employer most exposed in Sales — client list, IP, and territory disputes.",
}


def _build_heatmap_figure(matrix: list, custom_labels: list | None = None) -> go.Figure:
    """Build the interactive Plotly heatmap figure."""
    z = matrix
    y_labels = WORKFORCE_DEPARTMENTS
    x_labels = WORKFORCE_RISK_COLS

    # Build rich hover text
    hover = []
    for i, dept in enumerate(y_labels):
        row_hover = []
        for j, risk_col in enumerate(x_labels):
            val  = z[i][j]
            level = "High" if val >= 60 else "Medium" if val >= 35 else "Low"
            desc  = _CELL_DESCRIPTIONS.get((dept, risk_col), "")
            row_hover.append(
                f"<b>{dept} · {risk_col}</b><br>"
                f"Score: <b>{val}</b> / 100<br>"
                f"Level: <b>{level}</b><br>"
                f"<i>{desc}</i>"
            )
        hover.append(row_hover)

    # Custom diverging colorscale: green (low) → amber → red (high)
    colorscale = [
        [0.00, "#0D3320"],
        [0.25, "#27AE60"],
        [0.50, "#F39C12"],
        [0.75, "#E67E22"],
        [1.00, "#C0392B"],
    ]

    fig = go.Figure(go.Heatmap(
        z=z,
        x=x_labels,
        y=y_labels,
        colorscale=colorscale,
        zmin=0, zmax=100,
        showscale=True,
        hoverinfo="text",
        text=hover,
        xgap=4,
        ygap=4,
        colorbar=dict(
            title=dict(text="Risk Score", font=dict(color="#607D99", size=11)),
            tickfont=dict(color="#607D99", size=10),
            bgcolor="rgba(0,0,0,0)",
            bordercolor="#1E3448",
            borderwidth=1,
            thickness=14,
            len=0.85,
            tickvals=[0, 25, 50, 75, 100],
            ticktext=["0 Low", "25", "50 Med", "75", "100 High"],
        ),
    ))

    # Overlay score + level annotations
    for i, row in enumerate(z):
        for j, val in enumerate(row):
            level = "HIGH" if val >= 60 else "MED" if val >= 35 else "LOW"
            text_color = "#FFFFFF" if val >= 45 else "#E8E4DA"
            fig.add_annotation(
                x=x_labels[j], y=y_labels[i],
                text=f"<b>{val}</b><br><span style='font-size:9px'>{level}</span>",
                showarrow=False,
                font=dict(size=13, color=text_color),
                align="center",
            )

    fig.update_layout(
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        height=360,
        margin=dict(l=90, r=80, t=30, b=60),
        font=dict(color="#E8E4DA", family="Inter"),
        xaxis=dict(
            tickfont=dict(size=11, color="#C9A84C"),
            side="bottom",
            tickangle=0,
            gridcolor="rgba(0,0,0,0)",
        ),
        yaxis=dict(
            tickfont=dict(size=12, color="#C9A84C"),
            autorange="reversed",
            gridcolor="rgba(0,0,0,0)",
        ),
        hoverlabel=dict(
            bgcolor="#0D1F30",
            bordercolor="#C9A84C",
            font=dict(color="#E8E4DA", size=12),
        ),
    )
    return fig


def render_workforce_heatmap():
    """Feature 14 — Interactive Workforce Risk Heatmap."""
    h("""<div style="font-size:0.72rem;color:#607D99;margin-bottom:1rem;letter-spacing:0.04em;">
    Departmental risk profile across four key risk dimensions.
    Adjust sliders to model scenarios or import from Portfolio Analysis.</div>""")

    # ── Mode selector ──────────────────────────────────────────────────────────
    mode = st.radio(
        "Data Source",
        ["Default Benchmarks", "Manual Scenario Editor", "Import from Portfolio"],
        horizontal=True,
        key="heatmap_mode",
        label_visibility="collapsed",
    )

    h(f"""<div style="display:flex;gap:0.5rem;margin-bottom:1rem;">
        <span style="font-size:0.7rem;font-weight:700;letter-spacing:0.1em;
        text-transform:uppercase;color:#607D99;">Mode:</span>
        <span style="font-size:0.72rem;color:#C9A84C;font-weight:600;">{mode}</span>
    </div>""")

    # ── Build matrix ───────────────────────────────────────────────────────────
    matrix = [list(v) for v in _DEFAULT_HEATMAP.values()]  # deep copy

    if mode == "Manual Scenario Editor":
        section_title("Scenario Editor — Adjust Risk Scores", "🎛️")
        st.caption("Drag sliders to model different risk scenarios per department and risk type.")

        for i, dept in enumerate(WORKFORCE_DEPARTMENTS):
            with st.expander(f"🏢 {dept} Department", expanded=(i == 0)):
                cols = st.columns(4)
                for j, risk_col in enumerate(WORKFORCE_RISK_COLS):
                    with cols[j]:
                        val = st.slider(
                            risk_col,
                            min_value=0, max_value=100,
                            value=_DEFAULT_HEATMAP[dept][j],
                            key=f"hm_{dept}_{j}",
                            label_visibility="visible",
                        )
                        matrix[i][j] = val

    elif mode == "Import from Portfolio":
        portfolio_records = st.session_state.get("portfolio_records", [])
        if not portfolio_records:
            st.info("No portfolio data found. Upload contracts in the Portfolio Analysis tab first.")
        else:
            # Map portfolio records to departments by contract type heuristics
            dept_map = {
                "Employment Contract": "HR",
                "NDA":                 "Finance",
                "Service Agreement":   "Operations",
                "Consultancy":         "Sales",
                "Settlement":          "HR",
                "Fixed Term":          "Operations",
                "Zero Hours":          "Operations",
                "Other":               "HR",
            }
            dept_buckets: dict = {d: [] for d in WORKFORCE_DEPARTMENTS}
            for r in portfolio_records:
                dept = dept_map.get(r["contract_type"], "HR")
                dept_buckets[dept].append(r)

            for i, dept in enumerate(WORKFORCE_DEPARTMENTS):
                bucket = dept_buckets[dept]
                if bucket:
                    avg_r = int(sum(b["risk_score"] for b in bucket) / len(bucket))
                    comp_avg = int(sum(b["compliance_score"] for b in bucket) / len(bucket))
                    n_missing_notice = sum(1 for b in bucket if not b["has_notice"])
                    notice_ratio = int(n_missing_notice / len(bucket) * 100)

                    emp_risk   = min(100, int(avg_r * 0.85))
                    litig_risk = min(100, int(avg_r * 0.90 + notice_ratio * 0.1))
                    comp_risk  = min(100, 100 - comp_avg)
                    emplr_risk = min(100, int(avg_r * 0.75 + 10))
                    matrix[i] = [emp_risk, litig_risk, comp_risk, emplr_risk]

            st.success(f"Heatmap populated from {len(portfolio_records)} portfolio contracts.")

    # ── Main heatmap ───────────────────────────────────────────────────────────
    section_title("Workforce Risk Heatmap", "🌡️")
    fig = _build_heatmap_figure(matrix)
    st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})

    # ── Risk legend ────────────────────────────────────────────────────────────
    h("""<div style="display:flex;gap:1.2rem;justify-content:center;margin:0.5rem 0 1.2rem 0;">
        <span style="display:flex;align-items:center;gap:0.4rem;font-size:0.75rem;color:#27AE60;">
            <span style="width:12px;height:12px;border-radius:3px;background:#27AE60;display:inline-block;"></span>
            Low Risk (0–34)</span>
        <span style="display:flex;align-items:center;gap:0.4rem;font-size:0.75rem;color:#F39C12;">
            <span style="width:12px;height:12px;border-radius:3px;background:#F39C12;display:inline-block;"></span>
            Medium Risk (35–59)</span>
        <span style="display:flex;align-items:center;gap:0.4rem;font-size:0.75rem;color:#E74C3C;">
            <span style="width:12px;height:12px;border-radius:3px;background:#E74C3C;display:inline-block;"></span>
            High Risk (60–100)</span>
    </div>""")

    # ── Numerical summary table ────────────────────────────────────────────────
    section_title("Risk Matrix — Numerical Summary", "📊")

    # Header
    cols_h = st.columns([1.2] + [1] * 4 + [1])
    header_labels = ["Department"] + WORKFORCE_RISK_COLS + ["Avg Score"]
    header_colors = ["#607D99", "#E74C3C", "#F39C12", "#9B59B6", "#4A90D9", "#C9A84C"]
    for col, lbl, lc in zip(cols_h, header_labels, header_colors):
        with col:
            h(f'<div style="font-size:0.62rem;font-weight:700;letter-spacing:0.1em;'
              f'text-transform:uppercase;color:{lc};padding:0.3rem 0;'
              f'border-bottom:2px solid {lc};">{lbl}</div>')

    for i, dept in enumerate(WORKFORCE_DEPARTMENTS):
        row_vals = matrix[i]
        avg_row  = int(sum(row_vals) / len(row_vals))
        avg_col  = "#E74C3C" if avg_row >= 60 else "#F39C12" if avg_row >= 35 else "#27AE60"
        cols_r   = st.columns([1.2] + [1] * 4 + [1])

        dept_icons = {"HR": "👥", "Finance": "💰", "Operations": "⚙️", "Sales": "📈"}
        with cols_r[0]:
            h(f'<div style="font-size:0.82rem;font-weight:700;color:#C9A84C;'
              f'padding:0.45rem 0;">{dept_icons.get(dept, "🏢")} {dept}</div>')

        for j, val in enumerate(row_vals):
            vc = "#E74C3C" if val >= 60 else "#F39C12" if val >= 35 else "#27AE60"
            level = "HIGH" if val >= 60 else "MED" if val >= 35 else "LOW"
            with cols_r[j + 1]:
                h(f'<div style="background:rgba({_hex_to_rgb(vc)},0.08);'
                  f'border:1px solid rgba({_hex_to_rgb(vc)},0.25);border-radius:6px;'
                  f'padding:0.4rem 0.5rem;text-align:center;margin:0.15rem 0;">',)
                h(f'<span style="font-size:1rem;font-weight:700;color:{vc};">{val}</span>'
                  f'<br><span style="font-size:0.6rem;color:{vc};opacity:0.7;">{level}</span>'
                  f'</div>')

        with cols_r[5]:
            h(f'<div style="background:rgba({_hex_to_rgb(avg_col)},0.12);'
              f'border:1px solid rgba({_hex_to_rgb(avg_col)},0.4);border-radius:6px;'
              f'padding:0.4rem 0.5rem;text-align:center;margin:0.15rem 0;">'
              f'<span style="font-size:1rem;font-weight:700;color:{avg_col};">{avg_row}</span>'
              f'<br><span style="font-size:0.6rem;color:{avg_col};opacity:0.8;">AVG</span>'
              f'</div>')

    # ── Radial spider chart — per-department risk profile ──────────────────────
    section_title("Department Risk Profiles — Radar Chart", "🕸️")

    dept_colors = {"HR": "#27AE60", "Finance": "#4A90D9",
                   "Operations": "#F39C12", "Sales": "#E74C3C"}
    fig_radar = go.Figure()

    for i, dept in enumerate(WORKFORCE_DEPARTMENTS):
        vals = matrix[i] + [matrix[i][0]]  # close the polygon
        cats = WORKFORCE_RISK_COLS + [WORKFORCE_RISK_COLS[0]]
        color = dept_colors[dept]
        fig_radar.add_trace(go.Scatterpolar(
            r=vals, theta=cats,
            fill="toself",
            name=dept,
            line=dict(color=color, width=2),
            fillcolor=f"rgba({_hex_to_rgb(color)},0.12)",
            hovertemplate=f"<b>{dept}</b><br>%{{theta}}: %{{r}}<extra></extra>",
        ))

    fig_radar.update_layout(
        polar=dict(
            bgcolor="rgba(13,31,48,0.6)",
            radialaxis=dict(
                visible=True, range=[0, 100], showticklabels=True,
                tickfont=dict(size=8, color="#3A5570"),
                gridcolor="#1E3448", linecolor="#1E3448",
                tickvals=[0, 25, 50, 75, 100],
            ),
            angularaxis=dict(
                tickfont=dict(size=11, color="#C9A84C"),
                linecolor="#1E3448", gridcolor="#1E3448",
            ),
        ),
        paper_bgcolor="rgba(0,0,0,0)",
        height=420,
        margin=dict(l=60, r=60, t=40, b=40),
        font=dict(color="#E8E4DA", family="Inter"),
        legend=dict(
            font=dict(color="#607D99", size=11),
            bgcolor="rgba(13,31,48,0.7)",
            bordercolor="#1E3448", borderwidth=1,
            orientation="h", yanchor="bottom", y=-0.15, xanchor="center", x=0.5,
        ),
        hoverlabel=dict(bgcolor="#0D1F30", bordercolor="#C9A84C",
                        font=dict(color="#E8E4DA", size=12)),
    )
    st.plotly_chart(fig_radar, use_container_width=True, config={"displayModeBar": False})

    # ── Highest-risk cell spotlight ────────────────────────────────────────────
    section_title("Risk Spotlight — Highest Priority Actions", "🚨")

    cells_flat = []
    for i, dept in enumerate(WORKFORCE_DEPARTMENTS):
        for j, risk_col in enumerate(WORKFORCE_RISK_COLS):
            cells_flat.append((dept, risk_col, matrix[i][j]))
    cells_flat.sort(key=lambda x: x[2], reverse=True)

    spotlight_cols = st.columns(2)
    for idx, (dept, risk_col, val) in enumerate(cells_flat[:4]):
        vc = "#E74C3C" if val >= 60 else "#F39C12" if val >= 35 else "#27AE60"
        desc = _CELL_DESCRIPTIONS.get((dept, risk_col), "")
        with spotlight_cols[idx % 2]:
            h(f"""<div style="background:rgba({_hex_to_rgb(vc)},0.07);
            border:1px solid rgba({_hex_to_rgb(vc)},0.35);
            border-left:5px solid {vc};border-radius:0 10px 10px 0;
            padding:0.8rem 1rem;margin-bottom:0.6rem;">
                <div style="display:flex;justify-content:space-between;align-items:center;
                margin-bottom:0.3rem;">
                    <span style="font-size:0.78rem;font-weight:700;color:{vc};">
                    {dept} · {risk_col}</span>
                    <span style="font-family:'Cormorant Garamond',serif;font-size:1.6rem;
                    font-weight:700;color:{vc};">{val}</span>
                </div>
                <div style="font-size:0.78rem;color:#B0C4D8;line-height:1.6;">{desc}</div>
            </div>""")


# ═══════════════════════════════════════════════════════════════════════════════
# FEATURE 15 — CLAUSE BENCHMARKING ENGINE
# Benchmarks Salary, Notice Period, Annual Leave, Benefits, Probation,
# Restrictive Covenants against simulated UK market data.
# Shows radar chart, benchmark gauges, comparison cards,
# Competitiveness Score and Benchmark Score.
# ═══════════════════════════════════════════════════════════════════════════════

# ── Simulated UK market data by seniority tier ───────────────────────────────
BENCHMARK_MARKET_DATA = {
    "Junior (£20k–£35k)": {
        "Salary":                {"market_median": 27000, "market_25th": 22000, "market_75th": 33000, "unit": "£"},
        "Notice Period":         {"market_median": 4,     "market_25th": 2,     "market_75th": 8,     "unit": "weeks"},
        "Annual Leave":          {"market_median": 25,    "market_25th": 22,    "market_75th": 28,    "unit": "days"},
        "Pension Contribution":  {"market_median": 5,     "market_25th": 3,     "market_75th": 8,     "unit": "%"},
        "Probation Period":      {"market_median": 3,     "market_25th": 1,     "market_75th": 6,     "unit": "months"},
        "Restrictive Covenants": {"market_median": 3,     "market_25th": 1,     "market_75th": 6,     "unit": "months"},
    },
    "Mid-Level (£35k–£60k)": {
        "Salary":                {"market_median": 47000, "market_25th": 38000, "market_75th": 56000, "unit": "£"},
        "Notice Period":         {"market_median": 8,     "market_25th": 4,     "market_75th": 12,    "unit": "weeks"},
        "Annual Leave":          {"market_median": 28,    "market_25th": 25,    "market_75th": 33,    "unit": "days"},
        "Pension Contribution":  {"market_median": 7,     "market_25th": 5,     "market_75th": 10,    "unit": "%"},
        "Probation Period":      {"market_median": 3,     "market_25th": 3,     "market_75th": 6,     "unit": "months"},
        "Restrictive Covenants": {"market_median": 6,     "market_25th": 3,     "market_75th": 12,    "unit": "months"},
    },
    "Senior (£60k–£100k)": {
        "Salary":                {"market_median": 78000, "market_25th": 63000, "market_75th": 95000, "unit": "£"},
        "Notice Period":         {"market_median": 12,    "market_25th": 8,     "market_75th": 16,    "unit": "weeks"},
        "Annual Leave":          {"market_median": 30,    "market_25th": 28,    "market_75th": 35,    "unit": "days"},
        "Pension Contribution":  {"market_median": 10,    "market_25th": 7,     "market_75th": 15,    "unit": "%"},
        "Probation Period":      {"market_median": 6,     "market_25th": 3,     "market_75th": 6,     "unit": "months"},
        "Restrictive Covenants": {"market_median": 9,     "market_25th": 6,     "market_75th": 12,    "unit": "months"},
    },
    "Executive (£100k+)": {
        "Salary":                {"market_median": 145000,"market_25th": 105000,"market_75th": 200000,"unit": "£"},
        "Notice Period":         {"market_median": 20,    "market_25th": 12,    "market_75th": 26,    "unit": "weeks"},
        "Annual Leave":          {"market_median": 33,    "market_25th": 30,    "market_75th": 38,    "unit": "days"},
        "Pension Contribution":  {"market_median": 12,    "market_25th": 10,    "market_75th": 18,    "unit": "%"},
        "Probation Period":      {"market_median": 6,     "market_25th": 3,     "market_75th": 6,     "unit": "months"},
        "Restrictive Covenants": {"market_median": 12,    "market_25th": 9,     "market_75th": 18,    "unit": "months"},
    },
}


def _extract_benchmark_values(text: str) -> dict:
    """
    Extract numeric contract values for benchmarking.
    Returns dict of {clause: extracted_numeric_value_or_None}.
    """
    vals = {}

    # Salary
    sm = re.search(r"(?:salary|remuneration|basic pay)[^\n£]*£\s*([\d,]+)", text, re.I)
    if not sm:
        sm = re.search(r"£\s*([\d,]+)\s*(?:per annum|p\.?a\.?|annually)", text, re.I)
    if sm:
        try:
            vals["Salary"] = int(sm.group(1).replace(",", ""))
        except Exception:
            vals["Salary"] = None
    else:
        vals["Salary"] = None

    # Notice Period (convert to weeks)
    nm = re.search(r"notice[^\n]*?(\d+)\s*(week|month|day)", text, re.I)
    if nm:
        n_val = int(nm.group(1))
        n_unit = nm.group(2).lower()
        if "month" in n_unit:
            vals["Notice Period"] = n_val * 4
        elif "day" in n_unit:
            vals["Notice Period"] = round(n_val / 5, 1)
        else:
            vals["Notice Period"] = n_val
    else:
        vals["Notice Period"] = None

    # Annual Leave
    lm = re.search(r"(\d+)\s*(?:days?|working days?)\s*(?:annual|holiday|paid leave)", text, re.I)
    if not lm:
        lm = re.search(r"(?:annual leave|holiday entitlement)[^\n]*?(\d+)\s*days?", text, re.I)
    vals["Annual Leave"] = int(lm.group(1)) if lm else None

    # Pension
    pm = re.search(r"pension[^\n]*?(\d+\.?\d*)\s*%", text, re.I)
    vals["Pension Contribution"] = float(pm.group(1)) if pm else None

    # Probation (months)
    prob = re.search(r"probati(?:on|onary)[^\n]*?(\d+)\s*(week|month)", text, re.I)
    if prob:
        p_val = int(prob.group(1))
        p_unit = prob.group(2).lower()
        vals["Probation Period"] = p_val if "month" in p_unit else round(p_val / 4, 1)
    else:
        vals["Probation Period"] = None

    # Restrictive Covenants duration (months)
    rc = re.search(r"(?:non-compete|restrictive covenant|restraint)[^\n]*?(\d+)\s*(month|week|year)", text, re.I)
    if rc:
        rc_val = int(rc.group(1))
        rc_unit = rc.group(2).lower()
        if "year" in rc_unit:
            vals["Restrictive Covenants"] = rc_val * 12
        elif "week" in rc_unit:
            vals["Restrictive Covenants"] = round(rc_val / 4, 1)
        else:
            vals["Restrictive Covenants"] = rc_val
    else:
        vals["Restrictive Covenants"] = None

    return vals


def _compute_clause_score(contract_val, market_25th, market_median, market_75th, clause_name: str) -> float:
    """
    Return a 0–100 competitiveness score for a single clause.
    Higher = better for employee (more competitive offer).
    For probation & restrictive covenants: lower contract value = better.
    """
    if contract_val is None:
        return 35.0  # penalise missing clause

    reverse = clause_name in ("Probation Period", "Restrictive Covenants")
    if reverse:
        # Shorter = better for employee
        if contract_val <= market_25th:
            return 90.0
        elif contract_val <= market_median:
            return 70.0
        elif contract_val <= market_75th:
            return 50.0
        else:
            return 25.0
    else:
        if contract_val >= market_75th:
            return 90.0
        elif contract_val >= market_median:
            return 70.0
        elif contract_val >= market_25th:
            return 50.0
        else:
            return 25.0


def compute_benchmark_scores(contract_vals: dict, tier: str) -> dict:
    """
    Compare contract values against market benchmarks for the given tier.
    Returns dict with per-clause scores, competitiveness_score, benchmark_score.
    """
    market = BENCHMARK_MARKET_DATA[tier]
    clause_scores = {}
    for clause, mdata in market.items():
        cval = contract_vals.get(clause)
        score = _compute_clause_score(
            cval, mdata["market_25th"], mdata["market_median"], mdata["market_75th"], clause
        )
        clause_scores[clause] = {
            "contract_value":  cval,
            "market_25th":     mdata["market_25th"],
            "market_median":   mdata["market_median"],
            "market_75th":     mdata["market_75th"],
            "unit":            mdata["unit"],
            "score":           score,
        }

    scores_only = [v["score"] for v in clause_scores.values()]
    competitiveness_score = int(sum(scores_only) / len(scores_only))
    benchmark_score = int(
        sum(1 for v in clause_scores.values() if v["score"] >= 70) / len(clause_scores) * 100
    )
    return {
        "clause_scores":         clause_scores,
        "competitiveness_score": competitiveness_score,
        "benchmark_score":       benchmark_score,
    }


def _benchmark_gauge(title: str, value: int, color: str) -> go.Figure:
    """Small Plotly gauge for benchmark scores."""
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=value,
        number={"suffix": "", "font": {"size": 28, "color": color, "family": "Cormorant Garamond"}},
        gauge={
            "axis": {"range": [0, 100], "tickwidth": 1, "tickcolor": "#1E3448",
                     "tickfont": {"color": "#607D99", "size": 8}},
            "bar": {"color": color, "thickness": 0.22},
            "bgcolor": "#0D1F30",
            "borderwidth": 0,
            "steps": [
                {"range": [0,  39],  "color": "rgba(231,76,60,0.12)"},
                {"range": [39, 69],  "color": "rgba(243,156,18,0.12)"},
                {"range": [69, 100], "color": "rgba(39,174,96,0.12)"},
            ],
        },
        title={"text": title, "font": {"color": "#607D99", "size": 11}},
    ))
    fig.update_layout(
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        height=190, margin=dict(l=15, r=15, t=38, b=5),
    )
    return fig


def _benchmark_radar(clause_scores: dict) -> go.Figure:
    """Plotly radar chart: contract vs market median."""
    clauses = list(clause_scores.keys())
    # Normalise contract and market to 0-100 scale for radar display
    contract_pct = []
    market_pct   = []
    for c, v in clause_scores.items():
        m75 = v["market_75th"] or 1
        cval = v["contract_value"] if v["contract_value"] is not None else v["market_25th"] * 0.7
        # For reverse clauses, invert
        if c in ("Probation Period", "Restrictive Covenants"):
            max_val = max(v["market_75th"] * 1.5, cval * 1.1, 1)
            contract_pct.append(round(max(0, min(100, (1 - cval / max_val) * 100))))
            market_pct.append(round(max(0, min(100, (1 - v["market_median"] / max_val) * 100))))
        else:
            max_val = max(v["market_75th"] * 1.3, cval * 1.1, 1)
            contract_pct.append(round(max(0, min(100, cval / max_val * 100))))
            market_pct.append(round(max(0, min(100, v["market_median"] / max_val * 100))))

    # Close the polygon
    clauses_closed      = clauses + [clauses[0]]
    contract_pct_closed = contract_pct + [contract_pct[0]]
    market_pct_closed   = market_pct   + [market_pct[0]]

    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(
        r=market_pct_closed,
        theta=clauses_closed,
        fill="toself",
        name="Market Median",
        line=dict(color="#4A90D9", width=2, dash="dash"),
        fillcolor="rgba(74,144,217,0.08)",
        hovertemplate="%{theta}<br>Market Median: %{r:.0f}%<extra></extra>",
    ))
    fig.add_trace(go.Scatterpolar(
        r=contract_pct_closed,
        theta=clauses_closed,
        fill="toself",
        name="This Contract",
        line=dict(color="#C9A84C", width=2.5),
        fillcolor="rgba(201,168,76,0.12)",
        hovertemplate="%{theta}<br>Contract Position: %{r:.0f}%<extra></extra>",
    ))
    fig.update_layout(
        polar=dict(
            bgcolor="rgba(0,0,0,0)",
            radialaxis=dict(
                visible=True, range=[0, 100],
                gridcolor="#1E3448", tickcolor="#607D99",
                tickfont=dict(size=9, color="#607D99"),
                linecolor="#1E3448",
            ),
            angularaxis=dict(
                tickfont=dict(size=10, color="#C9A84C"),
                gridcolor="#1E3448", linecolor="#1E3448",
            ),
        ),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        height=400,
        margin=dict(l=60, r=60, t=40, b=40),
        legend=dict(
            font=dict(color="#607D99", size=10),
            bgcolor="rgba(13,31,48,0.9)",
            bordercolor="#1E3448", borderwidth=1,
            orientation="h", yanchor="bottom", y=-0.15, xanchor="center", x=0.5,
        ),
        hoverlabel=dict(bgcolor="#0D1F30", bordercolor="#C9A84C", font=dict(color="#E8E4DA")),
    )
    return fig


def render_clause_benchmarking(text: str):
    """Feature 15 — Clause Benchmarking Engine."""
    h("""<div style="font-size:0.72rem;color:#607D99;margin-bottom:1rem;letter-spacing:0.04em;">
    Benchmark contract clauses against UK market data.
    Compare salary, notice, leave, benefits, probation, and restrictive covenants
    against median market standards for your seniority tier.</div>""")

    if not text.strip():
        st.info("Upload or paste a contract to run clause benchmarking.")
        return

    # ── Tier selector ──────────────────────────────────────────────────────────
    tier_col, _ = st.columns([1, 2])
    with tier_col:
        h('<div style="font-size:0.7rem;font-weight:700;letter-spacing:0.12em;text-transform:uppercase;color:#607D99;margin-bottom:0.4rem;">Market Tier</div>')
        tier = st.selectbox(
            "Tier", list(BENCHMARK_MARKET_DATA.keys()),
            index=1, label_visibility="collapsed", key="bench_tier"
        )

    # ── Extract + compute ──────────────────────────────────────────────────────
    contract_vals = _extract_benchmark_values(text)
    bench = compute_benchmark_scores(contract_vals, tier)
    clause_scores = bench["clause_scores"]
    comp_score    = bench["competitiveness_score"]
    bench_score   = bench["benchmark_score"]

    # ── Summary gauges ─────────────────────────────────────────────────────────
    section_title("Overall Benchmark Scores", "🎯")
    g1, g2, g3 = st.columns([1, 1, 2])

    comp_color  = "#27AE60" if comp_score >= 70 else "#F39C12" if comp_score >= 45 else "#E74C3C"
    bench_color = "#27AE60" if bench_score >= 70 else "#F39C12" if bench_score >= 45 else "#E74C3C"

    with g1:
        fig_cs = _benchmark_gauge("Competitiveness Score", comp_score, comp_color)
        st.plotly_chart(fig_cs, use_container_width=True, config={"displayModeBar": False})
        level = "Above Market" if comp_score >= 70 else "At Market" if comp_score >= 45 else "Below Market"
        h(f"""<div style="text-align:center;background:#071219;border:1px solid #1E3448;
        border-radius:6px;padding:0.35rem;margin-top:-0.5rem;font-size:0.7rem;">
        <span style="color:{comp_color};font-weight:700;">{level}</span></div>""")

    with g2:
        fig_bs = _benchmark_gauge("Benchmark Score", bench_score, bench_color)
        st.plotly_chart(fig_bs, use_container_width=True, config={"displayModeBar": False})
        clauses_at_market = sum(1 for v in clause_scores.values() if v["score"] >= 70)
        h(f"""<div style="text-align:center;background:#071219;border:1px solid #1E3448;
        border-radius:6px;padding:0.35rem;margin-top:-0.5rem;font-size:0.7rem;">
        <span style="color:{bench_color};font-weight:700;">{clauses_at_market}/{len(clause_scores)} clauses at/above market</span></div>""")

    with g3:
        # Mini summary cards for each clause
        h('<div style="font-size:0.68rem;letter-spacing:0.1em;text-transform:uppercase;color:#607D99;margin-bottom:0.5rem;font-weight:700;">Clause Scores at a Glance</div>')
        mini_cols = st.columns(3)
        for idx, (clause, data) in enumerate(clause_scores.items()):
            sc = data["score"]
            sc_color = "#27AE60" if sc >= 70 else "#F39C12" if sc >= 45 else "#E74C3C"
            cval = data["contract_value"]
            cval_str = f"{data['unit']}{cval:,.0f}" if cval is not None and data['unit'] == '£' else \
                       f"{cval} {data['unit']}" if cval is not None else "Not found"
            with mini_cols[idx % 3]:
                h(f"""<div style="background:#0D1F30;border:1px solid #1E3448;
                border-top:2px solid {sc_color};border-radius:8px;
                padding:0.55rem 0.7rem;margin-bottom:0.4rem;">
                    <div style="font-size:0.65rem;font-weight:700;letter-spacing:0.08em;
                    text-transform:uppercase;color:{sc_color};margin-bottom:0.2rem;">
                    {clause}</div>
                    <div style="display:flex;justify-content:space-between;align-items:center;">
                        <span style="font-size:0.75rem;color:#B0C4D8;">{cval_str}</span>
                        <span style="font-family:'Cormorant Garamond',serif;font-size:1.2rem;
                        font-weight:700;color:{sc_color};">{int(sc)}</span>
                    </div>
                </div>""")

    # ── Radar chart ────────────────────────────────────────────────────────────
    section_title("Contract vs Market — Radar Chart", "📡")
    radar_col, legend_col = st.columns([1.6, 1])

    with radar_col:
        fig_radar = _benchmark_radar(clause_scores)
        st.plotly_chart(fig_radar, use_container_width=True, config={"displayModeBar": False})

    with legend_col:
        h("""<div style="font-size:0.68rem;letter-spacing:0.1em;text-transform:uppercase;
        color:#607D99;margin-bottom:0.6rem;font-weight:700;">How to Read This Chart</div>""")
        h("""<div style="font-size:0.78rem;color:#B0C4D8;line-height:1.75;">
        <span style="color:#C9A84C;font-weight:700;">Gold line</span> = This contract.<br>
        <span style="color:#4A90D9;font-weight:700;">Blue dashed</span> = Market median.<br><br>
        A larger gold area = more competitive contract for the employee.<br><br>
        Salary, Notice, Leave, Pension: higher = better.<br>
        Probation &amp; Restrictive Covenants: shorter = better for employee.
        </div>""")

        h("""<div style="margin-top:1rem;font-size:0.68rem;letter-spacing:0.1em;text-transform:uppercase;
        color:#607D99;margin-bottom:0.4rem;font-weight:700;">Score Key</div>""")
        for label, color, threshold in [
            ("90 — Above 75th Percentile", "#27AE60", ""),
            ("70 — At/Above Median",        "#27AE60", ""),
            ("50 — Between 25th & Median",  "#F39C12", ""),
            ("25 — Below 25th Percentile",  "#E74C3C", ""),
            ("35 — Missing Clause",          "#E74C3C", ""),
        ]:
            h(f"""<div style="display:flex;align-items:center;gap:0.5rem;
            font-size:0.72rem;color:#B0C4D8;margin-bottom:0.2rem;">
            <span style="width:8px;height:8px;border-radius:50%;background:{color};
            display:inline-block;flex-shrink:0;"></span>{label}</div>""")

    # ── Detailed comparison cards ──────────────────────────────────────────────
    section_title("Clause-by-Clause Comparison Cards", "📋")

    card_cols = st.columns(2)
    icons_map = {
        "Salary":                "💷",
        "Notice Period":         "📅",
        "Annual Leave":          "🏖️",
        "Pension Contribution":  "🏦",
        "Probation Period":      "⏱️",
        "Restrictive Covenants": "🔒",
    }
    reverse_clauses = {"Probation Period", "Restrictive Covenants"}

    for idx, (clause, data) in enumerate(clause_scores.items()):
        sc       = data["score"]
        sc_color = "#27AE60" if sc >= 70 else "#F39C12" if sc >= 45 else "#E74C3C"
        cval     = data["contract_value"]
        unit     = data["unit"]
        icon     = icons_map.get(clause, "📋")

        def fmt_val(v, u):
            if v is None:
                return "Not Found"
            return f"£{v:,.0f}" if u == "£" else f"{v} {u}"

        cval_str    = fmt_val(cval, unit)
        median_str  = fmt_val(data["market_median"],  unit)
        p25_str     = fmt_val(data["market_25th"],    unit)
        p75_str     = fmt_val(data["market_75th"],    unit)

        # Verdict
        if cval is None:
            verdict = "Missing — unable to benchmark"
            verdict_icon = "⚠️"
        elif sc >= 70:
            verdict = "Above / At Market — competitive"
            verdict_icon = "✅"
        elif sc >= 45:
            verdict = "Slightly Below Market — review recommended"
            verdict_icon = "🟡"
        else:
            verdict = "Below Market — consider renegotiating"
            verdict_icon = "🔴"

        reverse_note = " (shorter = better for employee)" if clause in reverse_clauses else ""

        with card_cols[idx % 2]:
            h(f"""<div style="background:linear-gradient(135deg,#0D1F30,#0B1825);
            border:1px solid #1E3448;border-left:4px solid {sc_color};
            border-radius:0 12px 12px 0;padding:1rem 1.1rem;margin-bottom:0.8rem;">

                <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:0.6rem;">
                    <div>
                        <span style="font-size:1rem;margin-right:0.4rem;">{icon}</span>
                        <span style="font-size:0.82rem;font-weight:700;color:#E8E4DA;">{clause}</span>
                        <div style="font-size:0.65rem;color:#3A5570;margin-top:0.1rem;">{reverse_note}</div>
                    </div>
                    <div style="text-align:right;">
                        <div style="font-family:'Cormorant Garamond',serif;font-size:1.7rem;
                        font-weight:700;color:{sc_color};line-height:1;">{int(sc)}</div>
                        <div style="font-size:0.6rem;color:#607D99;">score</div>
                    </div>
                </div>

                <div style="display:grid;grid-template-columns:1fr 1fr 1fr 1fr;
                gap:0.4rem;margin-bottom:0.6rem;">
                    <div style="background:#071219;border:1px solid #1E3448;border-radius:6px;
                    padding:0.4rem;text-align:center;">
                        <div style="font-size:0.6rem;color:#C9A84C;font-weight:700;
                        letter-spacing:0.08em;text-transform:uppercase;margin-bottom:0.15rem;">Contract</div>
                        <div style="font-size:0.82rem;font-weight:700;color:{sc_color};">{cval_str}</div>
                    </div>
                    <div style="background:#071219;border:1px solid #1E3448;border-radius:6px;
                    padding:0.4rem;text-align:center;">
                        <div style="font-size:0.6rem;color:#607D99;font-weight:700;
                        letter-spacing:0.08em;text-transform:uppercase;margin-bottom:0.15rem;">25th %ile</div>
                        <div style="font-size:0.82rem;color:#607D99;">{p25_str}</div>
                    </div>
                    <div style="background:#071219;border:1px solid #1E3448;border-radius:6px;
                    padding:0.4rem;text-align:center;">
                        <div style="font-size:0.6rem;color:#4A90D9;font-weight:700;
                        letter-spacing:0.08em;text-transform:uppercase;margin-bottom:0.15rem;">Median</div>
                        <div style="font-size:0.82rem;color:#4A90D9;">{median_str}</div>
                    </div>
                    <div style="background:#071219;border:1px solid #1E3448;border-radius:6px;
                    padding:0.4rem;text-align:center;">
                        <div style="font-size:0.6rem;color:#27AE60;font-weight:700;
                        letter-spacing:0.08em;text-transform:uppercase;margin-bottom:0.15rem;">75th %ile</div>
                        <div style="font-size:0.82rem;color:#27AE60;">{p75_str}</div>
                    </div>
                </div>

                <div style="display:flex;align-items:center;gap:0.4rem;font-size:0.78rem;
                color:#B0C4D8;background:rgba({_hex_to_rgb(sc_color)},0.06);
                border:1px solid rgba({_hex_to_rgb(sc_color)},0.2);border-radius:6px;padding:0.4rem 0.6rem;">
                    <span>{verdict_icon}</span>
                    <span>{verdict}</span>
                </div>
            </div>""")


# ═══════════════════════════════════════════════════════════════════════════════
# FEATURE 16 — OFFER ACCEPTANCE PREDICTOR
# Analyses Salary, Benefits, Leave, Notice Period, Flexible Working
# to produce Offer Acceptance Probability + Gauge + Recommendations
# ═══════════════════════════════════════════════════════════════════════════════

def _parse_offer_inputs(salary: float, leave: int, pension: float,
                        notice_weeks: int, flex_score: int,
                        healthcare: bool, bonus: bool, bonus_pct: float,
                        remote_days: int, relocation: bool) -> dict:
    """
    Compute offer acceptance probability from offer inputs.
    Uses a weighted scoring model calibrated against UK job market research.
    Returns score (0-100) and per-factor breakdown.
    """
    factors = {}

    # 1. Salary vs market expectation (user-provided expected salary)
    # We score based on absolute attractiveness bands
    if   salary >= 90000: factors["Salary Level"]      = (92, "Excellent — top quartile")
    elif salary >= 60000: factors["Salary Level"]      = (80, "Strong — above market median")
    elif salary >= 40000: factors["Salary Level"]      = (65, "Competitive — around median")
    elif salary >= 28000: factors["Salary Level"]      = (50, "Below average for most roles")
    else:                 factors["Salary Level"]      = (30, "Low — likely to deter candidates")

    # 2. Annual Leave
    if   leave >= 33:     factors["Annual Leave"]      = (90, "Excellent — highly attractive")
    elif leave >= 28:     factors["Annual Leave"]      = (75, "Above market standard")
    elif leave >= 25:     factors["Annual Leave"]      = (60, "Standard market offering")
    elif leave >= 20:     factors["Annual Leave"]      = (45, "Statutory minimum only")
    else:                 factors["Annual Leave"]      = (25, "Below statutory — review required")

    # 3. Pension
    if   pension >= 12:   factors["Pension"]           = (88, "Excellent employer contribution")
    elif pension >= 8:    factors["Pension"]           = (74, "Above market standard")
    elif pension >= 5:    factors["Pension"]           = (60, "Market standard")
    elif pension >= 3:    factors["Pension"]           = (45, "Below average contribution")
    else:                 factors["Pension"]           = (28, "Low — unattractive to candidates")

    # 4. Notice Period (candidate perspective — shorter notice is more flexible)
    if   notice_weeks <= 4:  factors["Notice Period"] = (80, "Short — easy to transition")
    elif notice_weeks <= 8:  factors["Notice Period"] = (68, "Standard — acceptable")
    elif notice_weeks <= 12: factors["Notice Period"] = (52, "Long — may deter candidates")
    elif notice_weeks <= 20: factors["Notice Period"] = (38, "Very long — significant deterrent")
    else:                    factors["Notice Period"] = (22, "Excessive — high candidate risk")

    # 5. Flexible working
    if   flex_score >= 4: factors["Flexible Working"] = (90, "Hybrid/remote — highly attractive")
    elif flex_score >= 3: factors["Flexible Working"] = (72, "Good flexibility offering")
    elif flex_score >= 2: factors["Flexible Working"] = (55, "Some flexibility offered")
    elif flex_score >= 1: factors["Flexible Working"] = (38, "Limited flexibility")
    else:                 factors["Flexible Working"] = (20, "Office-only — strong deterrent")

    # 6. Healthcare
    factors["Healthcare"]   = (82, "Private medical — strong attractor") if healthcare \
                             else (45, "No private healthcare included")

    # 7. Bonus scheme
    if bonus:
        if   bonus_pct >= 20: factors["Bonus Scheme"] = (90, f"{bonus_pct:.0f}% — highly attractive")
        elif bonus_pct >= 10: factors["Bonus Scheme"] = (75, f"{bonus_pct:.0f}% — competitive")
        elif bonus_pct >=  5: factors["Bonus Scheme"] = (60, f"{bonus_pct:.0f}% — standard")
        else:                 factors["Bonus Scheme"] = (45, f"{bonus_pct:.0f}% — modest")
    else:
        factors["Bonus Scheme"] = (35, "No bonus scheme")

    # 8. Remote working days
    if   remote_days >= 4: factors["Remote Working"] = (92, "Fully/mostly remote — top attractor")
    elif remote_days >= 3: factors["Remote Working"] = (78, "3 days remote — very attractive")
    elif remote_days >= 2: factors["Remote Working"] = (62, "2 days remote — attractive")
    elif remote_days >= 1: factors["Remote Working"] = (45, "1 day remote — limited appeal")
    else:                  factors["Remote Working"] = (22, "No remote working — deterrent")

    # 9. Relocation penalty
    if relocation:
        factors["Relocation Required"] = (35, "Relocation required — significant deterrent")
    else:
        factors["Relocation Required"] = (78, "No relocation needed")

    # Weighted average (salary and flex weighted more heavily)
    weights = {
        "Salary Level":       0.22,
        "Annual Leave":       0.09,
        "Pension":            0.08,
        "Notice Period":      0.10,
        "Flexible Working":   0.12,
        "Healthcare":         0.08,
        "Bonus Scheme":       0.10,
        "Remote Working":     0.13,
        "Relocation Required":0.08,
    }
    total_weight = sum(weights.values())
    probability = sum(factors[k][0] * weights[k] for k in factors) / total_weight
    probability = max(5, min(98, int(probability)))

    return {"probability": probability, "factors": factors}


def _acceptance_gauge(probability: int) -> go.Figure:
    """Large Plotly gauge for offer acceptance probability."""
    color = "#27AE60" if probability >= 70 else "#F39C12" if probability >= 45 else "#E74C3C"
    label = "Likely Accept" if probability >= 70 else "Uncertain" if probability >= 45 else "Likely Decline"
    fig = go.Figure(go.Indicator(
        mode="gauge+number+delta",
        value=probability,
        number={
            "suffix": "%",
            "font": {"size": 44, "color": color, "family": "Cormorant Garamond"},
        },
        delta={"reference": 65, "valueformat": ".0f",
               "increasing": {"color": "#27AE60"}, "decreasing": {"color": "#E74C3C"}},
        gauge={
            "axis": {"range": [0, 100], "tickwidth": 1, "tickcolor": "#1E3448",
                     "tickfont": {"color": "#607D99", "size": 10},
                     "tickvals": [0, 20, 40, 60, 80, 100],
                     "ticktext": ["0", "20", "40", "60", "80", "100"]},
            "bar": {"color": color, "thickness": 0.3},
            "bgcolor": "#0D1F30",
            "borderwidth": 0,
            "steps": [
                {"range": [0, 45],  "color": "rgba(231,76,60,0.12)"},
                {"range": [45, 70], "color": "rgba(243,156,18,0.10)"},
                {"range": [70, 100],"color": "rgba(39,174,96,0.12)"},
            ],
            "threshold": {
                "line": {"color": "#C9A84C", "width": 3},
                "thickness": 0.8,
                "value": 65
            },
        },
        title={
            "text": f"Offer Acceptance Probability<br><span style='font-size:12px;color:#607D99'>{label}</span>",
            "font": {"color": "#E8E4DA", "size": 13}
        },
    ))
    fig.update_layout(
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        height=300, margin=dict(l=30, r=30, t=50, b=20),
    )
    return fig


def _generate_acceptance_recommendations(factors: dict, probability: int) -> list:
    """Generate actionable recommendations based on weak scoring factors."""
    recs = []
    sorted_factors = sorted(factors.items(), key=lambda x: x[1][0])

    for factor_name, (score, _) in sorted_factors[:4]:
        if score < 50:
            if factor_name == "Salary Level":
                recs.append("💷 Consider increasing the base salary — it is the primary driver of acceptance. A 10–15% increase would substantially improve competitiveness.")
            elif factor_name == "Annual Leave":
                recs.append("🏖️ Annual leave is below market standard. Increasing to 28+ days would align with market and is a low-cost benefit with high perceived value.")
            elif factor_name == "Pension":
                recs.append("🏦 Pension contribution is below market. Consider increasing employer contribution to at least 5–8% to meet candidate expectations.")
            elif factor_name == "Notice Period":
                recs.append("📅 Long notice periods deter candidates. Consider reducing or adding a mutual flexibility clause to the notice provision.")
            elif factor_name == "Flexible Working":
                recs.append("🏠 Offering formal flexible working arrangements would significantly improve offer attractiveness without added cost.")
            elif factor_name == "Healthcare":
                recs.append("🏥 Adding private healthcare (BUPA/AXA) is a high-impact benefit that materially improves offer acceptance rates.")
            elif factor_name == "Bonus Scheme":
                recs.append("🎯 A structured bonus scheme, even at 5–10%, signals performance culture and increases offer attractiveness.")
            elif factor_name == "Remote Working":
                recs.append("💻 Offering even 2 days remote per week would significantly improve candidate appeal and acceptance rates.")
            elif factor_name == "Relocation Required":
                recs.append("📦 Relocation is a strong deterrent. Offering a relocation allowance or hybrid start arrangement could mitigate resistance.")

    if probability >= 70:
        recs.append("✅ This offer is competitive. Ensure offer letters are dispatched promptly — speed of offer is itself a positive signal to candidates.")
    elif probability >= 50:
        recs.append("⚠️ The offer is borderline. Address the lowest-scoring factors above before extending to maximise acceptance probability.")
    else:
        recs.append("🔴 This offer is likely to be declined. A significant package review is recommended before extending formally.")

    return recs


def render_offer_acceptance_predictor():
    """Feature 16 — Offer Acceptance Predictor."""
    h("""<div style="font-size:0.72rem;color:#607D99;margin-bottom:1rem;letter-spacing:0.04em;">
    Model the likelihood a candidate will accept this offer. Adjust offer parameters to
    find the optimal package. Based on UK labour market research and weighted scoring.</div>""")

    # ── Input form ─────────────────────────────────────────────────────────────
    section_title("Offer Parameters", "🎛️")

    col1, col2, col3 = st.columns(3)

    with col1:
        h('<div style="font-size:0.65rem;letter-spacing:0.1em;text-transform:uppercase;color:#607D99;margin-bottom:0.3rem;">Base Salary (£)</div>')
        salary = st.number_input("Salary", min_value=15000, max_value=500000,
                                  value=45000, step=1000, label_visibility="collapsed",
                                  key="oap_salary")
        h('<div style="font-size:0.65rem;letter-spacing:0.1em;text-transform:uppercase;color:#607D99;margin:0.6rem 0 0.3rem 0;">Annual Leave (days)</div>')
        leave = st.slider("Leave", 20, 45, 28, label_visibility="collapsed", key="oap_leave")
        h('<div style="font-size:0.65rem;letter-spacing:0.1em;text-transform:uppercase;color:#607D99;margin:0.6rem 0 0.3rem 0;">Employer Pension (%)</div>')
        pension = st.slider("Pension", 0.0, 25.0, 6.0, 0.5, label_visibility="collapsed", key="oap_pension")

    with col2:
        h('<div style="font-size:0.65rem;letter-spacing:0.1em;text-transform:uppercase;color:#607D99;margin-bottom:0.3rem;">Notice Period (weeks)</div>')
        notice_weeks = st.slider("Notice", 1, 52, 8, label_visibility="collapsed", key="oap_notice")
        h('<div style="font-size:0.65rem;letter-spacing:0.1em;text-transform:uppercase;color:#607D99;margin:0.6rem 0 0.3rem 0;">Flexible Working (0 = none, 5 = fully flexible)</div>')
        flex_score = st.slider("Flex", 0, 5, 3, label_visibility="collapsed", key="oap_flex")
        h('<div style="font-size:0.65rem;letter-spacing:0.1em;text-transform:uppercase;color:#607D99;margin:0.6rem 0 0.3rem 0;">Remote Days per Week</div>')
        remote_days = st.slider("Remote", 0, 5, 2, label_visibility="collapsed", key="oap_remote")

    with col3:
        h('<div style="font-size:0.65rem;letter-spacing:0.1em;text-transform:uppercase;color:#607D99;margin-bottom:0.3rem;">Benefits Included</div>')
        healthcare = st.toggle("🏥 Private Healthcare", value=False, key="oap_health")
        bonus = st.toggle("🎯 Bonus Scheme", value=True, key="oap_bonus")
        bonus_pct = 0.0
        if bonus:
            h('<div style="font-size:0.65rem;letter-spacing:0.1em;text-transform:uppercase;color:#607D99;margin:0.4rem 0 0.2rem 0;">Bonus % of Salary</div>')
            bonus_pct = st.slider("Bonus %", 0.0, 50.0, 10.0, 1.0,
                                   label_visibility="collapsed", key="oap_bonus_pct")
        relocation = st.toggle("📦 Relocation Required", value=False, key="oap_relocation")

    # ── Compute ────────────────────────────────────────────────────────────────
    result = _parse_offer_inputs(
        salary=salary, leave=leave, pension=pension,
        notice_weeks=notice_weeks, flex_score=flex_score,
        healthcare=healthcare, bonus=bonus, bonus_pct=bonus_pct,
        remote_days=remote_days, relocation=relocation,
    )
    probability = result["probability"]
    factors     = result["factors"]

    # ── Results layout ─────────────────────────────────────────────────────────
    section_title("Acceptance Probability", "🎯")
    gauge_col, breakdown_col = st.columns([1, 1.4])

    with gauge_col:
        fig_gauge = _acceptance_gauge(probability)
        st.plotly_chart(fig_gauge, use_container_width=True, config={"displayModeBar": False})

        prob_color = "#27AE60" if probability >= 70 else "#F39C12" if probability >= 45 else "#E74C3C"
        verdict = "Likely to Accept" if probability >= 70 else "Uncertain" if probability >= 45 else "Likely to Decline"
        h(f"""<div style="background:rgba({_hex_to_rgb(prob_color)},0.08);
        border:1px solid rgba({_hex_to_rgb(prob_color)},0.3);
        border-left:5px solid {prob_color};border-radius:0 10px 10px 0;
        padding:0.8rem 1rem;text-align:center;">
            <div style="font-size:0.65rem;letter-spacing:0.12em;text-transform:uppercase;
            color:#607D99;margin-bottom:0.2rem;">Verdict</div>
            <div style="font-size:1.1rem;font-weight:700;color:{prob_color};">{verdict}</div>
            <div style="font-size:0.72rem;color:#607D99;margin-top:0.3rem;">
            Gold line = 65% target threshold</div>
        </div>""")

    with breakdown_col:
        h('<div style="font-size:0.7rem;font-weight:700;letter-spacing:0.12em;text-transform:uppercase;color:#607D99;margin-bottom:0.6rem;">Factor Breakdown</div>')

        sorted_factors = sorted(factors.items(), key=lambda x: x[1][0], reverse=True)
        for factor_name, (score, desc) in sorted_factors:
            bar_color = "#27AE60" if score >= 70 else "#F39C12" if score >= 45 else "#E74C3C"
            h(f"""<div style="margin-bottom:0.6rem;">
                <div style="display:flex;justify-content:space-between;
                align-items:center;font-size:0.79rem;margin-bottom:0.15rem;">
                    <span style="color:#E8E4DA;font-weight:500;">{factor_name}</span>
                    <div style="display:flex;align-items:center;gap:0.5rem;">
                        <span style="font-size:0.68rem;color:#607D99;font-style:italic;
                        max-width:160px;text-align:right;overflow:hidden;
                        text-overflow:ellipsis;white-space:nowrap;">{desc}</span>
                        <span style="color:{bar_color};font-weight:700;min-width:28px;
                        text-align:right;">{score}</span>
                    </div>
                </div>
                <div style="background:#1E3448;border-radius:4px;height:6px;">
                    <div style="width:{score}%;background:linear-gradient(90deg,{bar_color},{bar_color}99);
                    height:100%;border-radius:4px;transition:width 0.4s ease;"></div>
                </div>
            </div>""")

    # ── Recommendations ────────────────────────────────────────────────────────
    section_title("Recommendations to Improve Offer Acceptance", "💡")
    recs = _generate_acceptance_recommendations(factors, probability)
    for rec in recs:
        rec_color = "#27AE60" if rec.startswith("✅") else "#F39C12" if rec.startswith("⚠️") else \
                    "#E74C3C" if rec.startswith("🔴") else "#C9A84C"
        h(f"""<div style="background:rgba({_hex_to_rgb(rec_color)},0.06);
        border:1px solid rgba({_hex_to_rgb(rec_color)},0.25);
        border-left:4px solid {rec_color};border-radius:0 8px 8px 0;
        padding:0.7rem 1rem;margin-bottom:0.5rem;font-size:0.82rem;
        color:#E8E4DA;line-height:1.6;">{rec}</div>""")

    # ── What-if scenario note ──────────────────────────────────────────────────
    h("""<div style="background:#0D1F30;border:1px solid #1E3448;border-radius:10px;
    padding:0.8rem 1.1rem;margin-top:1rem;">
        <div style="font-size:0.68rem;font-weight:700;letter-spacing:0.12em;text-transform:uppercase;
        color:#3A5570;margin-bottom:0.3rem;">💡 What-If Modelling</div>
        <div style="font-size:0.78rem;color:#607D99;line-height:1.6;">
        Adjust the sliders above in real time to model different offer scenarios.
        The probability and factor breakdown update instantly.
        Gold threshold line = 65% — the typical minimum acceptance probability
        for a competitive offer in the UK market.
        </div>
    </div>""")


# ═══════════════════════════════════════════════════════════════════════════════
# FEATURE 17 — EMPLOYMENT TRIBUNAL RISK PREDICTOR
# Predicts 5 tribunal risk categories + composite Tribunal Exposure Score
# ═══════════════════════════════════════════════════════════════════════════════

def compute_tribunal_risks(text: str, result: dict) -> dict:
    """
    Analyse contract text and API result to compute 5 tribunal risk scores
    and an overall Tribunal Exposure Score (0-100).
    Returns dict with individual risks and composite score.
    """
    text_l = (text or "").lower()
    clauses = result.get("clauses", {}) if result else {}
    base_risk = result.get("risk_score", 50) if result else 50

    def clause_missing(keys):
        return not any(any(k in c.lower() for k in keys) for c in clauses if clauses.get(c))

    # ── Unfair Dismissal Risk ──────────────────────────────────────────────────
    ud_score = 30
    if clause_missing(["notice", "terminat"]):           ud_score += 25
    if clause_missing(["disciplin"]):                    ud_score += 20
    if clause_missing(["grievance"]):                    ud_score += 15
    if "without cause" in text_l or "at will" in text_l: ud_score += 20
    if re.search(r"probat\w+\s+period", text_l):         ud_score = max(ud_score - 10, 10)
    ud_score = min(int(ud_score + base_risk * 0.15), 100)

    # ── Discrimination Risk ────────────────────────────────────────────────────
    disc_score = 20
    if clause_missing(["equal", "discriminat", "diversity"]): disc_score += 30
    if clause_missing(["harassment"]):                         disc_score += 20
    if not re.search(r"equality act|protected characteristic", text_l): disc_score += 20
    disc_score = min(int(disc_score + base_risk * 0.1), 100)

    # ── GDPR Risk ─────────────────────────────────────────────────────────────
    gdpr_score = 20
    if clause_missing(["gdpr", "data protection", "personal data"]): gdpr_score += 35
    if not re.search(r"ico|information commissioner|lawful basis", text_l): gdpr_score += 20
    if not re.search(r"data retention|subject access", text_l):       gdpr_score += 15
    gdpr_score = min(int(gdpr_score + base_risk * 0.1), 100)

    # ── Wage Dispute Risk ─────────────────────────────────────────────────────
    wage_score = 20
    if clause_missing(["salary", "remuneration", "wage", "pay"]):        wage_score += 30
    if clause_missing(["overtime", "working time"]):                      wage_score += 15
    if not re.search(r"national (minimum|living) wage|nmw|nlw", text_l): wage_score += 20
    if not re.search(r"pay review|salary review", text_l):               wage_score += 10
    wage_score = min(int(wage_score + base_risk * 0.1), 100)

    # ── Contract Breach Risk ──────────────────────────────────────────────────
    breach_score = 20
    if clause_missing(["confidential", "non-disclosur", "nda"]):  breach_score += 20
    if clause_missing(["intellectual property", "ip clause"]):    breach_score += 15
    if clause_missing(["restrictive covenant", "non-compet"]):    breach_score += 15
    if not re.search(r"governing law|jurisdiction|dispute resolution", text_l): breach_score += 20
    breach_score = min(int(breach_score + base_risk * 0.1), 100)

    # ── Tribunal Exposure Score (weighted composite) ───────────────────────────
    exposure = int(
        ud_score   * 0.30 +
        disc_score * 0.25 +
        gdpr_score * 0.15 +
        wage_score * 0.15 +
        breach_score * 0.15
    )

    def level(score):
        if score >= 75: return "Critical"
        if score >= 55: return "High"
        if score >= 35: return "Medium"
        return "Low"

    return {
        "Unfair Dismissal":  {"score": ud_score,     "level": level(ud_score)},
        "Discrimination":    {"score": disc_score,    "level": level(disc_score)},
        "GDPR":              {"score": gdpr_score,    "level": level(gdpr_score)},
        "Wage Dispute":      {"score": wage_score,    "level": level(wage_score)},
        "Contract Breach":   {"score": breach_score,  "level": level(breach_score)},
        "exposure_score":    exposure,
        "exposure_level":    level(exposure),
    }


_TRIBUNAL_LEVEL_CFG = {
    "Critical": {"color": "#C0392B", "bg": "rgba(192,57,43,0.12)",  "border": "rgba(192,57,43,0.4)",  "icon": "🔴", "bar": "#C0392B"},
    "High":     {"color": "#E74C3C", "bg": "rgba(231,76,60,0.09)",  "border": "rgba(231,76,60,0.35)", "icon": "🟠", "bar": "#E74C3C"},
    "Medium":   {"color": "#F39C12", "bg": "rgba(243,156,18,0.09)", "border": "rgba(243,156,18,0.35)","icon": "🟡", "bar": "#F39C12"},
    "Low":      {"color": "#27AE60", "bg": "rgba(39,174,96,0.09)",  "border": "rgba(39,174,96,0.3)",  "icon": "🟢", "bar": "#27AE60"},
}


def render_tribunal_risk_predictor(text: str, result: dict):
    """Feature 17 — Employment Tribunal Risk Predictor tab."""
    h("""<div style="font-size:0.72rem;color:#607D99;margin-bottom:1rem;letter-spacing:0.04em;">
    Predict employment tribunal exposure across five legal risk categories.
    Based on clause analysis, contract text, and UK employment law benchmarks.</div>""")

    if not text.strip():
        st.info("Upload or paste a contract to run tribunal risk prediction.")
        return

    risks = compute_tribunal_risks(text, result or {})
    exposure = risks["exposure_score"]
    exp_level = risks["exposure_level"]
    exp_cfg   = _TRIBUNAL_LEVEL_CFG[exp_level]

    # ── Tribunal Exposure Score banner ────────────────────────────────────────
    h(f"""
    <div style="background:linear-gradient(135deg,{exp_cfg['bg']},rgba(13,31,48,0.9));
    border:1.5px solid {exp_cfg['border']};border-left:6px solid {exp_cfg['color']};
    border-radius:0 12px 12px 0;padding:1.2rem 1.6rem;margin-bottom:1.5rem;
    display:flex;align-items:center;justify-content:space-between;">
        <div>
            <div style="font-size:0.6rem;letter-spacing:0.16em;text-transform:uppercase;
            color:#607D99;margin-bottom:0.25rem;">Tribunal Exposure Score</div>
            <div style="font-size:1.15rem;font-weight:700;color:{exp_cfg['color']};">
            {exp_cfg['icon']} &nbsp;{exp_level} Exposure</div>
            <div style="font-size:0.75rem;color:#607D99;margin-top:0.2rem;">
            Composite weighted score across 5 risk categories</div>
        </div>
        <div style="text-align:right;">
            <div style="font-family:'Cormorant Garamond',serif;font-size:3.5rem;
            font-weight:700;color:{exp_cfg['color']};line-height:1;">{exposure}</div>
            <div style="font-size:0.7rem;color:#607D99;">/ 100</div>
        </div>
    </div>""")

    # ── Exposure gauge ─────────────────────────────────────────────────────────
    gauge_col, detail_col = st.columns([1, 1.5])
    with gauge_col:
        fig = go.Figure(go.Indicator(
            mode="gauge+number",
            value=exposure,
            number={"font": {"size": 34, "color": exp_cfg["color"], "family": "Cormorant Garamond"}},
            gauge={
                "axis": {"range": [0, 100], "tickwidth": 1, "tickcolor": "#1E3448",
                         "tickfont": {"color": "#607D99", "size": 9}},
                "bar": {"color": exp_cfg["color"], "thickness": 0.24},
                "bgcolor": "#0D1F30", "borderwidth": 0,
                "steps": [
                    {"range": [0, 34],  "color": "rgba(39,174,96,0.1)"},
                    {"range": [34, 54], "color": "rgba(243,156,18,0.1)"},
                    {"range": [54, 74], "color": "rgba(231,76,60,0.1)"},
                    {"range": [74, 100],"color": "rgba(192,57,43,0.15)"},
                ],
                "threshold": {"line": {"color": "#C9A84C", "width": 2}, "thickness": 0.75, "value": exposure},
            },
            title={"text": "Tribunal Exposure Score", "font": {"color": "#607D99", "size": 11}},
        ))
        fig.update_layout(
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            height=220, margin=dict(l=20, r=20, t=35, b=5),
        )
        st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})

        # Legend
        h("""<div style="background:#071219;border:1px solid #1E3448;border-radius:8px;
        padding:0.6rem 0.8rem;font-size:0.7rem;color:#607D99;">
            <div style="margin-bottom:0.3rem;font-weight:700;color:#3A5570;letter-spacing:0.08em;
            text-transform:uppercase;font-size:0.62rem;">Score Bands</div>
            <div>🟢 <b style="color:#27AE60;">Low</b> &nbsp;0–34</div>
            <div>🟡 <b style="color:#F39C12;">Medium</b> &nbsp;35–54</div>
            <div>🟠 <b style="color:#E74C3C;">High</b> &nbsp;55–74</div>
            <div>🔴 <b style="color:#C0392B;">Critical</b> &nbsp;75–100</div>
        </div>""")

    # ── Individual risk category bars ─────────────────────────────────────────
    with detail_col:
        section_title("Risk Category Breakdown", "⚖️")
        risk_keys = ["Unfair Dismissal", "Discrimination", "GDPR", "Wage Dispute", "Contract Breach"]
        for rk in risk_keys:
            rd = risks[rk]
            cfg = _TRIBUNAL_LEVEL_CFG[rd["level"]]
            h(f"""<div style="background:{cfg['bg']};border:1px solid {cfg['border']};
            border-radius:8px;padding:0.6rem 0.9rem;margin-bottom:0.5rem;">
                <div style="display:flex;justify-content:space-between;align-items:center;
                margin-bottom:0.35rem;">
                    <span style="font-size:0.82rem;font-weight:600;color:#E8E4DA;">{rk} Risk</span>
                    <div style="display:flex;align-items:center;gap:0.5rem;">
                        <span style="background:rgba({_hex_to_rgb(cfg['color'])},0.15);
                        border:1px solid rgba({_hex_to_rgb(cfg['color'])},0.4);
                        border-radius:12px;padding:0.1rem 0.6rem;font-size:0.65rem;
                        font-weight:700;color:{cfg['color']};">{cfg['icon']} {rd['level']}</span>
                        <span style="font-family:'Cormorant Garamond',serif;font-size:1.4rem;
                        font-weight:700;color:{cfg['color']};min-width:30px;text-align:right;">
                        {rd['score']}</span>
                    </div>
                </div>
                <div style="background:#1E3448;border-radius:4px;height:7px;overflow:hidden;">
                    <div style="width:{rd['score']}%;background:linear-gradient(90deg,
                    {cfg['bar']},{cfg['bar']}99);height:100%;border-radius:4px;
                    transition:width 0.5s ease;"></div>
                </div>
            </div>""")

    # ── Risk category radar chart ──────────────────────────────────────────────
    section_title("Risk Profile Radar", "📡")
    categories = ["Unfair Dismissal", "Discrimination", "GDPR", "Wage Dispute", "Contract Breach"]
    values     = [risks[c]["score"] for c in categories]
    values_closed = values + [values[0]]
    cats_closed   = categories + [categories[0]]

    fig_radar = go.Figure(go.Scatterpolar(
        r=values_closed, theta=cats_closed,
        fill="toself",
        fillcolor="rgba(231,76,60,0.12)",
        line=dict(color="#E74C3C", width=2),
        marker=dict(color="#C9A84C", size=7),
        name="Tribunal Risk",
    ))
    fig_radar.update_layout(
        polar=dict(
            bgcolor="#071219",
            radialaxis=dict(visible=True, range=[0, 100], tickfont=dict(color="#607D99", size=8),
                            gridcolor="#1E3448", linecolor="#1E3448"),
            angularaxis=dict(tickfont=dict(color="#E8E4DA", size=10), gridcolor="#1E3448",
                             linecolor="#1E3448"),
        ),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        height=320,
        margin=dict(l=50, r=50, t=30, b=30),
        showlegend=False,
    )
    st.plotly_chart(fig_radar, use_container_width=True, config={"displayModeBar": False})


# ═══════════════════════════════════════════════════════════════════════════════
# FEATURE 18 — AUTOMATED LEGAL MEMO GENERATOR
# Generates a structured legal memo with PDF + DOCX export
# ═══════════════════════════════════════════════════════════════════════════════

def generate_legal_memo_content(text: str, result: dict) -> dict:
    """
    Build structured legal memo content from contract text and API result.
    Returns a dict with all memo sections.
    """
    clauses   = result.get("clauses", {}) if result else {}
    insights  = result.get("insights", []) if result else []
    recs      = result.get("recommendations", []) if result else []
    risk_score= result.get("risk_score", 0) if result else 0
    risk_level= result.get("risk_level", "Unknown") if result else "Unknown"
    label     = result.get("label", "Employment Contract") if result else "Employment Contract"
    scottish  = check_scottish_law(text) if text else {}
    comp      = extract_compensation(text) if text else {}
    trib      = compute_tribunal_risks(text, result or {})

    # ── Executive Summary ─────────────────────────────────────────────────────
    exec_summary = (
        f"This legal memorandum presents a comprehensive review of the submitted {label} "
        f"conducted by Caledonian HR Group on {datetime.now().strftime('%d %B %Y')}. "
        f"The contract has been assessed against UK and Scottish employment law requirements. "
        f"The overall risk score is {risk_score}/100, classified as {risk_level} Risk. "
        f"Tribunal exposure is assessed at {trib['exposure_score']}/100 ({trib['exposure_level']} level). "
        f"Immediate attention is required on {sum(1 for v in clauses.values() if not v)} missing clause(s) "
        f"and {sum(1 for s in scottish.values() if s=='missing')} Scottish law compliance item(s)."
    )

    # ── Key Findings ──────────────────────────────────────────────────────────
    key_findings = []
    for ins in insights[:6]:
        key_findings.append(ins)
    if risk_score >= 70:
        key_findings.append(f"CRITICAL: Overall risk score of {risk_score}/100 requires urgent legal review before contract execution.")
    if trib["exposure_level"] in ("High", "Critical"):
        key_findings.append(f"TRIBUNAL EXPOSURE: {trib['exposure_level']} tribunal risk ({trib['exposure_score']}/100) — consider pre-emptive legal counsel.")
    if not key_findings:
        key_findings = ["No critical findings detected. Standard review recommended.", "Contract appears broadly compliant — minor enhancements advised."]

    # ── Missing Clauses ───────────────────────────────────────────────────────
    missing_clauses = [f"{clause} — Not detected in contract" for clause, present in clauses.items() if not present]
    if not missing_clauses:
        missing_clauses = ["No critical clause omissions detected."]

    # ── Compliance Issues ─────────────────────────────────────────────────────
    compliance_issues = []
    for law, status in scottish.items():
        if status == "missing":
            compliance_issues.append(f"MISSING — {law}: Required under Scottish/UK employment law. Immediate inclusion recommended.")
        elif status == "review":
            compliance_issues.append(f"REVIEW REQUIRED — {law}: Clause present but requires legal scrutiny for compliance.")
    if trib["GDPR"]["level"] in ("High", "Critical"):
        compliance_issues.append(f"GDPR NON-COMPLIANCE RISK: Score {trib['GDPR']['score']}/100 — Data protection provisions are inadequate.")
    if not compliance_issues:
        compliance_issues = ["No significant compliance issues identified. Periodic review recommended."]

    # ── Recommendations ───────────────────────────────────────────────────────
    recommendations = list(recs[:8])
    if trib["Unfair Dismissal"]["score"] >= 55:
        recommendations.append("Strengthen dismissal procedures — add explicit fair reason and process requirements to reduce tribunal exposure.")
    if trib["Discrimination"]["score"] >= 55:
        recommendations.append("Add Equality Act 2010 compliance clause and explicit anti-harassment policy reference.")
    if not recommendations:
        recommendations = ["Conduct full legal review with qualified employment solicitor.", "Ensure annual contract review aligned with legislative changes."]

    return {
        "executive_summary":  exec_summary,
        "key_findings":       key_findings,
        "missing_clauses":    missing_clauses,
        "compliance_issues":  compliance_issues,
        "recommendations":    recommendations,
        "contract_type":      label,
        "risk_score":         risk_score,
        "risk_level":         risk_level,
        "exposure_score":     trib["exposure_score"],
        "exposure_level":     trib["exposure_level"],
        "date":               datetime.now().strftime('%d %B %Y'),
        "ref":                f"CHR-{datetime.now().strftime('%Y%m%d')}-{abs(hash(text[:50]))%9000+1000}",
    }


def _pdf_safe(text):
    if text is None:
        return ""

    text = str(text)

    text = text.replace("\r", "")
    text = text.replace("\t", " ")
    text = text.replace("\u2022", "-")
    text = text.replace("\u2013", "-")
    text = text.replace("\u2014", "-")
    text = text.replace("\u2018", "'")
    text = text.replace("\u2019", "'")
    text = text.replace("\u201c", '"')
    text = text.replace("\u201d", '"')

    return text.encode(
        "latin-1",
        errors="replace"
    ).decode("latin-1")


def generate_legal_memo_pdf(memo: dict) -> bytes:
    """Generate the Legal Memo as a dark-styled PDF using FPDF."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Cover header
    pdf.set_fill_color(7, 18, 25)
    pdf.rect(0, 0, 210, 40, "F")
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(201, 168, 76)
    pdf.set_y(8)
    pdf.cell(0, 8, "CALEDONIAN HR GROUP", ln=True, align="C")
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(150, 170, 190)
    pdf.cell(0, 5, "Scottish Employment Intelligence Platform", ln=True, align="C")
    pdf.set_draw_color(201, 168, 76)
    pdf.set_line_width(0.6)
    pdf.line(15, 26, 195, 26)
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(230, 228, 218)
    pdf.set_y(29)
    pdf.cell(0, 7, _pdf_safe("LEGAL MEMORANDUM \u2014 EMPLOYMENT CONTRACT REVIEW"), ln=True, align="C")
    pdf.ln(4)

    # Memo header block
    pdf.set_fill_color(13, 27, 42)
    pdf.set_draw_color(30, 52, 72)
    pdf.set_line_width(0.3)

    def memo_kv(k, v, lcolor=(201,168,76), vcolor=(200,200,190)):
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_text_color(*lcolor)
        pdf.set_x(15)
        pdf.cell(42, 5.5, _pdf_safe(f"  {k}:"), border=0, ln=0)
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(*vcolor)
        remaining_width = pdf.w - pdf.get_x() - 15
        pdf.multi_cell(remaining_width, 5.5, _pdf_safe(str(v)))

    memo_kv("Ref",            memo["ref"])
    memo_kv("Date",           memo["date"])
    memo_kv("Subject",        f"{memo['contract_type']} -- Risk Review")
    memo_kv("Risk Level",     f"{memo['risk_level']} ({memo['risk_score']}/100)")
    memo_kv("Tribunal Exp.",  f"{memo['exposure_level']} ({memo['exposure_score']}/100)")
    pdf.ln(3)
    pdf.set_draw_color(201, 168, 76)
    pdf.set_line_width(0.5)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(4)

    def sec(title):
        pdf.set_fill_color(13, 27, 42)
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(201, 168, 76)
        pdf.set_fill_color(20, 40, 58)
        pdf.cell(0, 7, _pdf_safe(f"  {title.upper()}"), ln=True, fill=True)
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(200, 200, 190)
        pdf.ln(1)

    def body(text_content):
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(200, 200, 190)
        pdf.multi_cell(0, 5.5, _pdf_safe(str(text_content)))
        pdf.ln(1)

    def bullets(items, bullet_color=(201, 168, 76)):
        if not items:
            return
        for item in items:
            text = _pdf_safe(str(item)).strip()

            if not text:
                continue

            # Always start from left margin
            pdf.set_x(pdf.l_margin)

            pdf.set_font("Helvetica", "B", 8.5)
            pdf.set_text_color(*bullet_color)

            pdf.cell(6, 5.5, "-", border=0)

            pdf.set_font("Helvetica", "", 8.5)
            pdf.set_text_color(200, 200, 190)

            available_width = (
                pdf.w
                - pdf.l_margin
                - pdf.r_margin
                - 6
            )

            pdf.multi_cell(
                available_width,
                5.5,
                text
            )

            pdf.ln(0.5)

    pdf.ln(1)

    # Set dark background for content pages
    pdf.set_fill_color(11, 24, 37)
    pdf.rect(0, pdf.get_y()-2, 210, 297, "F")

    sec("1. Executive Summary")
    body(memo["executive_summary"])

    sec("2. Key Findings")
    bullets(memo["key_findings"])

    sec("3. Missing Clauses")
    bullets(memo["missing_clauses"], bullet_color=(231, 76, 60))

    sec("4. Compliance Issues")
    bullets(memo["compliance_issues"], bullet_color=(243, 156, 18))

    sec("5. Recommendations")
    bullets(memo["recommendations"], bullet_color=(39, 174, 96))

    # Footer
    pdf.set_y(-20)
    pdf.set_font("Helvetica", "I", 7)
    pdf.set_text_color(60, 80, 100)
    pdf.set_draw_color(30, 52, 72)
    pdf.set_line_width(0.3)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(2)
    pdf.cell(0, 4, _pdf_safe(f"Caledonian HR Group · 16 Charlotte Square, Edinburgh EH2 4DR · Ref: {memo['ref']} · CONFIDENTIAL"), ln=True)

    return bytes(pdf.output())


def generate_legal_memo_docx(memo: dict) -> bytes:
    """Generate the Legal Memo as a DOCX file using python-docx."""
    if not DOCX_AVAILABLE:
        return b""

    doc = DocxDocument()

    # ── Title ──────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("CALEDONIAN HR GROUP")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(201, 168, 76)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("LEGAL MEMORANDUM — EMPLOYMENT CONTRACT REVIEW")
    sub_run.bold = True
    sub_run.font.size = Pt(12)

    doc.add_paragraph()

    # ── Memo header ────────────────────────────────────────────────────────────
    def add_kv(label, value):
        p = doc.add_paragraph()
        r_label = p.add_run(f"{label}: ")
        r_label.bold = True
        r_label.font.color.rgb = RGBColor(140, 100, 30)
        r_value = p.add_run(str(value))
        r_value.font.size = Pt(10)

    add_kv("Reference",           memo["ref"])
    add_kv("Date",                memo["date"])
    add_kv("Subject",             f"{memo['contract_type']} — Risk Review")
    add_kv("Overall Risk",        f"{memo['risk_level']} ({memo['risk_score']}/100)")
    add_kv("Tribunal Exposure",   f"{memo['exposure_level']} ({memo['exposure_score']}/100)")
    doc.add_paragraph()

    # ── Sections ───────────────────────────────────────────────────────────────
    def add_section(title, content_items, is_paragraph=False):
        h_para = doc.add_paragraph()
        h_run = h_para.add_run(title.upper())
        h_run.bold = True
        h_run.font.size = Pt(11)
        h_run.font.color.rgb = RGBColor(201, 168, 76)

        if is_paragraph:
            p = doc.add_paragraph(str(content_items))
            p.style.font.size = Pt(10)
        else:
            for item in content_items:
                bullet_para = doc.add_paragraph(style="List Bullet")
                bullet_run = bullet_para.add_run(str(item))
                bullet_run.font.size = Pt(10)

        doc.add_paragraph()

    add_section("1. Executive Summary",  memo["executive_summary"], is_paragraph=True)
    add_section("2. Key Findings",       memo["key_findings"])
    add_section("3. Missing Clauses",    memo["missing_clauses"])
    add_section("4. Compliance Issues",  memo["compliance_issues"])
    add_section("5. Recommendations",    memo["recommendations"])

    # ── Footer ─────────────────────────────────────────────────────────────────
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"Caledonian HR Group · 16 Charlotte Square, Edinburgh EH2 4DR · Ref: {memo['ref']} · CONFIDENTIAL"
    )
    footer_run.font.size = Pt(7)
    footer_run.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_legal_memo_tab(text: str, result: dict):
    """Feature 18 — Legal Memo Generator tab."""
    h("""<div style="font-size:0.72rem;color:#607D99;margin-bottom:1rem;letter-spacing:0.04em;">
    Generate a structured legal memorandum covering executive summary, key findings,
    missing clauses, compliance issues, and recommendations. Export to PDF or DOCX.</div>""")

    if not text.strip():
        st.info("Upload or paste a contract to generate a Legal Memo.")
        return

    # Auto-generate or on demand
    if "legal_memo_content" not in st.session_state or st.button("🔄 Regenerate Memo", key="regen_memo"):
        with st.spinner("Generating legal memorandum…"):
            st.session_state["legal_memo_content"] = generate_legal_memo_content(text, result or {})

    memo = st.session_state["legal_memo_content"]

    # ── Memo header display ────────────────────────────────────────────────────
    rl = memo["risk_level"]
    rc_cfg = _TRIBUNAL_LEVEL_CFG.get(
        "Critical" if memo["risk_score"] >= 75 else
        "High"     if memo["risk_score"] >= 55 else
        "Medium"   if memo["risk_score"] >= 35 else "Low"
    )
    h(f"""
    <div style="background:linear-gradient(135deg,#071219 0%,#0D1F30 100%);
    border:1px solid #1E3448;border-top:3px solid #C9A84C;border-radius:12px;
    padding:1.5rem;margin-bottom:1.2rem;">
        <div style="display:flex;justify-content:space-between;align-items:flex-start;flex-wrap:wrap;gap:1rem;">
            <div>
                <div style="font-size:0.58rem;letter-spacing:0.2em;text-transform:uppercase;
                color:#3A5570;margin-bottom:0.3rem;">LEGAL MEMORANDUM · CONFIDENTIAL</div>
                <div style="font-family:'Cormorant Garamond',serif;font-size:1.5rem;font-weight:700;
                color:#E8E4DA;">Employment Contract Review</div>
                <div style="font-size:0.78rem;color:#607D99;margin-top:0.3rem;">
                Caledonian HR Group · Scottish Employment Intelligence</div>
            </div>
            <div style="text-align:right;">
                <div style="font-size:0.65rem;color:#607D99;">Ref: {memo['ref']}</div>
                <div style="font-size:0.8rem;color:#C9A84C;font-weight:600;">{memo['date']}</div>
                <div style="margin-top:0.5rem;background:rgba({_hex_to_rgb(rc_cfg['color'])},0.12);
                border:1px solid rgba({_hex_to_rgb(rc_cfg['color'])},0.35);border-radius:6px;
                padding:0.25rem 0.7rem;font-size:0.75rem;font-weight:700;color:{rc_cfg['color']};">
                {rc_cfg['icon']} {rl} Risk — {memo['risk_score']}/100</div>
            </div>
        </div>
        <div style="display:flex;gap:0.5rem;margin-top:1rem;flex-wrap:wrap;">
            <span style="background:rgba(201,168,76,0.08);border:1px solid rgba(201,168,76,0.25);
            border-radius:4px;padding:0.2rem 0.6rem;font-size:0.68rem;color:#C9A84C;">
            {memo['contract_type']}</span>
            <span style="background:rgba(74,144,217,0.08);border:1px solid rgba(74,144,217,0.25);
            border-radius:4px;padding:0.2rem 0.6rem;font-size:0.68rem;color:#4A90D9;">
            Tribunal Exposure: {memo['exposure_level']} ({memo['exposure_score']})</span>
        </div>
    </div>""")

    # ── Memo sections ──────────────────────────────────────────────────────────
    def memo_section(number, title, icon, content_items, is_paragraph=False, item_color="#E8E4DA", border_color="#C9A84C"):
        h(f"""<div style="background:#0D1F30;border:1px solid #1E3448;
        border-left:4px solid {border_color};border-radius:0 10px 10px 0;
        padding:1rem 1.2rem;margin-bottom:0.8rem;">
            <div style="font-size:0.65rem;font-weight:700;letter-spacing:0.14em;
            text-transform:uppercase;color:#607D99;margin-bottom:0.5rem;">
            {icon} &nbsp;{number}. {title}</div>""")

        if is_paragraph:
            h(f'<div style="font-size:0.82rem;color:{item_color};line-height:1.7;">{content_items}</div>')
        else:
            items_html = ""
            for item in content_items:
                items_html += f"""<div style="display:flex;gap:0.5rem;margin-bottom:0.3rem;
                font-size:0.8rem;color:{item_color};line-height:1.6;">
                <span style="color:{border_color};margin-top:0.05rem;flex-shrink:0;">•</span>
                <span>{item}</span></div>"""
            h(items_html)
        h("</div>")

    memo_section("1", "Executive Summary",  "📑", memo["executive_summary"], is_paragraph=True)
    memo_section("2", "Key Findings",       "🔍", memo["key_findings"],       border_color="#C9A84C")
    memo_section("3", "Missing Clauses",    "✗",  memo["missing_clauses"],    item_color="#F0A8A8", border_color="#E74C3C")
    memo_section("4", "Compliance Issues",  "⚠️", memo["compliance_issues"],  item_color="#F5C888", border_color="#F39C12")
    memo_section("5", "Recommendations",    "✅", memo["recommendations"],    item_color="#D4EDE0", border_color="#27AE60")

    # ── Export buttons ─────────────────────────────────────────────────────────
    st.markdown("")
    section_title("Export Legal Memo", "📤")
    exp_col1, exp_col2, exp_col3 = st.columns([1, 1, 2])

    with exp_col1:
        pdf_bytes = generate_legal_memo_pdf(memo)
        st.download_button(
            "📄 Export as PDF",
            data=pdf_bytes,
            file_name=f"legal_memo_{memo['ref']}.pdf",
            mime="application/pdf",
            use_container_width=True,
            key="memo_pdf_dl",
        )

    with exp_col2:
        if DOCX_AVAILABLE:
            docx_bytes = generate_legal_memo_docx(memo)
            st.download_button(
                "📝 Export as DOCX",
                data=docx_bytes,
                file_name=f"legal_memo_{memo['ref']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="memo_docx_dl",
            )
        else:
            if st.button("📝 Export as DOCX", use_container_width=True, key="memo_docx_na"):
                st.warning("Install `python-docx` to enable DOCX export: `pip install python-docx`")

    with exp_col3:
        h("""<div style="background:#071219;border:1px solid #1E3448;border-radius:7px;
        padding:0.45rem 0.8rem;font-size:0.73rem;color:#607D99;line-height:1.5;">
        📌 Memo is auto-generated from contract analysis. Run contract analysis first
        for the most complete memo. Click Regenerate to refresh after new analysis.
        </div>""")


# ═══════════════════════════════════════════════════════════════════════════════
# NEW FEATURE A — HR DIRECTOR EXECUTIVE DASHBOARD
# KPI Cards · Risk Trends · Compliance Trends · Animated Gauges
# ═══════════════════════════════════════════════════════════════════════════════

def render_hr_director_dashboard():
    """HR Director Executive Dashboard — animated KPIs, trend charts, risk & compliance."""

    # ── Header ─────────────────────────────────────────────────────────────────
    h("""
    <div class="exec-dashboard-header">
        <div style="display:flex;align-items:flex-start;justify-content:space-between;flex-wrap:wrap;gap:1rem;">
            <div>
                <div style="font-size:0.58rem;letter-spacing:0.22em;text-transform:uppercase;
                color:#3A5570;margin-bottom:0.35rem;">CALEDONIAN HR GROUP · RESTRICTED</div>
                <div style="font-family:'Cormorant Garamond',serif;font-size:1.85rem;font-weight:700;
                color:#E8E4DA;line-height:1.2;letter-spacing:0.01em;">
                    HR Director Executive Dashboard
                </div>
                <div style="font-size:0.78rem;color:#607D99;margin-top:0.4rem;">
                    Portfolio-wide contract intelligence &nbsp;·&nbsp; Real-time risk &amp; compliance overview
                </div>
            </div>
            <div style="display:flex;flex-direction:column;align-items:flex-end;gap:0.5rem;">
                <div style="font-size:0.62rem;color:#607D99;letter-spacing:0.06em;">
                    Report generated: """ + datetime.now().strftime('%d %b %Y · %H:%M') + """
                </div>
                <div style="display:flex;gap:0.4rem;flex-wrap:wrap;justify-content:flex-end;">
                    <span style="background:rgba(201,168,76,0.1);border:1px solid rgba(201,168,76,0.3);
                    border-radius:20px;padding:0.22rem 0.8rem;font-size:0.65rem;color:#C9A84C;">
                    🏴󠁧󠁢󠁳󠁣󠁴󠁿 Scotland</span>
                    <span style="background:rgba(39,174,96,0.1);border:1px solid rgba(39,174,96,0.3);
                    border-radius:20px;padding:0.22rem 0.8rem;font-size:0.65rem;color:#27AE60;">
                    ⚡ Live Intelligence</span>
                    <span style="background:rgba(74,144,217,0.1);border:1px solid rgba(74,144,217,0.3);
                    border-radius:20px;padding:0.22rem 0.8rem;font-size:0.65rem;color:#4A90D9;">
                    🔒 GDPR</span>
                </div>
            </div>
        </div>
    </div>""")

    # ── Simulated portfolio KPI data ───────────────────────────────────────────
    # In production these would come from the database / API
    total_contracts   = 247
    high_risk_count   = 38
    compliance_rate   = 82
    avg_salary        = 54200
    missing_clauses   = 91
    risk_trend        = [55, 58, 62, 59, 53, 48, 44, 51, 47, 38, 42, 38]
    compliance_trend  = [71, 73, 74, 76, 77, 79, 80, 79, 81, 82, 83, 82]
    months            = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]

    # ── Row 1: Animated KPI Cards ──────────────────────────────────────────────
    section_title("Key Performance Indicators", "📊")
    k1, k2, k3, k4, k5, k6 = st.columns(6)

    kpi_data = [
        (k1, "Total Contracts",    str(total_contracts),  "#C9A84C",   "📁", "+12 this month",  True),
        (k2, "High Risk",          str(high_risk_count),  "#E74C3C",   "⚠️", f"{int(high_risk_count/total_contracts*100)}% of portfolio", True),
        (k3, "Compliance Rate",    f"{compliance_rate}%", "#27AE60",   "✅", "+3% vs last month", False),
        (k4, "Avg Salary",         f"£{avg_salary:,}",    "#4A90D9",   "💷", "UK market: £52,400", False),
        (k5, "Missing Clauses",    str(missing_clauses),  "#F39C12",   "⚡", "Across all contracts", True),
        (k6, "Low Risk",           str(total_contracts - high_risk_count - 42), "#607D99", "🟢", "Standard / compliant", False),
    ]

    for col, label, value, color, icon, sub, is_risk in kpi_data:
        with col:
            h(f"""
            <div class="exec-kpi-card" style="border-top-color:{color};">
                <div class="exec-kpi-label">{icon} {label}</div>
                <div class="exec-kpi-value" style="color:{color};">{value}</div>
                <div class="exec-kpi-trend" style="color:{'#E74C3C' if is_risk else '#607D99'};">
                    <span style="font-size:0.68rem;">{sub}</span>
                </div>
                <div style="margin-top:0.6rem;background:#1E3448;border-radius:3px;height:2px;">
                    <div style="width:{'100%' if label=='Total Contracts' else str(min(99,int(value.replace('%','').replace('£','').replace(',','').split('.')[0]) % 100 if value.replace('%','').replace('£','').replace(',','') else 50)) + '%' if not 'Avg' in label else '72%'};
                    background:linear-gradient(90deg,{color},{color}66);height:100%;border-radius:3px;
                    transition:width 1s ease;"></div>
                </div>
            </div>""")

    st.markdown("")

    # ── Row 2: Risk Trend + Compliance Trend Charts ────────────────────────────
    section_title("Trend Analytics", "📈")
    chart_col1, chart_col2 = st.columns(2)

    with chart_col1:
        h("""<div class="chart-container-glass">
        <div style="font-size:0.68rem;font-weight:700;letter-spacing:0.12em;text-transform:uppercase;
        color:#E74C3C;margin-bottom:0.6rem;">🔴 Portfolio Risk Score Trend</div>""")

        fig_risk = go.Figure()
        fig_risk.add_trace(go.Scatter(
            x=months, y=risk_trend,
            mode="lines+markers",
            name="Risk Score",
            line=dict(color="#E74C3C", width=2.5, shape="spline"),
            marker=dict(size=7, color="#E74C3C",
                        line=dict(color="#071219", width=2)),
            fill="tozeroy",
            fillcolor="rgba(231,76,60,0.06)",
            hovertemplate="<b>%{x}</b><br>Risk: %{y}<extra></extra>",
        ))
        fig_risk.add_hline(y=60, line_dash="dot", line_color="rgba(231,76,60,0.4)",
                           annotation_text="High Risk Threshold",
                           annotation_font_color="#E74C3C",
                           annotation_font_size=9)
        fig_risk.update_layout(
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            height=220, margin=dict(l=40, r=20, t=20, b=30),
            font=dict(color="#E8E4DA", size=10),
            xaxis=dict(gridcolor="rgba(30,52,72,0.5)", tickfont=dict(color="#607D99")),
            yaxis=dict(gridcolor="rgba(30,52,72,0.5)", range=[0, 100],
                       tickfont=dict(color="#607D99"), title="Score"),
            showlegend=False,
            hovermode="x unified",
            hoverlabel=dict(bgcolor="#0D1F30", bordercolor="#E74C3C",
                            font=dict(color="#E8E4DA", size=11)),
        )
        st.plotly_chart(fig_risk, use_container_width=True, config={"displayModeBar": False})
        h("</div>")

    with chart_col2:
        h("""<div class="chart-container-glass">
        <div style="font-size:0.68rem;font-weight:700;letter-spacing:0.12em;text-transform:uppercase;
        color:#27AE60;margin-bottom:0.6rem;">✅ Compliance Rate Trend</div>""")

        fig_comp = go.Figure()
        fig_comp.add_trace(go.Scatter(
            x=months, y=compliance_trend,
            mode="lines+markers",
            name="Compliance %",
            line=dict(color="#27AE60", width=2.5, shape="spline"),
            marker=dict(size=7, color="#27AE60",
                        line=dict(color="#071219", width=2)),
            fill="tozeroy",
            fillcolor="rgba(39,174,96,0.06)",
            hovertemplate="<b>%{x}</b><br>Compliance: %{y}%<extra></extra>",
        ))
        fig_comp.add_hline(y=85, line_dash="dot", line_color="rgba(201,168,76,0.4)",
                           annotation_text="Target 85%",
                           annotation_font_color="#C9A84C",
                           annotation_font_size=9)
        fig_comp.update_layout(
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            height=220, margin=dict(l=40, r=20, t=20, b=30),
            font=dict(color="#E8E4DA", size=10),
            xaxis=dict(gridcolor="rgba(30,52,72,0.5)", tickfont=dict(color="#607D99")),
            yaxis=dict(gridcolor="rgba(30,52,72,0.5)", range=[60, 100],
                       tickfont=dict(color="#607D99"), title="Compliance %"),
            showlegend=False,
            hovermode="x unified",
            hoverlabel=dict(bgcolor="#0D1F30", bordercolor="#27AE60",
                            font=dict(color="#E8E4DA", size=11)),
        )
        st.plotly_chart(fig_comp, use_container_width=True, config={"displayModeBar": False})
        h("</div>")

    # ── Row 3: Risk Distribution Donut + Missing Clauses Bar ──────────────────
    section_title("Risk Distribution & Clause Gaps", "🎯")
    d1, d2 = st.columns([1, 1.4])

    with d1:
        fig_donut = go.Figure(go.Pie(
            labels=["High Risk", "Medium Risk", "Low Risk"],
            values=[38, 42, 167],
            hole=0.62,
            marker=dict(
                colors=["#E74C3C", "#F39C12", "#27AE60"],
                line=dict(color="#071219", width=3),
            ),
            textinfo="label+percent",
            textfont=dict(color="#E8E4DA", size=10),
            hovertemplate="<b>%{label}</b><br>%{value} contracts<br>%{percent}<extra></extra>",
        ))
        fig_donut.update_layout(
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            height=240, margin=dict(l=20, r=20, t=20, b=20),
            showlegend=True,
            legend=dict(orientation="v", font=dict(color="#607D99", size=10),
                        bgcolor="rgba(0,0,0,0)"),
            annotations=[dict(
                text=f"<b style='color:#C9A84C'>{total_contracts}</b><br><span style='color:#607D99;font-size:9px'>TOTAL</span>",
                x=0.5, y=0.5, font_size=16, showarrow=False,
                font=dict(color="#E8E4DA"),
            )],
        )
        st.plotly_chart(fig_donut, use_container_width=True, config={"displayModeBar": False})

    with d2:
        missing_clause_data = {
            "GDPR / Data Protection": 42,
            "Grievance Procedure":    38,
            "Equality Clause":        31,
            "Disciplinary Procedure": 28,
            "Pension Reference":      22,
            "Holiday Entitlement":    14,
            "Notice Period":          9,
        }
        labels  = list(missing_clause_data.keys())
        vals    = list(missing_clause_data.values())
        pcts    = [int(v / total_contracts * 100) for v in vals]
        colors  = ["#E74C3C" if p >= 15 else "#F39C12" if p >= 10 else "#4A90D9" for p in pcts]

        fig_miss = go.Figure(go.Bar(
            x=vals, y=labels, orientation="h",
            marker=dict(color=colors, line=dict(color="#071219", width=1)),
            text=[f"{v} ({p}%)" for v, p in zip(vals, pcts)],
            textposition="outside",
            textfont=dict(color="#E8E4DA", size=9),
            hovertemplate="%{y}<br>Missing: %{x} contracts<extra></extra>",
        ))
        fig_miss.update_layout(
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            height=240, margin=dict(l=190, r=60, t=10, b=20),
            font=dict(color="#E8E4DA", size=10),
            xaxis=dict(gridcolor="rgba(30,52,72,0.5)", title="# Contracts"),
            yaxis=dict(gridcolor="rgba(0,0,0,0)", tickfont=dict(size=9, color="#B0C4D8")),
            bargap=0.25,
        )
        st.plotly_chart(fig_miss, use_container_width=True, config={"displayModeBar": False})

    # ── Row 4: Summary Intelligence Cards ─────────────────────────────────────
    h('<div class="exec-divider"></div>')
    section_title("Executive Intelligence Summary", "🧠")
    s1, s2, s3 = st.columns(3)

    with s1:
        h("""
        <div style="background:linear-gradient(135deg,rgba(231,76,60,0.08),rgba(13,31,48,0.95));
        border:1px solid rgba(231,76,60,0.25);border-top:2px solid #E74C3C;
        border-radius:12px;padding:1.1rem;animation:fadeInUp 0.5s ease forwards;">
            <div style="font-size:0.65rem;font-weight:700;letter-spacing:0.12em;
            text-transform:uppercase;color:#E74C3C;margin-bottom:0.7rem;">⚠️ Risk Alerts</div>
            <div style="font-size:0.8rem;color:#E8E4DA;line-height:1.7;">
                • 38 contracts exceed risk threshold<br>
                • Sales dept. highest exposure (72/100)<br>
                • Non-compete clauses in 31% of portfolio<br>
                • 3 contracts require immediate review
            </div>
        </div>""")

    with s2:
        h("""
        <div style="background:linear-gradient(135deg,rgba(39,174,96,0.08),rgba(13,31,48,0.95));
        border:1px solid rgba(39,174,96,0.25);border-top:2px solid #27AE60;
        border-radius:12px;padding:1.1rem;animation:fadeInUp 0.6s ease forwards;">
            <div style="font-size:0.65rem;font-weight:700;letter-spacing:0.12em;
            text-transform:uppercase;color:#27AE60;margin-bottom:0.7rem;">✅ Compliance Wins</div>
            <div style="font-size:0.8rem;color:#E8E4DA;line-height:1.7;">
                • Compliance rate +11% YTD improvement<br>
                • HR department: 94% compliant<br>
                • GDPR coverage up from 61% → 83%<br>
                • Notice period coverage: 96% complete
            </div>
        </div>""")

    with s3:
        h("""
        <div style="background:linear-gradient(135deg,rgba(201,168,76,0.08),rgba(13,31,48,0.95));
        border:1px solid rgba(201,168,76,0.25);border-top:2px solid #C9A84C;
        border-radius:12px;padding:1.1rem;animation:fadeInUp 0.7s ease forwards;">
            <div style="font-size:0.65rem;font-weight:700;letter-spacing:0.12em;
            text-transform:uppercase;color:#C9A84C;margin-bottom:0.7rem;">📋 Action Required</div>
            <div style="font-size:0.8rem;color:#E8E4DA;line-height:1.7;">
                • Review 42 contracts missing GDPR clause<br>
                • Update 38 grievance procedures<br>
                • Salary benchmarking: 67 below median<br>
                • Q1 board report due in 14 days
            </div>
        </div>""")


# ═══════════════════════════════════════════════════════════════════════════════
# NEW FEATURE B — CONTRACT VERSION HISTORY
# 4-Version Timeline · Added / Removed / Modified Clauses · Visual diff
# ═══════════════════════════════════════════════════════════════════════════════

# Default version history data (populated from real diffs if contracts uploaded)
_DEFAULT_VERSION_HISTORY = [
    {
        "version": "v1.0",
        "label":   "Original Draft",
        "date":    "15 Jan 2024",
        "author":  "Legal Dept.",
        "status":  "Archived",
        "added":   ["Commencement of employment", "Job title and duties", "Place of work", "Hours of work"],
        "removed": [],
        "modified":[],
        "risk_delta": 0,
        "compliance": 45,
        "notes":   "Initial contract template drafted by legal team.",
    },
    {
        "version": "v2.0",
        "label":   "HR Review",
        "date":    "3 Feb 2024",
        "author":  "HR Director",
        "status":  "Archived",
        "added":   ["Pension auto-enrolment clause", "GDPR data processing statement", "Remote working policy reference"],
        "removed": ["Generic confidentiality paragraph"],
        "modified":["Notice period (4 weeks → 8 weeks)", "Probation period (3 months → 6 months)"],
        "risk_delta": -12,
        "compliance": 62,
        "notes":   "HR review added mandatory UK employment clauses.",
    },
    {
        "version": "v3.0",
        "label":   "Legal Counsel Review",
        "date":    "19 Feb 2024",
        "author":  "Solicitor — Caledonian Law LLP",
        "status":  "Superseded",
        "added":   ["Restrictive covenant (6 months)", "Non-solicitation clause", "Garden leave provision", "IP assignment clause"],
        "removed": ["Ambiguous termination language"],
        "modified":["Confidentiality scope (narrow → broad)", "Governing law (England → Scotland)", "Dispute resolution (litigation → mediation first)"],
        "risk_delta": +8,
        "compliance": 74,
        "notes":   "Legal counsel strengthened post-employment restrictions.",
    },
    {
        "version": "v4.0",
        "label":   "Final Executed",
        "date":    "1 Mar 2024",
        "author":  "CEO & Employee",
        "status":  "Active",
        "added":   ["Equality and diversity commitment", "Grievance escalation procedure", "Whistleblowing reference"],
        "removed": ["Overly broad non-compete language"],
        "modified":["Restrictive covenant (6 months → 3 months)", "Salary review mechanism clarified"],
        "risk_delta": -15,
        "compliance": 88,
        "notes":   "Final negotiated version. Signed and countersigned 1 March 2024.",
    },
]


def render_version_history(uploaded_texts: list | None = None):
    """
    New Feature B — Contract Version History.
    Shows a visual timeline of v1→v4 with added/removed/modified clause diffs.
    If uploaded_texts (list of contract strings) is provided, diffs are computed live.
    """

    h("""
    <div style="background:linear-gradient(135deg,#071219,#0D1F30);
    border:1px solid #1E3448;border-top:2px solid #C9A84C;
    border-radius:12px;padding:1.2rem 1.4rem;margin-bottom:1.2rem;">
        <div style="font-size:0.62rem;letter-spacing:0.18em;text-transform:uppercase;
        color:#3A5570;margin-bottom:0.25rem;">CONTRACT LIFECYCLE MANAGEMENT</div>
        <div style="font-family:'Cormorant Garamond',serif;font-size:1.4rem;font-weight:700;
        color:#E8E4DA;">Version History &amp; Change Log</div>
        <div style="font-size:0.75rem;color:#607D99;margin-top:0.2rem;">
        Track every revision · + Added &nbsp; − Removed &nbsp; ~ Modified &nbsp;  · Visual timeline
        </div>
    </div>""")

    # Use live diffs if uploaded, else default demo data
    if uploaded_texts and len(uploaded_texts) >= 2:
        versions = list(_DEFAULT_VERSION_HISTORY)
        for i in range(min(len(uploaded_texts) - 1, 3)):
            d = diff_contracts(uploaded_texts[i], uploaded_texts[i + 1])
            versions[i + 1]["added"]    = d["added"][:6]
            versions[i + 1]["removed"]  = d["removed"][:4]
            versions[i + 1]["modified"] = [f"{m['old'][:80]}…" for m in d["modified"][:3]]
    else:
        versions = _DEFAULT_VERSION_HISTORY

    # ── Timeline legend ────────────────────────────────────────────────────────
    h("""<div style="display:flex;gap:1.2rem;margin-bottom:1.2rem;flex-wrap:wrap;">
        <span style="display:flex;align-items:center;gap:0.4rem;font-size:0.72rem;color:#27AE60;">
            <span style="width:10px;height:10px;border-radius:2px;background:rgba(39,174,96,0.25);
            border:1px solid #27AE60;display:inline-block;"></span>+ Added</span>
        <span style="display:flex;align-items:center;gap:0.4rem;font-size:0.72rem;color:#E74C3C;">
            <span style="width:10px;height:10px;border-radius:2px;background:rgba(231,76,60,0.25);
            border:1px solid #E74C3C;display:inline-block;"></span>− Removed</span>
        <span style="display:flex;align-items:center;gap:0.4rem;font-size:0.72rem;color:#F39C12;">
            <span style="width:10px;height:10px;border-radius:2px;background:rgba(243,156,18,0.25);
            border:1px solid #F39C12;display:inline-block;"></span>~ Modified</span>
    </div>""")

    # ── Horizontal version selector ────────────────────────────────────────────
    v_selector = st.radio(
        "Select version to inspect:",
        options=[f"{v['version']} — {v['label']}" for v in versions],
        horizontal=True,
        label_visibility="collapsed",
        key="vh_version_select",
    )
    selected_idx = next(
        (i for i, v in enumerate(versions) if v["version"] in v_selector), 0
    )

    # ── Visual progress bar showing version position ───────────────────────────
    progress_pct = int((selected_idx + 1) / len(versions) * 100)
    h(f"""<div style="margin-bottom:1.2rem;">
        <div style="display:flex;justify-content:space-between;font-size:0.65rem;
        color:#3A5570;margin-bottom:0.3rem;">
            <span>v1.0 · Original</span>
            <span style="color:#C9A84C;">Version {selected_idx+1} of {len(versions)} selected</span>
            <span>v4.0 · Final</span>
        </div>
        <div style="background:#1E3448;border-radius:6px;height:6px;overflow:hidden;">
            <div style="width:{progress_pct}%;
            background:linear-gradient(90deg,#C9A84C,#A8862A);
            height:100%;border-radius:6px;transition:width 0.6s cubic-bezier(0.4,0,0.2,1);">
            </div>
        </div>
        <div style="display:flex;justify-content:space-between;margin-top:0.3rem;">
            {''.join(f'<span style="font-size:0.6rem;color:{"#C9A84C" if i==selected_idx else "#3A5570"};">●</span>' for i in range(len(versions)))}
        </div>
    </div>""")

    # ── Two-column: Timeline + Detail ─────────────────────────────────────────
    tl_col, detail_col = st.columns([1, 1.6])

    with tl_col:
        section_title("Version Timeline", "🕐")
        h('<div class="version-timeline-container">')
        for i, v in enumerate(versions):
            is_selected = (i == selected_idx)
            is_latest   = (v["status"] == "Active")
            status_color = {"Active": "#27AE60", "Archived": "#3A5570",
                            "Superseded": "#607D99"}.get(v["status"], "#607D99")
            rd_color = "#27AE60" if v["risk_delta"] < 0 else "#E74C3C" if v["risk_delta"] > 0 else "#607D99"
            rd_icon  = "↓" if v["risk_delta"] < 0 else "↑" if v["risk_delta"] > 0 else "—"

            h(f"""
            <div class="version-node {'latest' if is_latest else ''}"
            style="background:{'rgba(201,168,76,0.07)' if is_selected else 'transparent'};
            border:1px solid {'rgba(201,168,76,0.3)' if is_selected else 'rgba(30,52,72,0.4)'};
            border-radius:10px;padding:0.75rem 0.9rem;transition:all 0.2s ease;">
                <div style="display:flex;justify-content:space-between;align-items:flex-start;margin-bottom:0.3rem;">
                    <div>
                        <span style="font-size:0.7rem;font-weight:700;color:#C9A84C;
                        letter-spacing:0.06em;">{v['version']}</span>
                        <span style="font-size:0.7rem;color:#E8E4DA;margin-left:0.4rem;">{v['label']}</span>
                    </div>
                    <span style="font-size:0.6rem;color:{status_color};background:rgba({_hex_to_rgb(status_color)},0.12);
                    border:1px solid rgba({_hex_to_rgb(status_color)},0.3);border-radius:10px;
                    padding:0.1rem 0.45rem;white-space:nowrap;">{v['status']}</span>
                </div>
                <div style="font-size:0.65rem;color:#607D99;margin-bottom:0.35rem;">
                    📅 {v['date']} &nbsp;·&nbsp; {v['author']}
                </div>
                <div style="display:flex;gap:0.5rem;font-size:0.65rem;flex-wrap:wrap;">
                    <span style="color:#27AE60;">+{len(v['added'])}</span>
                    <span style="color:#E74C3C;">−{len(v['removed'])}</span>
                    <span style="color:#F39C12;">~{len(v['modified'])}</span>
                    <span style="color:{rd_color};margin-left:auto;">Risk {rd_icon}{abs(v['risk_delta']) if v['risk_delta'] else '—'}</span>
                </div>
            </div>""")
        h("</div>")

    with detail_col:
        v = versions[selected_idx]
        compliance_color = "#27AE60" if v["compliance"] >= 75 else "#F39C12" if v["compliance"] >= 50 else "#E74C3C"

        section_title(f"Version Detail — {v['version']}: {v['label']}", "📄")

        # Meta strip
        h(f"""<div style="background:#071219;border:1px solid #1E3448;border-radius:8px;
        padding:0.7rem 1rem;margin-bottom:1rem;
        display:flex;gap:1.5rem;flex-wrap:wrap;align-items:center;">
            <div><div style="font-size:0.58rem;text-transform:uppercase;letter-spacing:0.1em;
            color:#3A5570;margin-bottom:0.15rem;">Date</div>
            <div style="font-size:0.82rem;color:#C9A84C;font-weight:600;">{v['date']}</div></div>
            <div><div style="font-size:0.58rem;text-transform:uppercase;letter-spacing:0.1em;
            color:#3A5570;margin-bottom:0.15rem;">Author</div>
            <div style="font-size:0.82rem;color:#E8E4DA;">{v['author']}</div></div>
            <div><div style="font-size:0.58rem;text-transform:uppercase;letter-spacing:0.1em;
            color:#3A5570;margin-bottom:0.15rem;">Compliance</div>
            <div style="font-size:0.82rem;color:{compliance_color};font-weight:700;">{v['compliance']}%</div></div>
            <div><div style="font-size:0.58rem;text-transform:uppercase;letter-spacing:0.1em;
            color:#3A5570;margin-bottom:0.15rem;">Risk Δ</div>
            <div style="font-size:0.82rem;color:{'#27AE60' if v['risk_delta'] < 0 else '#E74C3C' if v['risk_delta'] > 0 else '#607D99'};font-weight:700;">
            {'↓' if v['risk_delta'] < 0 else '↑' if v['risk_delta'] > 0 else '—'}{abs(v['risk_delta']) if v['risk_delta'] else '0'}
            </div></div>
        </div>""")

        # Notes
        if v["notes"]:
            h(f"""<div style="background:rgba(201,168,76,0.05);border:1px solid rgba(201,168,76,0.2);
            border-left:3px solid #C9A84C;border-radius:0 8px 8px 0;
            padding:0.5rem 0.9rem;margin-bottom:1rem;font-size:0.8rem;
            color:#E8E4DA;line-height:1.6;">📝 {v['notes']}</div>""")

        # Diff tabs
        da, dr, dm = st.tabs([
            f"  ✚ Added ({len(v['added'])})  ",
            f"  ✖ Removed ({len(v['removed'])})  ",
            f"  ⟳ Modified ({len(v['modified'])})  ",
        ])

        with da:
            if v["added"]:
                for item in v["added"]:
                    h(f"""<div class="diff-added" style="padding:0.45rem 0.8rem;
                    margin-bottom:0.3rem;font-size:0.8rem;color:#D4EDE0;">
                    <span style="color:#27AE60;font-weight:700;margin-right:0.5rem;">+</span>{item}</div>""")
            else:
                h('<div style="font-size:0.8rem;color:#3A5570;padding:0.5rem 0;">No clauses added in this version.</div>')

        with dr:
            if v["removed"]:
                for item in v["removed"]:
                    h(f"""<div class="diff-removed" style="padding:0.45rem 0.8rem;
                    margin-bottom:0.3rem;font-size:0.8rem;color:#F0A8A8;">
                    <span style="color:#E74C3C;font-weight:700;margin-right:0.5rem;">−</span>{item}</div>""")
            else:
                h('<div style="font-size:0.8rem;color:#3A5570;padding:0.5rem 0;">No clauses removed in this version.</div>')

        with dm:
            if v["modified"]:
                for item in v["modified"]:
                    h(f"""<div class="diff-modified" style="padding:0.5rem 0.8rem;
                    margin-bottom:0.3rem;font-size:0.8rem;color:#F5C888;">
                    <span style="color:#F39C12;font-weight:700;margin-right:0.5rem;">~</span>{item}</div>""")
            else:
                h('<div style="font-size:0.8rem;color:#3A5570;padding:0.5rem 0;">No clauses modified in this version.</div>')

        # Compliance progression bar
        st.markdown("")
        section_title("Compliance Progression Across Versions", "📈")
        fig_prog = go.Figure()
        comp_vals = [vv["compliance"] for vv in versions]
        v_labels  = [vv["version"] for vv in versions]
        bar_colors = ["#27AE60" if c >= 75 else "#F39C12" if c >= 50 else "#E74C3C" for c in comp_vals]
        # Highlight selected
        bar_colors[selected_idx] = "#C9A84C"

        fig_prog.add_trace(go.Bar(
            x=v_labels, y=comp_vals,
            marker=dict(color=bar_colors, line=dict(color="#071219", width=2)),
            text=[f"{c}%" for c in comp_vals],
            textposition="outside",
            textfont=dict(color="#E8E4DA", size=11),
            hovertemplate="%{x}<br>Compliance: %{y}%<extra></extra>",
        ))
        fig_prog.add_hline(y=85, line_dash="dot", line_color="rgba(201,168,76,0.5)",
                           annotation_text="Target",
                           annotation_font_color="#C9A84C",
                           annotation_font_size=9)
        fig_prog.update_layout(
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            height=200, margin=dict(l=20, r=20, t=30, b=20),
            font=dict(color="#E8E4DA", size=10),
            xaxis=dict(gridcolor="rgba(0,0,0,0)", tickfont=dict(color="#C9A84C", size=11)),
            yaxis=dict(gridcolor="rgba(30,52,72,0.5)", range=[0, 105],
                       tickfont=dict(color="#607D99")),
            bargap=0.35,
        )
        st.plotly_chart(fig_prog, use_container_width=True, config={"displayModeBar": False})

    # ── Multi-version upload for live diffs ────────────────────────────────────
    st.markdown("")
    h('<div class="exec-divider"></div>')
    section_title("Upload Multiple Versions for Live Diff", "📂")
    h("""<div style="font-size:0.75rem;color:#607D99;margin-bottom:0.6rem;">
    Upload up to 4 PDF contract versions to generate a real-time diff analysis.
    The timeline above will update with your actual contract changes.</div>""")

    up_cols = st.columns(4)
    uploaded_v = []
    for i, col in enumerate(up_cols):
        with col:
            h(f'<div style="font-size:0.65rem;font-weight:700;letter-spacing:0.1em;'
              f'text-transform:uppercase;color:#607D99;margin-bottom:0.3rem;">v{i+1} Contract</div>')
            f = st.file_uploader(f"v{i+1}", type=["pdf"], key=f"vh_upload_{i}",
                                 label_visibility="collapsed")
            if f:
                txt, _ = extract_pdf_text(f)
                uploaded_v.append(txt)
                h(f"""<div style="font-size:0.68rem;color:#27AE60;margin-top:0.2rem;">
                ✓ Loaded ({len(txt):,} chars)</div>""")

    if len(uploaded_v) >= 2:
        st.info(f"✓ {len(uploaded_v)} versions loaded. Switch to a version above to see live diffs.")
        # Store for re-render
        st.session_state["vh_texts"] = uploaded_v


# ═══════════════════════════════════════════════════════════════════════════════
# PDF EXPORT (original + board report variant)
# ═══════════════════════════════════════════════════════════════════════════════

def generate_pdf(result: dict) -> bytes:
    """Original standard PDF export — preserved from original app."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_fill_color(13, 27, 42)
    pdf.rect(0, 0, 210, 30, "F")
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(201, 168, 76)
    pdf.cell(0, 10, "", ln=True)
    pdf.cell(0, 10, "Caledonian HR Group", ln=True)
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(180, 180, 180)
    pdf.cell(0, 5, "Employment Law & Contract Intelligence  |  Edinburgh, Scotland", ln=True)
    pdf.ln(5)
    pdf.set_draw_color(201, 168, 76)
    pdf.set_line_width(0.5)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(4)

    def sec(t):
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(13, 27, 42)
        pdf.set_fill_color(240, 234, 210)
        pdf.cell(0, 7, f"  {t.upper()}", ln=True, fill=True)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(40, 40, 40)
        pdf.ln(1)

    def kv(k, v):
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(90, 70, 20)
        pdf.cell(55, 6, f"{k}:", border=0)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(30, 30, 30)
        pdf.multi_cell(0, 6, str(v))

    sec("Document Classification")
    kv("Type", result.get("label", "N/A"))
    conf = result.get("confidence", 0)
    kv("Confidence", f"{conf:.2%}" if isinstance(conf, float) else str(conf))
    kv("Risk Score", f"{result.get('risk_score', 0)}/100")
    kv("Risk Level", result.get("risk_level", "N/A"))
    kv("Processing Time", f"{result.get('processing_time_ms', 0)} ms")
    pdf.ln(3)
    sec("Detected Clauses")
    for clause, present in result.get("clauses", {}).items():
        pdf.set_text_color(0, 120, 60) if present else pdf.set_text_color(150, 150, 150)
        pdf.cell(0, 5, f"  {'✓' if present else '✗'}  {clause}", ln=True)
    pdf.ln(3)
    sec("Key Insights")
    pdf.set_text_color(40, 40, 40)
    for item in result.get("insights", []):
        pdf.multi_cell(0, 5, f"   • {item}")
    pdf.ln(2)
    sec("Recommendations")
    for item in result.get("recommendations", []):
        pdf.multi_cell(0, 5, f"   • {item}")
    pdf.ln(2)
    if result.get("executive_summary"):
        sec("Executive Summary")
        for k, v in result["executive_summary"].items():
            kv(k.replace("_", " ").title(), str(v))
    pdf.ln(5)
    pdf.set_font("Helvetica", "I", 7)
    pdf.set_text_color(160, 160, 160)
    pdf.cell(0, 5, "Caledonian HR Group · 16 Charlotte Square, Edinburgh EH2 4DR · SCO 487 321", ln=True)
    return bytes(pdf.output())


def generate_board_pdf(result: dict, text: str, compliance_score: int, scottish_law: dict) -> bytes:
    """Feature 8 — Board-level PDF report (McKinsey/Deloitte style)."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Cover page header
    pdf.set_fill_color(7, 18, 25)
    pdf.rect(0, 0, 210, 297, "F")
    pdf.set_draw_color(201, 168, 76)
    pdf.set_line_width(1)
    pdf.line(15, 35, 195, 35)
    pdf.line(15, 262, 195, 262)

    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(201, 168, 76)
    pdf.set_y(45)
    pdf.cell(0, 12, "CALEDONIAN HR GROUP", ln=True, align="C")

    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(150, 170, 190)
    pdf.cell(0, 8, "Scottish Employment Contract Intelligence", ln=True, align="C")

    pdf.ln(15)
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(230, 228, 218)
    pdf.cell(0, 10, "BOARD LEVEL CONTRACT RISK REPORT", ln=True, align="C")

    pdf.ln(8)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(100, 125, 153)
    pdf.cell(0, 6, f"Prepared: {datetime.now().strftime('%d %B %Y')}", ln=True, align="C")
    pdf.cell(0, 6, "CONFIDENTIAL  |  Not for distribution", ln=True, align="C")

    # Page 2 — content
    pdf.add_page()
    pdf.set_fill_color(13, 31, 48)
    pdf.rect(0, 0, 210, 297, "F")

    def board_sec(title):
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_fill_color(201, 168, 76)
        pdf.set_text_color(7, 18, 25)
        pdf.cell(0, 7, f"  {title.upper()}", ln=True, fill=True)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(220, 220, 210)
        pdf.ln(1)

    def board_kv(k, v):
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(201, 168, 76)
        pdf.cell(60, 6, f"{k}:", border=0)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(200, 200, 190)
        pdf.multi_cell(0, 6, str(v))

    board_sec("Executive Summary")
    board_kv("Contract Type",    result.get("label", "N/A"))
    board_kv("Risk Score",       f"{result.get('risk_score', 0)} / 100")
    board_kv("Risk Level",       result.get("risk_level", "N/A"))
    board_kv("Compliance Score", f"{compliance_score}%")
    board_kv("Report Date",      datetime.now().strftime('%d %B %Y'))
    pdf.ln(3)

    board_sec("Compliance Status")
    for law, status in scottish_law.items():
        icon = {"compliant": "[OK]", "review": "[!]", "missing": "[X]"}.get(status, "[-]")
        pdf.set_text_color(
            39 if status=="compliant" else 231 if status=="missing" else 243,
            174 if status=="compliant" else 76 if status=="missing" else 156,
            96 if status=="compliant" else 60 if status=="missing" else 18
        )
        pdf.cell(0, 5, f"  {icon}  {law}", ln=True)
    pdf.ln(3)

    board_sec("Key Findings")
    pdf.set_text_color(200, 200, 190)
    for item in result.get("insights", []):
        pdf.multi_cell(0, 5, f"  • {item}")
    pdf.ln(2)

    board_sec("Recommendations")
    for item in result.get("recommendations", []):
        pdf.multi_cell(0, 5, f"  • {item}")

    pdf.ln(8)
    pdf.set_font("Helvetica", "I", 7)
    pdf.set_text_color(60, 80, 100)
    pdf.cell(0, 5, "Caledonian HR Group · 16 Charlotte Square, Edinburgh EH2 4DR · SCO 487 321 · CONFIDENTIAL", ln=True)
    return bytes(pdf.output())


# ═══════════════════════════════════════════════════════════════════════════════
# CORE RENDER RESULTS — Preserved + enhanced with new features
# ═══════════════════════════════════════════════════════════════════════════════

def render_results(res: dict, contract_text: str = ""):
    label       = res.get("label", "N/A")
    confidence  = res.get("confidence", 0.0)
    risk_score  = res.get("risk_score", 0)
    risk_level  = res.get("risk_level", "Low")
    clauses     = res.get("clauses", {})
    insights    = res.get("insights", [])
    impact      = res.get("business_impact", [])
    recs        = res.get("recommendations", [])
    exec_sum    = res.get("executive_summary", {})
    entities    = res.get("entities", [])
    explanation = res.get("explanation", [])
    probs       = res.get("probabilities", {})
    top_preds   = res.get("top_predictions", [])
    proc_ms     = res.get("processing_time_ms", 0)

    rc = risk_config(risk_level)

    # Success banner
    h(f"""
    <div style="background:rgba(39,174,96,0.07);border:1px solid rgba(39,174,96,0.25);
    border-left:4px solid #27AE60;border-radius:0 8px 8px 0;
    padding:0.65rem 1rem;margin-bottom:1.2rem;font-size:0.83rem;color:#7DD8A0;
    display:flex;align-items:center;gap:0.6rem;">
        ✅ &nbsp; Contract analysis complete
        <span style="margin-left:auto;color:#607D99;font-size:0.75rem;">
        Processed in {proc_ms} ms</span>
    </div>""")

    # 4-metric row
    m1, m2, m3, m4 = st.columns(4)
    with m1: st.metric("Document Type", label)
    with m2:
        conf_str = f"{confidence:.1%}" if isinstance(confidence, float) else str(confidence)
        st.metric("Confidence", conf_str)
    with m3: st.metric("Risk Score", f"{risk_score}/100")
    with m4: st.metric("Risk Level", risk_level)

    # Risk banner
    h(f"""
    <div style="background:{rc['bg']};border:1px solid {rc['border']};
    border-left:5px solid {rc['color']};border-radius:0 8px 8px 0;
    padding:0.9rem 1.4rem;margin:1rem 0;display:flex;align-items:center;gap:1rem;">
        <span style="font-size:1.6rem;">{rc['icon']}</span>
        <div>
            <div style="font-size:0.68rem;letter-spacing:0.12em;text-transform:uppercase;
            color:#607D99;margin-bottom:0.15rem;">Risk Assessment</div>
            <div style="font-size:1.1rem;font-weight:700;color:{rc['color']};">
            {risk_level} Risk Contract</div>
        </div>
        <div style="margin-left:auto;text-align:right;">
            <div style="font-family:'Cormorant Garamond',serif;font-size:2.8rem;
            font-weight:700;color:{rc['color']};line-height:1;">{risk_score}</div>
            <div style="font-size:0.7rem;color:#607D99;">out of 100</div>
        </div>
    </div>""")

    # ── All tabs including new features
    (t_clauses, t_insights, t_compliance, t_scot, t_risks,
     t_summary, t_entities, t_model, t_highlight) = st.tabs([
        "  Clauses  ", "  Insights & Impact  ", "  UK Compliance  ",
        "  Scottish Law  ", "  Risk Gauges  ", "  Executive Summary  ",
        "  Entities & Explanation  ", "  Model Output  ", "  Clause Viewer  ",
    ])

    # ── Clauses tab (original)
    with t_clauses:
        section_title("Legal Clause Detection", "📋")
        present_clauses = {k: v for k, v in clauses.items() if v}
        absent_clauses  = {k: v for k, v in clauses.items() if not v}
        col_p, col_a = st.columns(2)
        with col_p:
            h(f'<div style="font-size:0.72rem;letter-spacing:0.1em;text-transform:uppercase;color:#27AE60;margin-bottom:0.5rem;font-weight:600;">✓ &nbsp;Detected ({len(present_clauses)})</div>')
            for clause in present_clauses:
                h(f'<div style="display:flex;align-items:center;gap:0.6rem;padding:0.5rem 0.7rem;margin-bottom:0.3rem;background:rgba(39,174,96,0.07);border:1px solid rgba(39,174,96,0.2);border-left:3px solid #27AE60;border-radius:0 7px 7px 0;font-size:0.83rem;color:#D4EDE0;"><span style="color:#27AE60;font-weight:700;">✓</span> {clause}</div>')
        with col_a:
            h(f'<div style="font-size:0.72rem;letter-spacing:0.1em;text-transform:uppercase;color:#607D99;margin-bottom:0.5rem;font-weight:600;">✗ &nbsp;Not Detected ({len(absent_clauses)})</div>')
            for clause in absent_clauses:
                h(f'<div style="display:flex;align-items:center;gap:0.6rem;padding:0.5rem 0.7rem;margin-bottom:0.3rem;background:rgba(30,52,72,0.3);border:1px solid rgba(30,52,72,0.6);border-left:3px solid #1E3448;border-radius:0 7px 7px 0;font-size:0.83rem;color:#607D99;"><span style="color:#2E4A60;">✗</span> {clause}</div>')

    # ── Insights tab (original)
    with t_insights:
        ic, bc, rc_col = st.columns([1, 1, 1])
        with ic:
            section_title("Key Insights", "💡")
            for item in insights:
                h(f'<div style="padding:0.5rem 0.8rem;margin-bottom:0.35rem;background:rgba(201,168,76,0.06);border-left:3px solid #C9A84C;border-radius:0 7px 7px 0;font-size:0.82rem;color:#E8E4DA;line-height:1.5;">{item}</div>')
            if not insights: st.info("No insights available.")
        with bc:
            section_title("Business Impact", "📊")
            for item in impact:
                h(f'<div style="padding:0.5rem 0.8rem;margin-bottom:0.35rem;background:rgba(243,156,18,0.06);border-left:3px solid #F39C12;border-radius:0 7px 7px 0;font-size:0.82rem;color:#E8E4DA;line-height:1.5;">{item}</div>')
            if not impact: st.info("No impact notes.")
        with rc_col:
            section_title("Recommendations", "✅")
            for i, item in enumerate(recs, 1):
                h(f'<div style="padding:0.5rem 0.8rem;margin-bottom:0.35rem;background:rgba(39,174,96,0.06);border-left:3px solid #27AE60;border-radius:0 7px 7px 0;font-size:0.82rem;color:#E8E4DA;line-height:1.5;"><b style="color:#C9A84C;">{i}.</b> {item}</div>')
            if not recs: st.info("No recommendations.")

    # ── Feature 1: UK Compliance tab
    with t_compliance:
        comp_score, comp_checklist = compute_compliance_score(clauses)
        render_compliance_gauge(comp_score, comp_checklist)

    # ── Feature 2: Scottish Law tab
    with t_scot:
        render_scottish_law_checker(contract_text)

    # ── Feature 3: Risk Gauges tab
    with t_risks:
        render_risk_gauges(res)

    # ── Executive Summary tab (original)
    with t_summary:
        sc, pc = st.columns([1, 1])
        with sc:
            section_title("Executive Summary", "📑")
            for k, v in exec_sum.items():
                kv_row(k.replace("_", " ").title(), str(v))
        with pc:
            section_title("Classification Breakdown", "🎯")
            if top_preds:
                for item in top_preds:
                    lbl = item.get("label", "N/A")
                    try:
                        p = float(item.get("confidence", 0))
                        bar = progress_bar(p)
                        ps = f"{p:.2%}"
                    except Exception:
                        bar, ps = "", str(item.get("confidence", ""))
                    h(f'<div style="margin-bottom:0.7rem;"><div style="display:flex;justify-content:space-between;font-size:0.8rem;margin-bottom:0.1rem;"><span style="color:#E8E4DA;">{lbl}</span><span style="color:#C9A84C;font-weight:600;">{ps}</span></div>{bar}</div>')

    # ── Entities & Explanation tab (original)
    with t_entities:
        ent_col, exp_col = st.columns(2)
        with ent_col:
            section_title("Named Entities", "🏷️")
            label_colors = {
                "ORG": "#4A90D9", "PERSON": "#9B59B6", "DATE": "#27AE60",
                "MONEY": "#C9A84C", "LAW": "#E67E22", "GPE": "#16A085",
                "LOC": "#16A085", "NORP": "#8E44AD", "EVENT": "#E91E63",
            }
            if entities:
                chips = "".join(
                    f'<span style="display:inline-flex;align-items:center;gap:0.35rem;background:#0D1F30;border:1px solid #1E3448;border-radius:20px;padding:0.22rem 0.7rem;margin:0.2rem;font-size:0.78rem;"><span style="color:#E8E4DA;">{e.get("text","")}</span><span style="background:{label_colors.get(e.get("label",""),"#607D99")};color:#fff;border-radius:10px;padding:0.05rem 0.42rem;font-size:0.6rem;font-weight:700;letter-spacing:0.04em;">{e.get("label","")}</span></span>'
                    for e in entities
                )
                h(f'<div style="line-height:2.2;">{chips}</div>')
            else:
                st.info("No entities detected.")
        with exp_col:
            section_title("Token Saliency", "📈")
            if explanation:
                max_s = max((e.get("score", 0) for e in explanation), default=1) or 1
                for item in explanation[:10]:
                    word = item.get("word", "")
                    score = item.get("score", 0)
                    p = score / max_s
                    h(f'<div style="display:grid;grid-template-columns:100px 1fr 40px;align-items:center;gap:0.7rem;margin-bottom:0.4rem;font-size:0.8rem;"><span style="color:#E8E4DA;font-weight:500;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">{word}</span><div style="background:#1E3448;border-radius:3px;height:5px;"><div style="width:{int(p*100)}%;background:#C9A84C;height:100%;border-radius:3px;"></div></div><span style="color:#607D99;text-align:right;">{score:.2f}</span></div>')
            else:
                st.info("No explanation data.")

    # ── Model Output tab (original)
    with t_model:
        section_title("Probability Distribution", "📊")
        if probs:
            for cls_lbl, prob_val in sorted(probs.items(), key=lambda x: x[1], reverse=True):
                try:
                    p = float(prob_val)
                    bar = progress_bar(p, "#4A90D9", 8)
                    ps = f"{p:.3%}"
                except Exception:
                    bar, ps = "", str(prob_val)
                h(f'<div style="margin-bottom:0.65rem;"><div style="display:flex;justify-content:space-between;font-size:0.82rem;margin-bottom:0.2rem;"><span style="color:#E8E4DA;">{cls_lbl}</span><span style="color:#4A90D9;font-weight:600;">{ps}</span></div>{bar}</div>')
        else:
            st.info("No probability data.")

    # ── Feature 9: Clause Viewer tab
    with t_highlight:
        render_clause_highlighter(contract_text)


# ═══════════════════════════════════════════════════════════════════════════════
# SIDEBAR (preserved + enhanced)
# ═══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    # ── Premium Glassmorphism Sidebar Logo
    h("""
    <div class="sidebar-logo-area">
        <div style="position:relative;z-index:1;">
            <div style="font-size:2rem;margin-bottom:0.5rem;
            filter:drop-shadow(0 0 12px rgba(201,168,76,0.7));
            animation:glowPulse 3s infinite;">⚖️</div>
            <div style="font-family:'Cormorant Garamond',serif;font-size:1.15rem;
            font-weight:700;color:#E8E4DA;letter-spacing:0.03em;
            text-shadow:0 2px 8px rgba(0,0,0,0.5);">Caledonian HR Group</div>
            <div style="font-size:0.55rem;letter-spacing:0.2em;text-transform:uppercase;
            color:#C9A84C;margin-top:0.2rem;font-weight:600;">
            Scottish Workforce Intelligence</div>
            <div style="font-size:0.56rem;letter-spacing:0.08em;
            color:#3A5570;margin-top:0.12rem;">Employment Risk &amp; Contract Analytics</div>
            <div style="display:flex;justify-content:center;gap:0.35rem;
            margin-top:0.8rem;flex-wrap:wrap;">
                <span style="background:rgba(201,168,76,0.1);
                border:1px solid rgba(201,168,76,0.25);
                border-radius:10px;padding:0.12rem 0.55rem;
                font-size:0.56rem;color:#C9A84C;">🏴󠁧󠁢󠁳󠁣󠁴󠁿 Scotland</span>
                <span style="background:rgba(39,174,96,0.1);
                border:1px solid rgba(39,174,96,0.25);
                border-radius:10px;padding:0.12rem 0.55rem;
                font-size:0.56rem;color:#27AE60;">⚡ AI</span>
                <span style="background:rgba(74,144,217,0.1);
                border:1px solid rgba(74,144,217,0.25);
                border-radius:10px;padding:0.12rem 0.55rem;
                font-size:0.56rem;color:#4A90D9;">🔒 GDPR</span>
            </div>
        </div>
    </div>""")

    # Stats
    stats = fetch_stats(API_URL)
    docs  = stats.get("docs_analyzed")
    flags = stats.get("high_risk_flags")

    flags_color = "#E74C3C" if isinstance(flags, int) and flags > 0 else "#C9A84C"
    h(f"""
    <div style="display:grid;grid-template-columns:1fr 1fr;gap:0.6rem;margin-bottom:1rem;">
        <div class="exec-kpi-card" style="border-top-color:#C9A84C;padding:0.8rem;text-align:center;">
            <div class="exec-kpi-label" style="font-size:0.55rem;">📁 Reviewed</div>
            <div class="exec-kpi-value" style="font-size:1.5rem;color:#C9A84C;">
            {docs if docs is not None else '—'}</div>
        </div>
        <div class="exec-kpi-card" style="border-top-color:{flags_color};padding:0.8rem;text-align:center;">
            <div class="exec-kpi-label" style="font-size:0.55rem;">⚠️ High Risk</div>
            <div class="exec-kpi-value" style="font-size:1.5rem;color:{flags_color};">
            {flags if flags is not None else '—'}</div>
        </div>
    </div>""")

    # API Status
    section_title("System Status", "🔌")
    api_ok, _ = check_api_health(API_URL)
    status_color = "#27AE60" if api_ok else "#E74C3C"
    status_label = "Online" if api_ok else "Offline"
    h(f"""
    <div style="display:flex;align-items:center;justify-content:space-between;
    background:#0D1F30;border:1px solid #1E3448;border-radius:7px;
    padding:0.55rem 0.8rem;font-size:0.8rem;margin-bottom:0.9rem;">
        <span style="color:#607D99;">API Server</span>
        <span style="display:flex;align-items:center;gap:0.4rem;color:{status_color};font-weight:600;">
            <span style="width:7px;height:7px;background:{status_color};border-radius:50%;
            display:inline-block;box-shadow:0 0 5px {status_color};"></span>
            {status_label}
        </span>
    </div>""")

    # Trusted by branding
    h("""
    <div style="background:rgba(201,168,76,0.04);border:1px solid rgba(201,168,76,0.15);
    border-radius:8px;padding:0.7rem;margin-bottom:1rem;text-align:center;">
        <div style="font-size:0.6rem;color:#3A5570;letter-spacing:0.1em;text-transform:uppercase;
        margin-bottom:0.3rem;">Trusted by Scottish Employers</div>
        <div style="font-size:0.65rem;color:#607D99;line-height:1.6;">
        Employment Law Intelligence<br>AI-Powered Contract Analytics
        </div>
    </div>""")

    # History
    section_title("Recent Contracts", "📁")
    history = fetch_history(API_URL)
    if history:
        for rec in history[:6]:
            rs = rec.get("risk_score", 0)
            rc = "#E74C3C" if rs >= 70 else "#F39C12" if rs >= 40 else "#27AE60"
            ts = rec.get("timestamp", "")[:16].replace("T", " ")
            preview = rec.get("preview", "")[:48]
            h(f"""
            <div style="background:#0D1F30;border:1px solid #1E3448;border-radius:7px;
            padding:0.55rem 0.75rem;margin-bottom:0.4rem;">
                <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:0.25rem;">
                    <span style="color:#C9A84C;font-size:0.78rem;font-weight:600;">
                    {rec.get('label','Unknown')}</span>
                    <span style="background:rgba({_hex_to_rgb(rc)},0.15);
                    border:1px solid rgba({_hex_to_rgb(rc)},0.4);border-radius:4px;
                    padding:0.05rem 0.45rem;font-size:0.65rem;font-weight:700;color:{rc};">
                    Risk {rs}</span>
                </div>
                <div style="font-size:0.7rem;color:#607D99;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">
                {preview}…</div>
                <div style="font-size:0.65rem;color:#3A5570;margin-top:0.2rem;">{ts}</div>
            </div>""")
    else:
        h('<div style="font-size:0.78rem;color:#607D99;padding:0.5rem 0;font-style:italic;">No contracts reviewed yet.</div>')

    # Footer
    h("""
    <div style="margin-top:1.5rem;padding:1rem;
    background:linear-gradient(135deg,rgba(7,18,25,0.9),rgba(13,31,48,0.8));
    border:1px solid rgba(201,168,76,0.15);border-radius:10px;
    font-size:0.66rem;color:#3A5570;line-height:1.8;text-align:center;">
        <div style="color:#C9A84C;font-weight:700;margin-bottom:0.4rem;
        font-family:'Cormorant Garamond',serif;font-size:0.85rem;
        letter-spacing:0.05em;">⚖️ Caledonian HR Group</div>
        <div style="color:#607D99;">16 Charlotte Square<br>Edinburgh EH2 4DR<br>Scotland</div>
        <div style="margin:0.4rem 0;height:1px;
        background:linear-gradient(90deg,transparent,rgba(201,168,76,0.25),transparent);"></div>
        <div style="color:#3A5570;font-size:0.6rem;">SCO 487 321 &nbsp;·&nbsp; Reg. Scotland</div>
        <div style="margin-top:0.4rem;font-size:0.58rem;color:#2A3F52;">
        Harvey AI / Deloitte calibre legal intelligence<br>
        Powered by BERT · FAISS · NLI
        </div>
    </div>""")


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN LAYOUT — Premium Header (Feature 10 branding)
# ═══════════════════════════════════════════════════════════════════════════════

h(f"""
<div class="exec-dashboard-header" style="margin-bottom:1.5rem;">
    <div style="display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:1rem;">
        <div>
            <div style="font-size:0.58rem;letter-spacing:0.24em;text-transform:uppercase;
            color:#3A5570;margin-bottom:0.35rem;font-weight:600;">
            SCOTTISH WORKFORCE COMPLIANCE PLATFORM &nbsp;·&nbsp; ENTERPRISE EDITION</div>
            <div style="font-family:'Cormorant Garamond',serif;font-size:2.1rem;
            font-weight:700;color:#E8E4DA;letter-spacing:0.01em;line-height:1.1;">
                Contract Risk Intelligence
            </div>
            <div style="font-size:0.75rem;letter-spacing:0.06em;
            color:#607D99;margin-top:0.4rem;">
                Caledonian HR Group &nbsp;·&nbsp; Edinburgh &nbsp;·&nbsp;
                Employment Law &amp; Contract Analytics
            </div>
        </div>
        <div style="display:flex;flex-direction:column;align-items:flex-end;gap:0.5rem;">
            <div style="display:flex;gap:0.4rem;flex-wrap:wrap;justify-content:flex-end;">
                <span style="background:rgba(201,168,76,0.1);border:1px solid rgba(201,168,76,0.3);
                border-radius:20px;padding:0.22rem 0.9rem;font-size:0.65rem;color:#C9A84C;
                letter-spacing:0.05em;">🏴󠁧󠁢󠁳󠁣󠁴󠁿 Registered Scotland</span>
                <span style="background:rgba(39,174,96,0.1);border:1px solid rgba(39,174,96,0.3);
                border-radius:20px;padding:0.22rem 0.9rem;font-size:0.65rem;color:#27AE60;
                letter-spacing:0.05em;">⚡ AI Powered</span>
                <span style="background:rgba(74,144,217,0.1);border:1px solid rgba(74,144,217,0.3);
                border-radius:20px;padding:0.22rem 0.9rem;font-size:0.65rem;color:#4A90D9;
                letter-spacing:0.05em;">🔒 GDPR Compliant</span>
            </div>
            <div style="font-size:0.6rem;color:#3A5570;letter-spacing:0.08em;text-align:right;">
            Trusted by Scottish Employers &nbsp;·&nbsp; {datetime.now().strftime('%d %b %Y')}
            </div>
        </div>
    </div>
    <div style="margin-top:1.2rem;display:flex;gap:0.6rem;flex-wrap:wrap;">
        <span style="background:rgba(201,168,76,0.05);border:1px solid rgba(201,168,76,0.15);
        border-radius:6px;padding:0.2rem 0.7rem;font-size:0.65rem;color:#C9A84C;">
        BERT Classification</span>
        <span style="background:rgba(201,168,76,0.05);border:1px solid rgba(201,168,76,0.15);
        border-radius:6px;padding:0.2rem 0.7rem;font-size:0.65rem;color:#C9A84C;">
        NER Extraction</span>
        <span style="background:rgba(201,168,76,0.05);border:1px solid rgba(201,168,76,0.15);
        border-radius:6px;padding:0.2rem 0.7rem;font-size:0.65rem;color:#C9A84C;">
        Zero-shot NLI</span>
        <span style="background:rgba(201,168,76,0.05);border:1px solid rgba(201,168,76,0.15);
        border-radius:6px;padding:0.2rem 0.7rem;font-size:0.65rem;color:#C9A84C;">
        RAG Vector Search</span>
        <span style="background:rgba(201,168,76,0.05);border:1px solid rgba(201,168,76,0.15);
        border-radius:6px;padding:0.2rem 0.7rem;font-size:0.65rem;color:#C9A84C;">
        Scottish Law Engine</span>
        <span style="background:rgba(201,168,76,0.05);border:1px solid rgba(201,168,76,0.15);
        border-radius:6px;padding:0.2rem 0.7rem;font-size:0.65rem;color:#C9A84C;">
        Tribunal Risk AI</span>
    </div>
</div>""")

# ═══════════════════════════════════════════════════════════════════════════════
# MAIN CONTENT — Three-tab layout: Analyse | Advanced Tools | Board Report
# ═══════════════════════════════════════════════════════════════════════════════

main_tab1, main_tab2, main_tab3, main_tab4, main_tab5, main_tab6, main_tab7, main_tab8, main_tab9, main_tab10, main_tab11, main_tab12, main_tab13, main_tab14, main_tab15 = st.tabs([
    "  ⚖️ Contract Analysis  ",
    "  💰 Compensation Intelligence  ",
    "  🔄 Contract Comparison  ",
    "  👥 Recruitment Intelligence  ",
    "  📋 Board Report  ",
    "  🤖 RAG Assistant  ",
    "  🔍 Clause Highlighter  ",
    "  📂 Portfolio Analysis  ",
    "  🌡️ Workforce Heatmap  ",
    "  📊 Clause Benchmarking  ",
    "  🎯 Offer Predictor  ",
    "  ⚖️ Tribunal Risk  ",
    "  📝 Legal Memo  ",
    "  📊 HR Dashboard  ",
    "  🕐 Version History  ",
])

with main_tab1:
    left_col, right_col = st.columns([1.05, 1.45], gap="large")

    with left_col:
        h('<div style="font-size:0.7rem;font-weight:700;letter-spacing:0.12em;text-transform:uppercase;color:#607D99;margin-bottom:0.5rem;">📎 Upload Contract (PDF)</div>')
        uploaded = st.file_uploader("Upload PDF", type=["pdf"], label_visibility="collapsed", key="main_upload")

        extracted = ""
        if uploaded:
            extracted, err = extract_pdf_text(uploaded)
            if err:
                st.error(f"PDF extraction failed: {err}")
            elif extracted:
                char_count = f"{len(extracted):,}"
                h(f"""<div style="background:rgba(39,174,96,0.07);border:1px solid rgba(39,174,96,0.25);
                border-left:3px solid #27AE60;border-radius:0 6px 6px 0;
                padding:0.45rem 0.8rem;margin-bottom:0.5rem;font-size:0.8rem;color:#7DD8A0;">
                ✓ &nbsp;Extracted {char_count} characters</div>""")
            else:
                st.warning("PDF uploaded but no readable text found.")

        h('<div style="font-size:0.7rem;font-weight:700;letter-spacing:0.12em;text-transform:uppercase;color:#607D99;margin:0.8rem 0 0.5rem 0;">📝 Contract Text</div>')
        text = st.text_area(
            "text",
            value=extracted,
            height=300,
            label_visibility="collapsed",
            placeholder="Paste employment contract, NDA, service agreement, or any legal document text here…",
        )

        # Store text in session for cross-tab use
        if text.strip():
            st.session_state["contract_text"] = text.strip()

        b1, b2 = st.columns([3, 2])
        with b1:
            run = st.button("⚖️  Analyse Contract", use_container_width=True)
        with b2:
            if st.session_state.get("last_result"):
                pdf_bytes = generate_pdf(st.session_state["last_result"])
                st.download_button("📄 Export PDF", data=pdf_bytes,
                                   file_name="caledonian_hr_report.pdf",
                                   mime="application/pdf", use_container_width=True)

        if run:
            if not text.strip():
                st.error("Please provide contract text before running analysis.")
            else:
                with st.spinner("Running contract intelligence analysis…"):
                    result, err = call_api(API_URL, text)
                if err:
                    st.error(err)
                    st.warning("Backend may be cold-starting on Render. Wait 10 seconds and retry.")
                else:
                    st.session_state["last_result"] = result

        # Pipeline info (idle state)
        if not st.session_state.get("last_result"):
            h("""
            <div style="background:#0D1F30;border:1px solid #1E3448;border-radius:10px;
            padding:1.1rem 1.2rem;margin-top:1rem;">
                <div style="font-size:0.68rem;font-weight:700;letter-spacing:0.14em;
                text-transform:uppercase;color:#607D99;margin-bottom:0.8rem;">Analysis Pipeline</div>
                <div style="display:flex;flex-wrap:wrap;gap:0.35rem;align-items:center;">
                    <span style="background:rgba(201,168,76,0.08);border:1px solid rgba(201,168,76,0.2);border-radius:5px;padding:0.2rem 0.6rem;font-size:0.7rem;color:#C9A84C;">PDF Intake</span>
                    <span style="color:#3A5570;font-size:0.7rem;">→</span>
                    <span style="background:rgba(201,168,76,0.08);border:1px solid rgba(201,168,76,0.2);border-radius:5px;padding:0.2rem 0.6rem;font-size:0.7rem;color:#C9A84C;">BERT Classification</span>
                    <span style="color:#3A5570;font-size:0.7rem;">→</span>
                    <span style="background:rgba(201,168,76,0.08);border:1px solid rgba(201,168,76,0.2);border-radius:5px;padding:0.2rem 0.6rem;font-size:0.7rem;color:#C9A84C;">NER Extraction</span>
                    <span style="color:#3A5570;font-size:0.7rem;">→</span>
                    <span style="background:rgba(201,168,76,0.08);border:1px solid rgba(201,168,76,0.2);border-radius:5px;padding:0.2rem 0.6rem;font-size:0.7rem;color:#C9A84C;">Zero-shot Clauses</span>
                    <span style="color:#3A5570;font-size:0.7rem;">→</span>
                    <span style="background:rgba(201,168,76,0.08);border:1px solid rgba(201,168,76,0.2);border-radius:5px;padding:0.2rem 0.6rem;font-size:0.7rem;color:#C9A84C;">Risk Scoring</span>
                    <span style="color:#3A5570;font-size:0.7rem;">→</span>
                    <span style="background:rgba(201,168,76,0.08);border:1px solid rgba(201,168,76,0.2);border-radius:5px;padding:0.2rem 0.6rem;font-size:0.7rem;color:#C9A84C;">Insights</span>
                </div>
                <div style="margin-top:0.9rem;font-size:0.74rem;color:#3A5570;line-height:1.6;">
                    Powered by BERT · Zero-shot NLI · Gradient saliency · Scottish Employment Law aligned
                </div>
            </div>""")

    with right_col:
        # ── Feature 7 — Chat widget alongside results
        chat_toggle = st.toggle("💬 Contract AI Assistant", value=False)

        if chat_toggle:
            ct = st.session_state.get("contract_text", text)
            render_chat_assistant(ct)
        else:
            h('<div style="font-size:0.7rem;font-weight:700;letter-spacing:0.12em;text-transform:uppercase;color:#607D99;margin-bottom:0.5rem;">👁 Analysis Results</div>')
            if st.session_state.get("last_result"):
                ct = st.session_state.get("contract_text", text)
                render_results(st.session_state["last_result"], ct)
            else:
                preview = text.strip() if text.strip() else (
                    "Upload or paste a contract to begin analysis.\n\n"
                    "Caledonian HR Group specialises in:\n\n"
                    "  •  Employment Contracts\n"
                    "  •  Non-Disclosure Agreements\n"
                    "  •  Service & Vendor Agreements\n"
                    "  •  Settlement Agreements\n"
                    "  •  Scottish Jurisdiction Clauses\n"
                    "  •  TUPE Transfer Documents\n"
                    "  •  Restrictive Covenant Clauses"
                )
                st.text_area("Preview", value=preview, height=560, label_visibility="collapsed")

# ── Feature 4: Compensation Intelligence Tab
with main_tab2:
    ct = st.session_state.get("contract_text", "")
    render_compensation_intelligence(ct)

# ── Feature 5: Contract Comparison Tab
with main_tab3:
    render_contract_comparison()

# ── Feature 6: Recruitment Intelligence Tab
with main_tab4:
    ct = st.session_state.get("contract_text", "")
    render_recruitment_intelligence(ct)

# ── Feature 8: Board Report Tab
with main_tab5:
    if st.session_state.get("last_result"):
        ct = st.session_state.get("contract_text", "")
        render_board_report(st.session_state["last_result"], ct)
    else:
        h("""<div style="background:#0D1F30;border:1px solid #1E3448;border-radius:12px;
        padding:3rem;text-align:center;">
            <div style="font-size:2rem;margin-bottom:0.8rem;">📋</div>
            <div style="font-family:'Cormorant Garamond',serif;font-size:1.2rem;
            color:#E8E4DA;margin-bottom:0.5rem;">Board Report Awaiting Analysis</div>
            <div style="font-size:0.82rem;color:#607D99;">
            Run a contract analysis first to generate the executive board report.
            </div>
        </div>""")

# ── Feature 11: Enterprise RAG Contract Assistant
with main_tab6:
    ct = st.session_state.get("contract_text", "")
    render_rag_assistant(ct)

# ── Feature 12: Enhanced Clause Highlighter
with main_tab7:
    ct = st.session_state.get("contract_text", "")
    render_enhanced_clause_highlighter(ct)

# ── Feature 13: Multi-Document Portfolio Analysis
with main_tab8:
    render_portfolio_analysis()

# ── Feature 14: Workforce Risk Heatmap
with main_tab9:
    render_workforce_heatmap()

# ── Feature 15: Clause Benchmarking Engine
with main_tab10:
    ct = st.session_state.get("contract_text", "")
    render_clause_benchmarking(ct)

# ── Feature 16: Offer Acceptance Predictor
with main_tab11:
    render_offer_acceptance_predictor()

# ── Feature 17: Employment Tribunal Risk Predictor
with main_tab12:
    ct = st.session_state.get("contract_text", "")
    res = st.session_state.get("last_result")
    render_tribunal_risk_predictor(ct, res)

# ── Feature 18: Legal Memo Generator
with main_tab13:
    ct = st.session_state.get("contract_text", "")
    res = st.session_state.get("last_result")
    render_legal_memo_tab(ct, res)

# ── New Feature A: HR Director Executive Dashboard
with main_tab14:
    render_hr_director_dashboard()

# ── New Feature B: Contract Version History
with main_tab15:
    vh_texts = st.session_state.get("vh_texts", None)
    render_version_history(uploaded_texts=vh_texts)